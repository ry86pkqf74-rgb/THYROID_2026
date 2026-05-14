# -*- coding: utf-8 -*-
"""Build the BigQuery Studio Feature Integration Plan as a .docx (python-docx)."""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x53, 0x95)
GREY = RGBColor(0x55, 0x55, 0x55)
LGREY = RGBColor(0x77, 0x77, 0x77)
HEADER_FILL = "1F3864"
ALT_FILL = "EEF2F8"
CALLOUT_FILL = "E2EFDA"
GREEN = RGBColor(0x37, 0x56, 0x23)

doc = Document()

# ---- base style ----
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ---- page setup ----
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))

# header / footer
hdr = sec.header.paragraphs[0]
r = hdr.add_run("BigQuery Studio Feature Integration Plan  —  thyroid-canonical-pub-2026")
r.font.size = Pt(8); r.font.color.rgb = LGREY
ftr = sec.footer.paragraphs[0]
r = ftr.add_run("Draft for team review  ·  May 14, 2026")
r.font.size = Pt(8); r.font.color.rgb = LGREY


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 11)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15 if level == 1 else 12.5)
    run.font.color.rgb = NAVY if level == 1 else BLUE
    p.style = doc.styles["Heading %d" % level]
    # re-apply (Heading style may override)
    for rr in p.runs:
        rr.bold = True
        rr.font.size = Pt(15 if level == 1 else 12.5)
        rr.font.color.rgb = NAVY if level == 1 else BLUE
    return p


def para(text="", bold=False, italic=False, size=11, color=None, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
    return p


def rich(parts, after=6, before=0):
    """parts: list of (text, {bold,italic,color,size})"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    for item in parts:
        if isinstance(item, str):
            text, opt = item, {}
        else:
            text, opt = item
        run = p.add_run(text)
        run.bold = opt.get("bold", False)
        run.italic = opt.get("italic", False)
        run.font.size = Pt(opt.get("size", 11))
        if "color" in opt:
            run.font.color.rgb = opt["color"]
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, list):
            for item in it:
                text, opt = (item, {}) if isinstance(item, str) else item
                run = p.add_run(text)
                run.bold = opt.get("bold", False)
                run.font.size = Pt(11)
        else:
            p.add_run(it).font.size = Pt(11)


def numbered(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(it).font.size = Pt(11)


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr_cells[i], HEADER_FILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if ri % 2 == 1:
                _set_cell_bg(cells[ci], ALT_FILL)
    if widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(widths):
                t.rows[ri].cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def callout(title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _set_cell_bg(cell, CALLOUT_FILL)
    cell.text = ""
    p0 = cell.paragraphs[0]
    r0 = p0.add_run(title)
    r0.bold = True
    r0.font.size = Pt(11)
    r0.font.color.rgb = GREEN
    for ln in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if isinstance(ln, list):
            for item in ln:
                text, opt = (item, {}) if isinstance(item, str) else item
                run = p.add_run(text)
                run.bold = opt.get("bold", False)
                run.font.size = Pt(10)
        else:
            p.add_run(ln).font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ============================ TITLE ============================
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("BigQuery Studio Feature Integration Plan")
r.bold = True; r.font.size = Pt(23); r.font.color.rgb = NAVY
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Reducing data-quality, lineage, and legacy-migration friction in the Thyroid Canonical Publication database")
r.font.size = Pt(13); r.font.color.rgb = GREY

rich([("Project: ", {"bold": True}), "thyroid-canonical-pub-2026     ",
      ("Prepared for: ", {"bold": True}), "the research data team     ",
      ("Date: ", {"bold": True}), "May 14, 2026"], after=3)
rich([("Status: ", {"bold": True}), "Draft for team review — first deliverables already executed (see Section 2)"], after=3)
rich([("Scope: ", {"bold": True}), "BigQuery Studio (Data Preparation, Data Canvas, Notebooks, Visualizations, Apache Spark, "
      "Pipelines, custom Agents, Gemini) plus the broader Google Cloud analytics ecosystem"], after=10)

# ============================ 1. EXEC SUMMARY ============================
heading("1. Executive summary")
para("Our BigQuery project is large, mature, and carefully governed — roughly 10,871 patients spread across nine "
     "datasets and several hundred tables, with a real sign-off registry, a query-cost log, and a snapshot-everything "
     "discipline. The problem is not that the data is bad. The problem is that almost all of our quality control happens "
     "after the fact: we discover a broken link, a disagreeing column, or a builder that did not re-run only once it has "
     "already propagated into a cohort or a manuscript number, and then we spend days backtracking.")
para("BigQuery Studio — together with the wider Google Cloud analytics stack — now ships a set of features that "
     "let us move that quality control to the front of the process and make the whole dataset legible to the team. This "
     "plan maps each feature to a specific pain point we are actually hitting, records the first deliverables that have "
     "already been built, proposes a new “global evaluation layer” of patient-level census views, evaluates the "
     "broader GCP toolset, and lays out a phased rollout.")
para("The single highest-leverage change is to stop treating our QC rules as a script we run when we remember to, and "
     "instead make them a Pipeline gate that canonical tables must pass before they are published. Everything else in "
     "this plan supports or extends that idea.", bold=True)

# ============================ 2. WHAT HAS BEEN EXECUTED ============================
heading("2. What has already been executed")
para("This is a working plan, not just a proposal. The directly-executable piece — the global evaluation layer "
     "— has already been built in the live project, and the data issues it surfaced have been logged. The rest of "
     "this section records what is in place as of May 14, 2026.")
callout("Live in the project today", [
    [("New dataset ", {}), ("pub_eval", {"bold": True}), (" — isolated, views only, no underlying data mutated.", {})],
    [("vw_patient_workup_census_v1", {"bold": True}),
     (" — one row per patient (10,871). Preop and postop performed-flags and intervals for ultrasound, CT, MRI, FNA, "
      "and nuclear medicine; reoperation signals; and a prior-thyroid-procedure pathology-gap review flag.", {})],
    [("vw_workup_census_summary_v1", {"bold": True}),
     (" — long-format aggregate roll-up, built to drive a dashboard.", {})],
    [("vw_nuclear_med_dated_v1", {"bold": True}),
     (" — nuclear-medicine scans with scan dates recovered (see below).", {})],
])

heading("2.1 Data issue found and fixed: nuclear medicine had no usable dates", 2)
para("The first build of the census view showed nuclear medicine as entirely absent — zero patients, preop or "
     "postop. The cause was not missing data: pub_workspace.nuclear_med_clean.scandate_parsed was NULL for all 2,220 "
     "rows because the parser step never ran, even though the raw scandate column is a clean MM/DD/YYYY string that "
     "parses for 2,218 of the 2,220 rows. A new view, vw_nuclear_med_dated_v1, parses those dates, and the census now "
     "correctly resolves nuclear medicine to 1,148 patients with scans on file — 511 with a preoperative scan "
     "(median 174 days before surgery) and 686 with a postoperative scan (median 99 days after). Two rows (a typo'd "
     "year and one blank) were routed to chart review.")

heading("2.2 Issues logged to Linear", 2)
para("Every issue that surfaced during this work was filed in Linear under the Thyroid Database team so nothing is lost:")
make_table(
    ["Issue", "Title", "Disposition"],
    [
        ["THY-86", "Nuclear medicine scan dates unparsed — scandate_parsed NULL for all 2,220 rows",
         "High priority. Interim fix shipped in pub_eval; recommends an upstream backfill plus a QC assertion."],
        ["THY-87", "Three competing surgery-date columns on canonical_patient_master (SURG01)",
         "Medium. The census view had to pick first_surgery_date as its anchor; one canonical column should be chosen."],
        ["THY-88", "pub_eval global evaluation layer created — wire into a QC gate / Pipeline refresh",
         "Medium. Tracks promoting the views, adding QC assertions, and resolving the US/CT/MRI lineage caveats."],
    ],
    widths=[0.85, 4.6, 3.9])

heading("2.3 What the census shows", 2)
para("Early numbers from the census, useful as a baseline for the dataset’s completeness:")
make_table(
    ["Modality", "Preop coverage", "Median days pre-op", "Postop coverage"],
    [
        ["FNA", "46.6% (5,063)", "91", "3.1% (336)"],
        ["Ultrasound", "32.6% (3,539)", "212", "15.6% (1,699)"],
        ["CT", "23.8% (2,592)", "181", "8.6% (930)"],
        ["Nuclear medicine", "4.7% (511)", "174", "6.3% (686)"],
        ["MRI", "2.9% (319)", "339", "1.0% (113)"],
    ],
    widths=[2.3, 2.2, 1.9, 2.95])
para("Patient-level flags: reoperation 10.1%, completion thyroidectomy 6.1%, prior thyroid procedure documented 8.7%, "
     "and 42.8% of patients in the “sparse” preop-workup tier (no preoperative ultrasound and no preoperative "
     "FNA on record) — itself a finding worth a closer look.")

# ============================ 3. WHERE THE PROJECT STANDS ============================
heading("3. Where the project stands today")
para("A walk through the live project shows a great deal of infrastructure already in place — and a clear set of "
     "recurring failure modes that the new Studio features are well suited to close.")
heading("3.1 What is already working well", 2)
bullets([
    "A layered dataset structure exists: pub_raw, pub_staging, pub_legacy_source_20260416, pub_canonical, pub_semantic, "
    "pub_views_readable, pub_workspace, pub_signoff, and pub_archive.",
    "A genuine governance layer exists in pub_signoff: a table sign-off registry (239 tables), a column verification "
    "registry (6,819 columns), a deprecation registry, a 111-row migration log, a master drift gate, and daily spend / "
    "expensive-query logs.",
    "A QC rule framework exists: 20 documented rules in qc_rules_v1, 368 logged violations, 18 run-based assertions, "
    "and daily/weekly QC summaries.",
    "Discipline is strong: every mutation is snapshotted into pub_archive, conventions are written down in a "
    "__conventions table, and a __readme change-log records what each build script did.",
])
heading("3.2 The recurring pain points", 2)
para("These are not hypothetical — each one is visible in the project today:")
make_table(
    ["Pattern", "What it looks like in the project", "Cost to us"],
    [
        ["Validation is reactive",
         "QC rules live in pub_workspace and are run as scripts; violations are logged but do not block a table from "
         "being published or read.",
         "Bad values reach cohorts and manuscript numbers before anyone notices."],
        ["Competing sources of truth",
         "Rules flag three disagreeing surgery-date columns, two LN-positive columns that disagree, and a cross-source "
         "T/N/M disagreement table with 4,256 rows.",
         "Analysts must know which column to trust; different analyses silently pick differently."],
        ["Legacy-migration scarring",
         "~135 frozen tables in pub_legacy_source, ~175 pre-migration snapshots in pub_archive, dozens of MIG_* tables, "
         "and a __readme full of post-migration patches.",
         "Work done in the legacy system breaks when carried over; fixes are one-off."],
        ["Broken / opaque lineage",
         "Orphan-reference and unresolved-pointer audit tables exist; research_id is VARCHAR in some tables and INTEGER "
         "in others, which silently produces 0% join overlap.",
         "Backtracking to find why something is not linked is a frequent, manual task."],
        ["Builder gaps found late",
         "The __readme repeatedly notes “builder did not retrigger” and “pipeline gap.” The "
         "nuclear-medicine date gap in Section 2 is exactly this pattern.",
         "Derived tables go stale silently; nobody is told a dependency moved."],
        ["Workspace sprawl",
         "pub_workspace holds ~330 tables: many near-duplicate cohort versions, dry-run tables, and snapshots with no "
         "lifecycle policy.",
         "Hard to tell live from dead; pub_staging sits empty while pub_raw holds only two tables."],
    ],
    widths=[1.6, 4.3, 3.45])
rich([("Read together, the theme is clear: ", {}),
      ("we have the discipline and the audit trail, but not the automation and the visibility.", {"bold": True}),
      (" The features below are aimed squarely at that gap.", {})])

# ============================ 4. FEATURE MAP ============================
heading("4. BigQuery Studio feature map")
para("Each feature below is matched to the pain point it most directly addresses, with a concrete first use in our "
     "project.")

heading("4.1 Data Preparation", 2)
rich([("Addresses: ", {"bold": True}), "data validation & quality; legacy migration."], after=3)
para("Data Preparation profiles a table, surfaces anomalies and schema drift, and lets you build a Gemini-assisted, "
     "reusable cleaning recipe instead of a one-off SQL script.")
bullets([
    "Build prep recipes for the known normalization problems: histology vocabulary, free-text extrathyroidal-extension "
    "strings, the VARCHAR-vs-INTEGER research_id mismatch, and exactly the kind of unparsed date column that the "
    "nuclear-medicine fix in Section 2 had to handle by hand.",
    "Use it as the standard intake path for pub_legacy_source_20260416 — profile each legacy table, capture the "
    "transform as a recipe, and land the result in pub_staging (currently empty, and meant to be the staging layer it "
    "was named for).",
    "Recipes are versioned and re-runnable, so a legacy fix is documented once and applied consistently rather than "
    "re-discovered.",
])

heading("4.2 Data Canvas", 2)
rich([("Addresses: ", {"bold": True}), "broken / opaque lineage; visibility & collaboration."], after=3)
para("Data Canvas is a visual, node-based workspace for joining, exploring, and tracing data with natural-language "
     "assistance.")
bullets([
    "Prototype new evaluation views on a canvas before committing SQL — the joins across imaging, FNA, operative, "
    "and nuclear-medicine tables behind the census view are exactly what a canvas is for.",
    "When two columns disagree (for example the LN-positive columns), trace each one back to its feeder tables visually, "
    "so the “which source do we trust” conversation happens over a picture.",
    "Use saved canvases as onboarding artifacts: a canvas of the canonical patient spine and its main feeders is worth "
    "more than a paragraph of documentation.",
])

heading("4.3 Notebooks (and the Notebook gallery)", 2)
rich([("Addresses: ", {"bold": True}), "validation & quality; visibility & collaboration."], after=3)
para("Notebooks combine SQL, Python, BigQuery ML, and inline charts in one shareable, re-runnable document. The Notebook "
     "gallery lets us start from standardized templates.")
bullets([
    "Convert the ad hoc decision-log JSON files into notebooks: a data-quality report notebook that runs the QC rules, "
    "charts the violation trend, and is readable by a non-engineer.",
    "Create a gallery template for “new manuscript cohort” so every cohort starts with the same cohort-scoping "
    "checks and RID-cast conventions — this directly attacks the cohort_m* sprawl.",
    "Make notebooks the home of cohort feasibility analysis, so the reasoning behind an N is attached to the cohort, "
    "not buried in commit history.",
])

heading("4.4 Visualizations", 2)
rich([("Addresses: ", {"bold": True}), "visibility & collaboration."], after=3)
bullets([
    "Stand up a QC dashboard: violation counts by rule and severity over time, drawing on qc_violations_v1 and the "
    "daily/weekly summaries.",
    "Publish vw_workup_census_summary_v1 as a standing completeness dashboard — it was built in long format "
    "specifically to feed this.",
    "Build a cohort-attrition waterfall so the team can see where the dataset is thin before designing an analysis.",
])

heading("4.5 Apache Spark (serverless)", 2)
rich([("Addresses: ", {"bold": True}), "broken lineage; legacy migration at scale."], after=3)
bullets([
    "Run cross-source disagreement detection at scale — the kind of work behind the 4,256-row T/N/M disagreement "
    "table — as a repeatable job rather than a one-off script.",
    "Use it for fuzzy matching when repairing broken linkages (orphaned pointers, unresolved references) where exact "
    "joins fail.",
    "Use it for batch reconciliation of LLM-extracted entity tables against the structured canonical tables.",
])

heading("4.6 Pipelines", 2)
rich([("Addresses: ", {"bold": True}), "all four priorities — especially “builder did not retrigger.”"], after=3)
bullets([
    "Model the canonical build as a DAG: ingest → stage → validate (gate) → canonical → rollup "
    "→ QC → publish. When an upstream table changes, every dependent rebuild is triggered automatically "
    "— the nuclear-medicine date gap is precisely the failure this prevents.",
    "Insert the QC rules as a blocking gate: a canonical table is not promoted until critical-severity rules pass.",
    "Schedule the QC dashboard refresh and the pub_eval evaluation-layer refresh so they are never stale (tracked in "
    "THY-88).",
])

heading("4.7 Custom Agents and Gemini in queries", 2)
rich([("Addresses: ", {"bold": True}), "visibility & collaboration; validation."], after=3)
bullets([
    "Build a project-aware agent seeded with the rules already written in __conventions — always CAST research_id, "
    "use canonical_patient_master for cohort-wide questions, never read pub_workspace as production.",
    "Let clinicians and analysts ask questions like “how many patients had a preoperative FNA but no pathology on "
    "file?” without hand-writing joins — the pub_eval census view answers exactly this class of question.",
    "Use Gemini in queries for day-to-day SQL assistance, with the agent as the guardrail that keeps answers consistent "
    "with our conventions.",
])

heading("4.8 Connections and Files explorer", 2)
rich([("Addresses: ", {"bold": True}), "legacy migration; project hygiene."], after=3)
bullets([
    "Use Connections to bring legacy Excel and Parquet sources in through managed external tables rather than one-off "
    "load scripts, so re-ingestion is reproducible.",
    "Use the Files explorer to keep notebook and pipeline assets organized as the team grows, instead of tracking them "
    "by script number.",
])

# ============================ 5. GLOBAL EVALUATION LAYER ============================
heading("5. The global evaluation layer (built)")
para("A recurring request is a way to evaluate the dataset globally — to ask, for any slice of patients, what "
     "work-up actually happened and where the record is thin. That answer used to be scattered across imaging, FNA, "
     "operative, nuclear-medicine, and past-surgical-history tables, each with its own grain. It is now consolidated in "
     "the pub_eval dataset.")
heading("5.1 vw_patient_workup_census_v1", 2)
para("One row per patient (10,871, validated: no null IDs, no null anchors, no negative intervals). For each imaging / "
     "tissue modality, the same pattern: was it done before surgery, how long before, was it done after surgery, how "
     "long after. The surgery anchor is first_surgery_date, exposed as surgery_anchor_date for transparency (see THY-87).")
make_table(
    ["Field group", "Type", "Definition / source"],
    [
        ["us_preop_performed / _first_date / _interval_days; us_postop_performed / _last_date / _interval_days",
         "BOOL / DATE / INT",
         "Ultrasound, derived from canonical_patient_master us_first_exam_date / us_last_exam_date."],
        ["ct_* and mri_* (same pattern) plus ct_n_exams / mri_n_exams", "BOOL / DATE / INT",
         "CT and MRI, derived from the canonical *_first_date / *_last_date columns."],
        ["fna_preop_performed / _first_date / _interval_days / _n; fna_postop_*", "BOOL / DATE / INT",
         "FNA, event-level from canonical_fna_events_v1 — exact preop/postop counts and first dates."],
        ["nucmed_any_on_file / nucmed_n_total / nucmed_date_resolved; nucmed_preop_* / nucmed_postop_*",
         "BOOL / INT / DATE",
         "Nuclear medicine, event-level from pub_eval.vw_nuclear_med_dated_v1 (dates recovered — see Section 2.1)."],
        ["n_preop_modalities / preop_core_workup_score / preop_workup_tier", "INT / FLOAT / STRING",
         "Derived completeness summary over the preoperative modalities (tier: core_complete / partial / sparse)."],
        ["n_surgeries / n_completion_thyroidectomies / any_reoperative_field / reoperation_flag", "INT / FLOAT / BOOL",
         "From canonical_operative_patient_rollup_v1_1."],
        ["prior_thyroidectomy_documented / prior_neck_surgery_documented / prior_thyroid_procedure_documented", "BOOL",
         "From canonical_psh_patient_rollup_v1 (past-surgical-history evidence flags)."],
        ["has_surgical_pathology_on_file / prior_procedure_path_gap_flag", "BOOL",
         "prior_procedure_path_gap_flag = prior procedure documented AND no surgical pathology on file — a "
         "chart-review trigger, not a definitive determination."],
        ["lobectomy_first_flag / completion_path_definite_flag / lobectomy_first_no_completion_path_flag", "BOOL",
         "From patient_completion_oed_path_linkage_v1 — completion-thyroidectomy pathology documentation gaps."],
    ],
    widths=[3.0, 1.5, 4.85])
heading("5.2 vw_workup_census_summary_v1", 2)
para("A long-format aggregate roll-up — one row per metric — built to drive the Visualization dashboard. "
     "imaging_modality rows give preop/postop coverage and interval medians per modality; cohort_flag rows give "
     "patient-level prevalences (reoperation, prior-procedure path gap, workup tiers, and the nuclear-medicine coverage "
     "flags).")
heading("5.3 Known caveats (tracked in THY-88)", 2)
bullets([
    "US, CT, and MRI pre/post are derived from patient-level first/last date columns, so n_exams is a patient total and "
    "the postop reference is the last exam, not the first postop exam. FNA and nuclear medicine are event-level and "
    "exact. Building event-level US/CT/MRI feeds would make all five modalities symmetric.",
    "These are views, not materialized tables. Promotion plus a scheduled Pipeline refresh and census-specific QC "
    "assertions are the next step.",
    "prior_procedure_path_gap_flag should be validated against a chart-reviewed sample before it is used in any "
    "manuscript.",
])

# ============================ 6. BROADER GCP ECOSYSTEM ============================
heading("6. The broader Google Cloud analytics ecosystem")
para("BigQuery Studio is the core, but several adjacent Google Cloud services map cleanly onto our pain points. The "
     "table below rates each on fit — Recommended, Situational, or Not applicable — for this project "
     "specifically. The guiding principle is to keep the stack consolidated around BigQuery and only add a service "
     "where it closes a real gap.")
heading("6.1 Recommended", 2)
make_table(
    ["Service", "How we would use it"],
    [
        ["Looker / Looker Studio",
         "The natural front end for the work already built. Point Looker at vw_workup_census_summary_v1 and the "
         "pub_signoff QC tables to give the team a standing completeness dashboard and a QC-trend dashboard without "
         "anyone re-running a query. Looker Studio is the lightweight, no-cost starting point; Looker proper if we need "
         "governed, shared models."],
        ["Data Catalog / Dataplex (Knowledge Catalog)",
         "Directly targets the broken-lineage pain. A searchable catalog with table- and column-level tags and lineage "
         "means people stop guessing which table is canonical, and the orphan-reference / unresolved-pointer problem "
         "becomes visible instead of discovered by accident. This is the metadata backbone the project is missing."],
        ["Managed Airflow (Cloud Composer)",
         "Workflow orchestration for the ingest → stage → validate → canonical → rollup → QC "
         "→ publish DAG. BigQuery Pipelines may be enough on its own; Composer is the answer if we outgrow it or "
         "need cross-service orchestration. Either way, the numbered-script chain becomes a real dependency graph."],
        ["Managed Apache Spark / Dataproc",
         "Serverless PySpark for the heavy reconciliation work — cross-source disagreement detection, fuzzy "
         "linkage repair, LLM-output reconciliation — that is awkward in pure SQL. Same recommendation as Section "
         "4.5, delivered as a managed service."],
        ["Document AI",
         "High value for a chart-heavy dataset. Document AI can parse scanned or free-text pathology reports, operative "
         "notes, and nuclear-medicine narratives into structured fields — directly relevant to recovering the two "
         "undateable nuclear-medicine rows in THY-86 and to the broader LLM-extraction effort."],
        ["Agent Platform / Gemini Enterprise",
         "Where the project-aware custom agent from Section 4.7 would actually be built and deployed — an agent "
         "that knows our conventions and can answer work-up and cohort questions in plain language against the pub_eval "
         "views."],
    ],
    widths=[1.9, 7.45])
heading("6.2 Situational — adopt if a specific need appears", 2)
make_table(
    ["Service", "When it would make sense"],
    [
        ["Healthcare API / Healthcare data storage",
         "If we ever ingest EHR data directly or need a FHIR/HL7 representation of the cohort. For a research database "
         "built from periodic Excel/Parquet extracts, this is not needed today."],
        ["NotebookLM for Enterprise",
         "A useful adjunct for the team to query and understand its own protocols, manuscript drafts, and "
         "documentation — not part of the data pipeline, but a low-cost productivity aid."],
        ["Data Fusion / Alteryx Designer Cloud",
         "Visual ETL / data wrangling. These overlap with BigQuery Data Preparation, which is the lighter-weight native "
         "option we should try first; reach for these only if Data Preparation proves too limited."],
        ["Lakehouse (Iceberg on Cloud Storage)",
         "Could serve as the landing format for the legacy Parquet sources feeding pub_raw / pub_staging. A nice-to-have "
         "for storage hygiene, not a priority."],
        ["Datastream",
         "Only relevant if live EHR replication ever becomes a requirement; not applicable to the current batch-extract "
         "model."],
    ],
    widths=[1.9, 7.45])
heading("6.3 Not applicable to this project", 2)
para("Pub/Sub, Dataflow, Managed Service for Apache Kafka, and Confluent Cloud are real-time streaming and messaging "
     "tools. This is a static, batch research dataset assembled from periodic data drops — there is no streaming "
     "workload to justify them. Databricks and Elastic Cloud are capable platforms but would fragment a stack that "
     "BigQuery plus serverless Spark already covers. Earth Engine, Search for commerce, Talent Solution, AI Edge "
     "Portal, and the CCAI Platform are aimed at domains (geospatial, retail, recruiting, on-device ML, contact "
     "centers) unrelated to clinical thyroid research.")

# ============================ 7. PHASED ROLLOUT ============================
heading("7. Phased rollout")
para("Four phases over roughly ten weeks. Phase 0 is partly done; each remaining phase produces something the team can "
     "use immediately.")
heading("Phase 0 — Foundations and quick wins (Weeks 1–2)", 2)
numbered([
    "DONE — pub_eval global evaluation layer built; nuclear-medicine dates recovered; issues filed as THY-86, "
    "THY-87, THY-88.",
    "Build the first data-quality notebook from a Notebook gallery template: run the 20 QC rules, chart the violation "
    "trend, share it.",
    "Recreate the canonical patient spine and its main feeders as a saved Data Canvas — the first shared lineage "
    "artifact.",
    "Agree a naming and lifecycle convention for pub_workspace (what gets a TTL, what gets archived).",
])
heading("Phase 1 — Validation as a gate (Weeks 3–6)", 2)
numbered([
    "Consolidate the QC rules into one location and wrap them in a Pipeline step; add the census-specific assertions "
    "from THY-88.",
    "Make that step a blocking gate: critical-severity rules must pass before a canonical or safe-view table is "
    "published.",
    "Build Data Preparation recipes for histology vocabulary, ETE strings, the research_id type mismatch, and unparsed "
    "date columns.",
    "Publish the QC dashboard and the workup-census dashboard in Looker Studio.",
])
heading("Phase 2 — Lineage and metadata (Weeks 6–10)", 2)
numbered([
    "Stand up Data Catalog / Dataplex: tag canonical vs workspace vs archive, and turn on lineage tracking.",
    "Promote the pub_eval views to materialized tables with a scheduled Pipeline refresh; resolve the US/CT/MRI "
    "event-level caveat.",
    "Adopt Data Canvas lineage tracing as the standard first step whenever two columns disagree.",
    "Stand up the first serverless Spark job for cross-source disagreement detection; route legacy intake through "
    "Connections + Data Preparation into pub_staging.",
])
heading("Phase 3 — Agents and self-serve (Weeks 10+)", 2)
numbered([
    "Build the project-aware custom Agent on Agent Platform, seeded with the __conventions rules, and give the team "
    "access.",
    "Pilot Document AI on a batch of pathology / nuclear-medicine narratives to recover structured fields.",
    "Publish the “new manuscript cohort” Notebook gallery template; automate the pub_workspace lifecycle "
    "policy.",
    "Review: re-run the Section 3.2 pain-point table and confirm each pattern has a gate, a view, or a catalog entry "
    "behind it.",
])

# ============================ 8. GOVERNANCE ============================
heading("8. Governance and ownership")
para("The features only help if a few habits change with them:")
bullets([
    "One QC home. The rules, the violations, and the gate live together and are versioned; new rules are added there, "
    "not in ad hoc tables.",
    "The gate is real. A canonical or safe-view table that fails a critical rule is not published — the Pipeline "
    "enforces it, not a person remembering to check.",
    "pub_staging is used. Legacy and raw intake lands there first and is promoted only after a prep recipe and the QC "
    "gate.",
    "Workspace has a lifecycle. Every pub_workspace table has an owner and a sunset trigger, or it gets a TTL; dry-run "
    "and snapshot tables are swept on a schedule.",
    "Lineage is a picture and a catalog. Disagreements are investigated on a Data Canvas and recorded in Data Catalog, "
    "so the next person does not start from zero.",
    "Findings are tracked. Issues surfaced by the QC gate or by analysis are filed in Linear — as THY-86/87/88 "
    "already are — not left in comments.",
])

# ============================ APPENDICES ============================
heading("Appendix A — Feature-to-pain matrix")
make_table(
    ["Feature", "Validation & quality", "Legacy migration", "Lineage / linking", "Visibility"],
    [
        ["Data Preparation", "Primary", "Primary", "Supporting", "—"],
        ["Data Canvas", "—", "Supporting", "Primary", "Primary"],
        ["Notebooks", "Primary", "Supporting", "—", "Primary"],
        ["Visualizations / Looker", "Supporting", "—", "—", "Primary"],
        ["Apache Spark / Dataproc", "Supporting", "Primary", "Primary", "—"],
        ["Pipelines / Composer", "Primary", "Primary", "Primary", "Supporting"],
        ["Custom Agents / Gemini", "Supporting", "—", "Supporting", "Primary"],
        ["Data Catalog / Dataplex", "Supporting", "Supporting", "Primary", "Primary"],
        ["Document AI", "Supporting", "Primary", "—", "—"],
    ],
    widths=[2.05, 1.9, 1.8, 1.8, 1.8])

heading("Appendix B — pub_eval objects created")
make_table(
    ["Object", "Description"],
    [
        ["pub_eval", "New dataset (us-central1). Views only — non-destructive."],
        ["vw_patient_workup_census_v1",
         "One row per patient (10,871). Preop/postop modality flags + intervals, reoperation signals, prior-procedure "
         "pathology-gap flags."],
        ["vw_workup_census_summary_v1", "Long-format aggregate roll-up of the census, for dashboards."],
        ["vw_nuclear_med_dated_v1",
         "Nuclear-medicine scans with scan dates parsed from the raw scandate string; carries a scandate_quality flag."],
    ],
    widths=[2.6, 6.75])

heading("Appendix C — First-two-weeks checklist")
bullets([
    "DONE — pub_eval evaluation layer built and validated.",
    "DONE — nuclear-medicine dates recovered; THY-86/87/88 filed.",
    "Data-quality notebook live and shared.",
    "Canonical-spine Data Canvas saved.",
    "pub_workspace naming + lifecycle convention agreed in writing.",
    "Owners named for Phase 1 (QC gate) and Phase 2 (metadata & lineage).",
])
para("Prepared as a draft for team review. Figures (dataset, table, rule, and census counts) reflect a live inspection "
     "of thyroid-canonical-pub-2026 on May 14, 2026 and should be re-checked before this plan is circulated widely.",
     italic=True, size=9, color=LGREY, before=8)

out = sys.argv[1] if len(sys.argv) > 1 else "BigQuery_Studio_Integration_Plan.docx"
doc.save(out)
print("written %s" % out)
