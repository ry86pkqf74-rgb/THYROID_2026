"""05_build_d4_emethods.py — Deliverable 4
M038_GOITER_eMethods.docx — 8-section statistical methods document.
US Letter, Arial 11pt body / 14pt headings.
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
PKG = HERE.parent
OUT = PKG / "M038_GOITER_eMethods.docx"


def set_run(run, font="Arial", size=11, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=16, bold=True, color=RGBColor(0x1F, 0x4E, 0x78))
        p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run(run, size=13, bold=True, color=RGBColor(0x1F, 0x4E, 0x78))
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    else:
        set_run(run, size=11, bold=True, color=RGBColor(0x1F, 0x4E, 0x78))
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, italic=False, size=11, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size=size, italic=italic)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = ""  # clear style placeholder
    run = p.add_run(text)
    set_run(run, size=11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_kv_table(doc, rows, col_widths=(2.5, 4.0)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        c1 = table.cell(i, 0); c2 = table.cell(i, 1)
        c1.width = Inches(col_widths[0]); c2.width = Inches(col_widths[1])
        for c in c1.paragraphs[0].runs: c.text = ""
        for c in c2.paragraphs[0].runs: c.text = ""
        r1 = c1.paragraphs[0].add_run(k); set_run(r1, size=10, bold=True)
        r2 = c2.paragraphs[0].add_run(v); set_run(r2, size=10)
    return table


def main():
    doc = Document()

    # Page setup: US Letter, 1" margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0); section.bottom_margin = Inches(1.0)

    # Default style → Arial 11pt
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ---------- Title ----------
    p = doc.add_paragraph()
    run = p.add_run("Supplemental eMethods")
    set_run(run, size=20, bold=True, color=RGBColor(0x1F, 0x4E, 0x78))
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run("Massive Goiter at a Tertiary Referral Center: A Composite-Definition Descriptive Cohort of 2,501 Patients")
    set_run(run, size=12, bold=True, italic=False)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(f"Manuscript ID: M038  •  Document date: {date.today().isoformat()}  •  Database release: pub_v1_1_20260504")
    set_run(run, size=10, italic=True, color=RGBColor(0x55, 0x55, 0x55))
    p.paragraph_format.space_after = Pt(14)

    # ---------- §1 ----------
    add_heading(doc, "1. Study Design and Data Source", level=1)
    add_para(doc,
        "This was a single-institution, retrospective cohort study of all adult patients undergoing thyroid surgery "
        "at Emory University Hospital and affiliated facilities between January 1, 1999 and December 31, 2025. Source "
        "data were drawn from the THYROID_2026 canonical lakehouse (MotherDuck database "
        "thyroid_canonical_publication_v1_0, release pub_v1_1_20260504, most-recent applied migration "
        "mig_255_cohort_m038_complication_temporality_columns_20260502). The lakehouse integrates structured EHR "
        "extracts, NSQIP institutional submission records, synoptic pathology reports, operative notes (with "
        "domain-specific NLP enrichment), and imaging-report NLP extractions.")
    add_para(doc,
        "All analytic queries are reproducible against the post-mig_255 cohort view "
        "manuscript_workspace.cohort_m038_massive_goiter_v1, which materializes the per-patient analytic record. "
        "The reproducible SQL package accompanies this submission as 08_analysis_code/M038_descriptive_analysis.sql; "
        "the per-patient parquet extract used for the four submission-package deliverables was pulled on the date "
        "above using the script 01_pull_parquet.py.")

    # ---------- §2 ----------
    add_heading(doc, "2. Cohort Definition", level=1)
    add_para(doc,
        "The analytic cohort comprises all patients with at least one thyroid surgery in the canonical "
        "patient master and at least one resolved histologic record. The total cohort is n = 10,871. Patients were "
        "stratified into a binary exposure based on a composite-massive flag (see §3).")
    add_para(doc,
        "There were no exclusions for age, prior thyroid surgery, indication, or follow-up duration; the goal of "
        "this descriptive cohort was to characterize the population presenting for thyroid surgery, including those "
        "with antecedent thyroid history, in order to estimate true clinical prevalence of the massive-goiter "
        "phenotype.")

    # ---------- §3 ----------
    add_heading(doc, "3. Exposure Definition (Composite Massive Flag)", level=1)
    add_para(doc,
        "The composite-massive flag (is_massive) is defined by the disjunction of three component groups, each "
        "derived from a distinct data source:")
    add_bullet(doc, "Weight component (W): synoptic pathology gland weight ≥100 g (gland_weight_final_g >= 100).")
    add_bullet(doc,
        "Substernal component (S): substernal extension documented on cross-sectional imaging — "
        "ct_substernal_extension_any OR mri_substernal_any.")
    add_bullet(doc,
        "Airway component (A): airway compromise documented on CT — ct_tracheal_deviation_any OR "
        "ct_tracheal_narrowing_any OR ct_airway_compromise_any.")
    add_para(doc,
        "A patient is classified as composite-massive (is_massive = TRUE) if any of W, S, or A is TRUE; that is, "
        "is_massive = (W ∨ S ∨ A). The disjunction was chosen because each component captures a distinct clinical "
        "axis of the massive-goiter phenotype (mass burden, mediastinal extension, and tracheal compromise) and "
        "because patients can present with any one as the dominant feature. Component overlap is reported in "
        "Tables 1 and 5 and Figure 1.")

    # ---------- §4 ----------
    add_heading(doc, "4. Variables", level=1)
    add_heading(doc, "4.1 Demographics & comorbidities", level=2)
    add_para(doc,
        "Age at first surgery, sex, race (9-bucket), and combined BMI (NSQIP-first with vitals/derived fallback; "
        "source documented in bmi_source) were extracted from the canonical patient master. Comorbidity flags were "
        "extracted via two independent sources: NSQIP linkage (nsqip_diabetes, nsqip_hypertension, nsqip_copd, "
        "nsqip_heart_failure, nsqip_bleeding_disorder, nsqip_disseminated_cancer, nsqip_functional_status) and a "
        "narrative-NLP layer over past medical history sections (pmhx_nlp_diabetes, pmhx_nlp_hypertension, "
        "pmhx_nlp_cad, pmhx_nlp_ckd, pmhx_nlp_copd, pmhx_nlp_n_comorbidities, pmhx_nlp_autoimmune_thyroid_hx). "
        "ASA physical status class (nsqip_asa_class) is reported on the NSQIP-linked subset only.")

    add_heading(doc, "4.2 Surgical context", level=2)
    add_para(doc,
        "First-surgery date (surg_first_date) was backfilled in mig_254 to maximize era-attribution coverage. "
        "Procedure-type completeness was finalized in mig_253 and reaches 100% in the massive arm and 99.98% in "
        "the non-massive arm. Operative duration (nsqip_operative_duration_min), length of stay "
        "(nsqip_length_of_stay_days), drain usage, vessel sealant, RLN monitoring, and central/lateral neck "
        "dissection were sourced from NSQIP records. Operative-note NLP supplied the difficult-airway flag "
        "(ops_difficult_airway). Tracheostomy was extracted from procedure-NLP (proc_nlp_tracheostomy).")

    add_heading(doc, "4.3 Pathology", level=2)
    add_para(doc,
        "Resolved histologic diagnosis (histology_final), malignancy flag (is_malignant), bilateral disease "
        "(bilateral_path_flag from pathology, bilateral_disease_flag combining clinical + imaging), and closest "
        "margin (closest_margin_mm) were drawn from synoptic pathology with manual reconciliation of conflicting "
        "reports.")

    add_heading(doc, "4.4 Complication outcomes (strict-definition)", level=2)
    add_para(doc,
        "Complication outcomes use the strict-definition canonical rollup (post-mig_252) that requires both a "
        "definite/probable status assertion AND temporal evidence the event occurred postoperatively. The columns "
        "are any_confirmed_complication_flag (composite), comp_hematoma_confirmed, comp_seroma_confirmed, "
        "comp_chyle_leak_confirmed, comp_rln_injury_confirmed, comp_vc_paresis_confirmed, "
        "comp_vc_paralysis_confirmed, comp_hypocalcemia_confirmed, comp_hypoparathyroidism_confirmed, "
        "comp_airway_complication_definitive, comp_pneumothorax_definitive, and comp_mortality_definitive.")
    add_para(doc,
        "Per the standing rule (memory/feedback_complications_transient_vs_permanent.md), hypoparathyroidism is "
        "split into transient (<6 months postop), permanent (>6 months postop), preexisting (preop), and new "
        "postop using the temporality columns added in mig_255: comp_hypoparathyroidism_transient, "
        "comp_hypoparathyroidism_permanent, comp_hypoparathyroidism_timing_window, "
        "comp_hypoparathyroidism_preexisting, and comp_hypoparathyroidism_new_postop. Hypocalcemia carries a "
        "preexisting flag derived from comp_hypocalcemia_timing_window = 'pre_surgery' OR "
        "comp_hypocalcemia_clinical_preexisting; the postop confirmed events are reported separately. The same "
        "transient/permanent decomposition is applied to RLN injury, vocal cord paresis, and vocal cord paralysis "
        "where the temporality columns are populated.")

    add_heading(doc, "4.5 NSQIP 30-day complications", level=2)
    add_para(doc,
        "NSQIP-flagged 30-day complications (transfusion, neck hematoma, RLN injury flag, hypocalcemia flag, "
        "unplanned intubation, unplanned return to OR, 30-day readmission, 30-day mortality, pneumonia, DVT, PE, "
        "sepsis, superficial/deep/organ-space SSI) are reported on the NSQIP-linked subset only and are presented "
        "separately from the strict-definition canonical complications (Table 4).")

    add_heading(doc, "4.6 Era binning", level=2)
    add_para(doc,
        "Surgery era was binned in two ways: a 5-year bucket scheme (1999-2004, 2005-2009, 2010-2014, 2015-2019, "
        "2020-2025) for component-coverage analysis (Figure 4 and Supp S2), and the 3-bucket headline scheme used "
        "in the manuscript abstract (pre-2015, 2015-2019, 2020-2025). The upper-bound rule sweeps any surgery "
        "with surg_first_date earlier than 1999-01-01 into the 1999-2004 bucket. Patients with NULL surg_first_date "
        "are reported as 'unknown' and excluded from era-trend tests.")

    # ---------- §5 ----------
    add_heading(doc, "5. Statistical Methods", level=1)
    add_para(doc,
        "Continuous variables are summarized as median (IQR) with mean ± SD where indicated; categorical "
        "variables as n (%). Comparisons between massive and non-massive arms used Mann-Whitney U tests for "
        "continuous variables and chi-squared (without Yates continuity correction) for categorical variables. "
        "Fisher exact tests substituted for chi-squared whenever any expected cell count fell below 5, when chi-"
        "squared computation was undefined (zero-row or zero-column tables), or when low-event scenarios warranted "
        "exact inference. Two-sided p-values are reported throughout; p < 0.05 was considered nominally "
        "statistically significant. No multiplicity correction was applied because of the descriptive intent of "
        "this study.")
    add_para(doc,
        "Risk ratios (RR) for binary complication outcomes were computed as massive prevalence divided by "
        "non-massive prevalence, with 95% confidence intervals derived from the Wald approximation on the natural "
        "log of the RR using SE = √[(1−p₁)/a + (1−p₂)/b], where a and b are event counts and p₁/p₂ are arm-specific "
        "prevalences. For zero-cell or near-zero-cell scenarios a Haldane-Anscombe 0.5 continuity correction was "
        "applied to the contingency cells before the SE calculation. Wilson 95% CIs were used for arm-prevalence "
        "estimates in Table 5 to avoid pathology at low or high prevalences.")
    add_para(doc,
        "Era-trend analysis used a chi-squared test of independence on the 3 × 2 contingency table of era × "
        "(massive / non-massive). Component-coverage trends across era are reported descriptively (Figure 4 / "
        "Supp S2) without a formal trend test because component documentation rates rather than disease prevalence "
        "are the primary signal.")

    # ---------- §6 ----------
    add_heading(doc, "6. Complication Definitions and Standing Rule", level=1)
    add_para(doc,
        "The strict-definition complication rollup (suffix _confirmed for soft-tissue/RLN/biochemical events; "
        "suffix _definitive for airway/pulmonary/mortality events) requires (a) a definite or probable status "
        "assertion in the source documentation and (b) temporal evidence that the event occurred postoperatively. "
        "Events confined to the preoperative window are reported in dedicated preexisting rows and are not counted "
        "in the postop confirmed numerator. NSQIP 30-day complication flags are a parallel signal reported "
        "separately because their case definitions, denominator (NSQIP-linked subset), and 30-day window differ "
        "from the canonical rollup.")
    add_para(doc,
        "Hypoparathyroidism standing-rule details (per memory/feedback_complications_transient_vs_permanent.md, "
        "ratified 2026-05-01):")
    add_bullet(doc,
        "Postop transient hypoparathyroidism (<6 months): comp_hypoparathyroidism_confirmed AND "
        "comp_hypoparathyroidism_transient.")
    add_bullet(doc,
        "Postop permanent hypoparathyroidism (>6 months): comp_hypoparathyroidism_confirmed AND "
        "comp_hypoparathyroidism_permanent.")
    add_bullet(doc,
        "Preexisting hypoparathyroidism: comp_hypoparathyroidism_preexisting (preop diagnosis present in source "
        "documentation, irrespective of postop status).")
    add_bullet(doc,
        "New postop hypoparathyroidism: comp_hypoparathyroidism_new_postop (postop event with no preop "
        "evidence).")
    add_para(doc,
        "Hypocalcemia is split into postop confirmed events (transient/permanent if temporality is documented) "
        "and a preexisting row defined by (comp_hypocalcemia_timing_window = 'pre_surgery' OR "
        "comp_hypocalcemia_clinical_preexisting), to mirror the hypoparathyroidism schema. The known carry-forward "
        "items CF-RLN-PREOP-FLAG and CF-VC-PARALYSIS-PREOP-FLAG (no preop status encoding for RLN injury or vocal "
        "cord paralysis) are documented in the validation report and remain open.")

    # ---------- §7 ----------
    add_heading(doc, "7. Missing Data", level=1)
    add_para(doc,
        "Missing data are reported in Supplementary Table S5. Salient coverage notes:")
    add_bullet(doc,
            "Surgery date (surg_first_date): cohort-wide 80.3% known (8,731 / 10,871); massive arm 69.6% known "
            "(1,740 / 2,501). Pre-1999 surgery dates are subject to the upper-bound binning rule.")
    add_bullet(doc,
            "BMI (bmi_combined): predominantly NSQIP-sourced; vitals/derived fallback documented in bmi_source.")
    add_bullet(doc,
            "NSQIP-linkage subset: variables with prefix nsqip_ are populated only for the linked subset (used as "
            "a denominator-defining proxy in Table 3 and Supp S3). Approximately 13% of the full cohort has NSQIP "
            "linkage; massive-arm linkage rate is comparable.")
    add_bullet(doc,
            "Histology (histology_final): missing values are flagged '(missing)' in Tables 2 and pub Table 2; "
            "no imputation was performed.")
    add_para(doc,
        "All missing-data analyses use a complete-case approach within each variable. The descriptive intent of "
        "the study and the absence of a primary inferential endpoint preclude the need for multiple imputation.")

    # ---------- §8 ----------
    add_heading(doc, "8. Software & Reproducibility", level=1)
    add_para(doc,
        "All analyses were performed in Python 3.9 with the following packages: pandas 2.3, numpy, scipy "
        "(stats.chi2_contingency, stats.fisher_exact, stats.mannwhitneyu, stats.ttest_ind), openpyxl "
        "(Excel deliverables), python-docx (Word eMethods). Data extraction from MotherDuck used the duckdb "
        "Python client (1.4) authenticated to the logan.glosser.eras account that owns the canonical publication "
        "share. The full extraction-to-deliverable pipeline is reproducible from the scripts in "
        "data_extraction_v2_20260504/ via:")
    p = doc.add_paragraph()
    run = p.add_run(
        "  python3 01_pull_parquet.py\n"
        "  python3 02_build_d1_patient_dataset.py\n"
        "  python3 03_build_d2_analysis_workbook.py\n"
        "  python3 04_build_d3_tables_figures.py\n"
        "  python3 05_build_d4_emethods.py")
    set_run(run, font="Courier New", size=10)
    p.paragraph_format.space_after = Pt(8)
    add_para(doc,
        "All numeric cells reproduce against the live MotherDuck cohort view at extract time. The independent "
        "validation report (09_validation_report.md) audited 156 manuscript cells against MotherDuck on 2026-05-01 "
        "with 153 PASS / 3 patched / 0 FAIL; the present deliverables re-derive those cells from the post-mig_255 "
        "cohort view as of the document date above. Any cell-level differences from the 2026-05-01 validation "
        "report are traceable to mig_254 (surg_first_date backfill) and mig_255 (complication temporality columns), "
        "which slightly shifted era attribution and complication classification.")

    # ---------- Footer table: deliverable index ----------
    add_heading(doc, "Appendix A — Deliverable Index", level=2)
    add_kv_table(doc, [
        ("Deliverable 1 — Per-patient dataset",
         "M038_GOITER_patient_level_dataset.xlsx (Cover / Patient Data / Data Dictionary)"),
        ("Deliverable 2 — Analysis workbook",
         "M038_GOITER_analysis_workbook.xlsx (9 tabs: overview, T1-T5, NSQIP, component subgroup, exploratory)"),
        ("Deliverable 3 — Pub-formatted tables/figures",
         "M038_GOITER_tables_figures.xlsx (T1-T5 + Fig 1-4 data + Supp S1-S6)"),
        ("Deliverable 4 — eMethods",
         "M038_GOITER_eMethods.docx (this document)"),
        ("Reproducible scripts",
         "data_extraction_v2_20260504/01-05_*.py + _stats.py"),
        ("Standing rule reference",
         "memory/feedback_complications_transient_vs_permanent.md"),
        ("Validation reference",
         "09_validation_report.md (2026-05-01 — 153/156 PASS)"),
    ])

    doc.save(OUT)
    print(f"→ Wrote {OUT}  ({OUT.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
