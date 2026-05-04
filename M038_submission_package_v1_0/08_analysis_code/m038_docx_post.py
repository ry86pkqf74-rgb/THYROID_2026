"""Post-process pandoc-generated 02_manuscript.docx:
- Force Arial 11pt for default + headings
- Set US Letter page size + 1 inch margins
- Embed the 4 figures inline at their callout markers
"""
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from copy import deepcopy
from pathlib import Path
import re

# mig_299: portable paths — script lives at <PKG>/08_analysis_code/<this>.py
PKG = Path(__file__).resolve().parents[1]
DOCX = str(PKG / "02_manuscript.docx")
FIGS_DIR = PKG / "06_figures"

doc = Document(DOCX)

# 1. Page size + margins
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.orientation = WD_ORIENT.PORTRAIT

# 2. Set Arial 11pt for Normal + headings + table styles
styles = doc.styles

def set_font(style, name="Arial", size_pt=11, bold=None):
    try:
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        for attr in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
            rfonts.set(qn(attr), name)
        # set size
        sz = rpr.find(qn("w:sz"))
        if sz is None:
            sz = rpr.makeelement(qn("w:sz"), {qn("w:val"): str(size_pt*2)})
            rpr.append(sz)
        else:
            sz.set(qn("w:val"), str(size_pt*2))
        szcs = rpr.find(qn("w:szCs"))
        if szcs is None:
            szcs = rpr.makeelement(qn("w:szCs"), {qn("w:val"): str(size_pt*2)})
            rpr.append(szcs)
        else:
            szcs.set(qn("w:val"), str(size_pt*2))
    except Exception as e:
        print(f"  warn: could not style {style.name}: {e}")

# Style key paragraph styles
for sname in ["Normal", "Body Text", "Body Text First Indent",
              "Heading 1", "Heading 2", "Heading 3", "Heading 4",
              "Title", "Subtitle", "Caption", "Quote", "Compact"]:
    if sname in [s.name for s in styles]:
        try:
            set_font(styles[sname], "Arial",
                     size_pt={"Heading 1":16, "Heading 2":14, "Heading 3":12, "Heading 4":11,
                              "Title":18, "Subtitle":14}.get(sname, 11))
        except Exception:
            pass

# 3. Embed figures inline at the *[Figure N — ...] callout markers
fig_files = {
    1: FIGS_DIR / "fig1_composite_flag_venn.png",
    2: FIGS_DIR / "fig2_era_prevalence.png",
    3: FIGS_DIR / "fig3_complications_bar.png",
    4: FIGS_DIR / "fig4_component_coverage.png",
}

# Walk paragraphs; when one matches "*[Figure N — ...]*", replace its content with an inline image
for p in list(doc.paragraphs):
    txt = p.text.strip()
    m = re.match(r"^\*?\[Figure (\d) [—-].*\]\*?$", txt)
    if m:
        n = int(m.group(1))
        fig_path = fig_files.get(n)
        if fig_path and fig_path.exists():
            # clear the paragraph
            for r in list(p.runs):
                r.text = ""
            # add image to first run, sized to ~6.5 inches wide
            r = p.add_run()
            r.add_picture(str(fig_path), width=Inches(6.5))
            # add caption below
            new_p = p.insert_paragraph_before("")
            # Actually we want the caption AFTER the image; use add_paragraph trick via XML
            # Simpler: just overwrite text in current paragraph after the image (use run for caption text)
            cap_run = p.add_run()
            cap_run.add_break()
            cap_run.text = ""

doc.save(DOCX)
print(f"Post-processed {DOCX}")
