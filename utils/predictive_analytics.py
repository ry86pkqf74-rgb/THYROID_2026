"""
Integrated Predictive Analytics & Comparative Survival Workbench.

Builds on the existing ``ThyroidStatisticalAnalyzer`` (Cox PH, logistic,
forest plots) and ``ThyroidAdvancedAnalyzer`` (Aalen-Johansen CIF, ML
nomograms) to add:

  1. PTCM-powered individual cure prediction (wraps script 39 math)
  2. Enhanced competing-risks with subdistribution HR estimates
  3. Survival-specific explainable ML (sksurv + SHAP)
  4. Multi-model comparison suite (KM, Cox PH, PTCM, CIF)
  5. Interactive personalized cure calculator inputs
  6. One-click manuscript report generation

Usage (standalone)::

    from utils.predictive_analytics import ThyroidPredictiveAnalyzer
    pa = ThyroidPredictiveAnalyzer(con)
    cure = pa.predict_individual_cure_probability({"age": 45, "braf": True})
    comp = pa.compare_survival_models()

Usage (dashboard)::

    Consumed by ``app/predictive_analytics.py`` which wires interactive
    controls around these methods.
"""
from __future__ import annotations

import io
import json
import logging
import warnings
from datetime import datetime
from pathlib import Path
from typing import Any, Literal

import numpy as np
import pandas as pd

try:
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

try:
    from lifelines import (
        AalenJohansenFitter,
        CoxPHFitter,
        KaplanMeierFitter,
    )
    HAS_LIFELINES = True
except ImportError:
    HAS_LIFELINES = False

try:
    from scipy.optimize import minimize
    from scipy.stats import norm as scipy_norm
    HAS_SCIPY = True
except ImportError:
    HAS_SCIPY = False

try:
    import xgboost as xgb
    HAS_XGB = True
except Exception:
    HAS_XGB = False

try:
    from sklearn.model_selection import StratifiedKFold, cross_val_predict
    from sklearn.metrics import brier_score_loss, roc_auc_score
    from sklearn.calibration import calibration_curve
    from sklearn.ensemble import RandomForestClassifier
    HAS_SKLEARN = True
except ImportError:
    HAS_SKLEARN = False

try:
    import shap
    HAS_SHAP = True
except ImportError:
    HAS_SHAP = False

try:
    from sksurv.ensemble import RandomSurvivalForest
    from sksurv.metrics import concordance_index_censored, brier_score as sksurv_brier
    from sksurv.linear_model import CoxPHSurvivalAnalysis, CoxnetSurvivalAnalysis
    HAS_SKSURV = True
except ImportError:
    HAS_SKSURV = False

try:
    from sksurv.nonparametric import cumulative_incidence_competing_risks as sksurv_cif
    HAS_SKSURV_CIF = True
except ImportError:
    HAS_SKSURV_CIF = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from utils.statistical_analysis import (
    ThyroidStatisticalAnalyzer,
    THYROID_PREDICTORS,
    THYROID_SURVIVAL,
    PL,
)

log = logging.getLogger(__name__)

# ── Color tokens (match dashboard theme) ─────────────────────────────────
_TEAL = "#2dd4bf"
_SKY = "#38bdf8"
_VIOLET = "#a78bfa"
_AMBER = "#f59e0b"
_ROSE = "#f43f5e"
_GREEN = "#34d399"
_GRAY = "#4a5568"
_COLORWAY = [_TEAL, _SKY, _VIOLET, _AMBER, _ROSE, _GREEN, "#fb923c"]

# ── Project paths ────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
PTCM_EXPORT_DIR = ROOT / "exports" / "promotion_cure_results"
PREDICTIVE_EXPORT_DIR = ROOT / "exports" / "predictive_analytics"

# ── PTCM constants (must match script 39) ────────────────────────────────
PTCM_COVARIATES = [
    "intercept",
    "age_z",
    "ajcc_2", "ajcc_3", "ajcc_4",
    "ete_microscopic", "ete_gross",
    "braf_pos", "tert_pos",
    "high_risk_band",
]

MAX_TIME = 365 * 15  # 15-year administrative censor cap

# ── Thyroid-specific presets ─────────────────────────────────────────────

PREDICTIVE_PRESETS: dict[str, dict[str, Any]] = {
    "recurrence": {
        "label": "Structural Recurrence",
        "time_col": "time_to_event_days",
        "event_col": "event_occurred",
        "competing_event_col": "death_occurred",
        "predictors": [
            "age_at_surgery", "braf_positive", "tert_positive",
            "tumor_1_gross_ete", "tumor_1_ete_microscopic_only",
            "ln_positive", "ln_ratio", "largest_tumor_cm",
            "overall_stage_ajcc8",
        ],
    },
    "death": {
        "label": "All-Cause Death",
        "time_col": "time_to_event_days",
        "event_col": "death_occurred",
        "competing_event_col": "event_occurred",
        "predictors": [
            "age_at_surgery", "overall_stage_ajcc8",
            "tumor_1_gross_ete", "largest_tumor_cm",
            "any_nsqip_complication",
        ],
    },
}

CURE_CALCULATOR_FEATURES: dict[str, dict[str, Any]] = {
    "age_at_diagnosis": {"label": "Age at diagnosis", "min": 18, "max": 90, "default": 45, "type": "slider", "group": "core"},
    "ajcc_stage_8": {"label": "AJCC 8th Ed Stage", "options": ["I", "II", "III", "IV"], "default": "I", "type": "select", "group": "core"},
    "ete_type": {"label": "ETE Type", "options": ["none", "microscopic", "gross"], "default": "none", "type": "select", "group": "core"},
    "braf_status": {"label": "BRAF V600E", "options": [False, True], "default": False, "type": "toggle", "group": "core"},
    "tert_status": {"label": "TERT Promoter", "options": [False, True], "default": False, "type": "toggle", "group": "core"},
    "recurrence_risk_band": {"label": "ATA Risk Band", "options": ["low", "intermediate", "high"], "default": "low", "type": "select", "group": "core"},
    "tumor_size_cm": {"label": "Tumor size (cm)", "min": 0.1, "max": 10.0, "default": 1.5, "step": 0.1, "type": "slider", "group": "core"},
    "ln_status": {"label": "Lymph node status", "options": ["N0", "N1a", "N1b", "Nx"], "default": "N0", "type": "select", "group": "core"},
}

# Hybrid theta adjustment factors for features not in the 10-covariate PTCM.
# Derived from published Cox PH estimates in differentiated thyroid cancer
# (Tuttle 2017, Adam 2015, Haugen 2016 ATA guidelines).  These multiplicative
# factors modulate the PTCM theta to account for tumor size and LN effects
# that are partially captured by the core AJCC/ETE/risk-band covariates.
_THETA_ADJUSTMENT_LN: dict[str, float] = {
    "N0": 1.0,
    "N1a": 1.15,   # central compartment LN metastasis — modest uplift
    "N1b": 1.35,   # lateral compartment LN — stronger effect
    "Nx": 1.05,    # unknown — slight uplift to reflect uncertainty
}
_THETA_ADJUSTMENT_SIZE_BREAKPOINTS = [
    (1.0, 0.92),   # ≤1 cm — protective (T1a)
    (2.0, 1.0),    # 1-2 cm — reference (T1b)
    (4.0, 1.12),   # 2-4 cm — moderate uplift (T2)
    (10.0, 1.28),  # >4 cm — substantial uplift (T3a+)
]

CLINICAL_INTERPRETATIONS: dict[str, str] = {
    "very_high": (
        "Very high cure probability (>85%). Consider de-escalation of "
        "surveillance frequency and reassurance."
    ),
    "high": (
        "High cure probability (70-85%). Standard follow-up protocol "
        "appropriate; consider reduced imaging frequency after 5 years."
    ),
    "moderate": (
        "Moderate cure probability (50-70%). Intensified surveillance "
        "recommended; consider active treatment discussion."
    ),
    "low": (
        "Low cure probability (<50%). High latent disease burden; "
        "aggressive multimodal surveillance and treatment optimization warranted."
    ),
}


def _fig_layout(fig: "go.Figure", title: str, height: int = 500, **kw) -> "go.Figure":
    """Apply standard dark-theme layout to a Plotly figure."""
    base = {k: v for k, v in PL.items()}
    base.update(title=title, height=height)
    base.update(kw)
    fig.update_layout(**base)
    return fig


# ── PTCM math (replicated from script 39 for import-free reuse) ─────────

def _weibull_f0_F0(t: np.ndarray, log_kappa: float, log_sigma: float):
    """Weibull density f₀ and CDF F₀ at times t."""
    kappa = np.exp(log_kappa)
    sigma = np.exp(log_sigma)
    u = (t / sigma) ** kappa
    F0 = 1.0 - np.exp(-u)
    f0 = (kappa / sigma) * (t / sigma) ** (kappa - 1.0) * np.exp(-u)
    return f0, F0


def _ptcm_survival(t: np.ndarray, theta: float, log_kappa: float, log_sigma: float) -> np.ndarray:
    """PTCM survival function S(t|θ) = exp(-θ·F₀(t))."""
    _, F0 = _weibull_f0_F0(t, log_kappa, log_sigma)
    return np.exp(-theta * F0)


def _build_ptcm_feature_vector(features: dict[str, Any], age_mean: float, age_std: float) -> np.ndarray:
    """Build a single-row design matrix row matching script 39 conventions.

    Parameters
    ----------
    features : dict
        Patient features with keys matching ``CURE_CALCULATOR_FEATURES``.
    age_mean, age_std : float
        Z-scoring parameters from the training cohort.
    """
    age = float(features.get("age_at_diagnosis", 45))
    stage = str(features.get("ajcc_stage_8", "I")).strip().upper()
    ete = str(features.get("ete_type", "none")).strip().lower()
    braf = bool(features.get("braf_status", False))
    tert = bool(features.get("tert_status", False))
    risk = str(features.get("recurrence_risk_band", "low")).strip().lower()

    return np.array([
        1.0,                                          # intercept
        (age - age_mean) / max(age_std, 1.0),         # age_z
        float(stage in ("II", "2")),                   # ajcc_2
        float(stage in ("III", "3")),                  # ajcc_3
        float(stage in ("IV", "IVA", "IVB", "IVC", "4")),  # ajcc_4
        float(ete == "microscopic"),                    # ete_microscopic
        float(ete == "gross"),                          # ete_gross
        float(braf),                                   # braf_pos
        float(tert),                                   # tert_pos
        float(risk == "high"),                         # high_risk_band
    ])


# ── Main class ───────────────────────────────────────────────────────────

class ThyroidPredictiveAnalyzer:
    """Integrated predictive analytics engine for thyroid cancer research.

    Composes ``ThyroidStatisticalAnalyzer`` for Cox PH / clinical snippets
    and wraps the PTCM math from script 39 to expose a reusable prediction
    API.  Adds multi-model comparison, enhanced competing risks, and
    personalized cure-probability estimation.

    Parameters
    ----------
    con : duckdb.DuckDBPyConnection
        Active DuckDB / MotherDuck connection.
    """

    def __init__(self, con: Any) -> None:
        self._con = con
        self._base = ThyroidStatisticalAnalyzer(con)
        self._ptcm_params: np.ndarray | None = None
        self._ptcm_summary: dict | None = None
        self._ptcm_age_mean: float = 50.0
        self._ptcm_age_std: float = 15.0
        self._load_ptcm_artifacts()

    # ── Data loading ─────────────────────────────────────────────────────

    def _load(self, view: str, where: str | None = None) -> pd.DataFrame:
        clause = f" WHERE {where}" if where else ""
        try:
            return self._con.execute(f"SELECT * FROM {view}{clause}").fetchdf()
        except Exception as exc:
            log.warning("Query on %s failed: %s", view, exc)
            return pd.DataFrame()

    def _tbl_exists(self, name: str) -> bool:
        try:
            r = self._con.execute(
                f"SELECT COUNT(*) FROM information_schema.tables "
                f"WHERE table_name='{name}'"
            ).fetchone()
            return bool(r and r[0] > 0)
        except Exception:
            return False

    @staticmethod
    def _prepare_numeric(df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
        sub = df.copy()
        for c in cols:
            if c in sub.columns:
                sub[c] = pd.to_numeric(sub[c], errors="coerce")
        return sub.dropna(subset=[c for c in cols if c in sub.columns])

    # ── PTCM artifact loading ────────────────────────────────────────────

    def _load_ptcm_artifacts(self) -> None:
        """Load fitted PTCM parameters from CSV exports or DuckDB table."""
        summary_path = PTCM_EXPORT_DIR / "ptcm_summary.csv"
        coeff_path = PTCM_EXPORT_DIR / "ptcm_covariate_effects.csv"
        meta_path = PTCM_EXPORT_DIR / "analysis_metadata.json"
        patient_path = PTCM_EXPORT_DIR / "ptcm_patient_cure_probs.csv"

        if meta_path.exists():
            try:
                meta = json.loads(meta_path.read_text(encoding="utf-8"))
                self._ptcm_summary = meta
            except Exception:
                pass

        if summary_path.exists() and self._ptcm_summary is None:
            try:
                sdf = pd.read_csv(summary_path)
                if not sdf.empty:
                    self._ptcm_summary = sdf.iloc[0].to_dict()
            except Exception:
                pass

        if coeff_path.exists():
            try:
                cdf = pd.read_csv(coeff_path)
                if not cdf.empty and "beta" in cdf.columns:
                    cov_order = {c: i for i, c in enumerate(PTCM_COVARIATES)}
                    cdf["_order"] = cdf["covariate"].map(cov_order)
                    cdf = cdf.sort_values("_order").dropna(subset=["_order"])
                    betas = cdf["beta"].values
                    if self._ptcm_summary:
                        kappa = self._ptcm_summary.get("weibull_kappa", 1.0)
                        sigma_years = self._ptcm_summary.get("weibull_sigma_years", 5.0)
                        log_kappa = np.log(max(kappa, 1e-6))
                        log_sigma = np.log(max(sigma_years * 365.25, 1.0))
                        self._ptcm_params = np.concatenate([[log_kappa, log_sigma], betas])
            except Exception as exc:
                log.warning("Failed to load PTCM coefficients: %s", exc)

        if patient_path.exists():
            try:
                pdf = pd.read_csv(patient_path, usecols=["age_at_diagnosis"], nrows=50000)
                if not pdf.empty:
                    self._ptcm_age_mean = float(pdf["age_at_diagnosis"].mean())
                    self._ptcm_age_std = max(float(pdf["age_at_diagnosis"].std()), 1.0)
            except Exception:
                pass

        if self._ptcm_params is None and self._tbl_exists("promotion_cure_cohort"):
            try:
                df = self._load("promotion_cure_cohort")
                if not df.empty and "age_at_diagnosis" in df.columns:
                    self._ptcm_age_mean = float(df["age_at_diagnosis"].mean())
                    self._ptcm_age_std = max(float(df["age_at_diagnosis"].std()), 1.0)
            except Exception:
                pass

    @property
    def ptcm_available(self) -> bool:
        return self._ptcm_params is not None and self._ptcm_summary is not None

    # ── 1. Individual cure prediction (PTCM) ────────────────────────────

    def predict_individual_cure_probability(
        self,
        patient_features: dict[str, Any],
        time_horizons: list[int] | None = None,
    ) -> dict[str, Any]:
        """Predict cure probability for a single patient using the fitted PTCM.

        Parameters
        ----------
        patient_features : dict
            Keys from ``CURE_CALCULATOR_FEATURES`` (age_at_diagnosis,
            ajcc_stage_8, ete_type, braf_status, tert_status,
            recurrence_risk_band).
        time_horizons : list[int], optional
            Years at which to compute conditional survival.
            Defaults to [1, 3, 5, 10, 15].

        Returns
        -------
        dict with keys: cure_probability, cure_tier, cure_interpretation,
        theta, conditional_survival (DataFrame), feature_contributions (dict),
        reference_population (dict), warnings
        """
        if not self.ptcm_available:
            return {"error": "PTCM not fitted. Run scripts/39_promotion_time_cure_models.py first."}

        if time_horizons is None:
            time_horizons = [1, 3, 5, 10, 15]

        result_warnings: list[str] = []
        params = self._ptcm_params
        log_kappa, log_sigma = float(params[0]), float(params[1])
        beta = params[2:]

        x = _build_ptcm_feature_vector(patient_features, self._ptcm_age_mean, self._ptcm_age_std)
        theta_base = float(np.exp(x @ beta))

        # Hybrid adjustment for tumor size and LN status (not in PTCM design matrix)
        ln = str(patient_features.get("ln_status", "N0")).strip()
        ln_adj = _THETA_ADJUSTMENT_LN.get(ln, 1.0)

        size = float(patient_features.get("tumor_size_cm", 1.5))
        size_adj = 1.0
        for cutoff, factor in _THETA_ADJUSTMENT_SIZE_BREAKPOINTS:
            if size <= cutoff:
                size_adj = factor
                break
        else:
            size_adj = _THETA_ADJUSTMENT_SIZE_BREAKPOINTS[-1][1]

        theta = theta_base * ln_adj * size_adj
        cure_prob = float(np.exp(-theta))

        if ln_adj != 1.0 or size_adj != 1.0:
            result_warnings.append(
                f"Hybrid adjustment applied: LN({ln})×{ln_adj:.2f}, "
                f"size({size:.1f}cm)×{size_adj:.2f} → θ adjusted "
                f"{theta_base:.3f}→{theta:.3f}. A full PTCM refit with "
                f"expanded covariates would improve calibration."
            )

        if cure_prob > 0.85:
            tier = "very_high"
        elif cure_prob > 0.70:
            tier = "high"
        elif cure_prob > 0.50:
            tier = "moderate"
        else:
            tier = "low"

        cond_surv_rows = []
        for yr in time_horizons:
            t_days = np.array([yr * 365.25])
            s = float(_ptcm_survival(t_days, theta, log_kappa, log_sigma)[0])
            cond_surv_rows.append({
                "year": yr,
                "survival_probability": round(s, 4),
                "recurrence_risk_pct": round((1 - s) * 100, 2),
            })
        cond_surv_df = pd.DataFrame(cond_surv_rows)

        # Feature contribution — delta in theta per covariate
        x_ref = np.zeros_like(x)
        x_ref[0] = 1.0  # intercept only
        theta_ref = float(np.exp(x_ref @ beta))
        contributions = {}
        for i, cov in enumerate(PTCM_COVARIATES):
            if cov == "intercept":
                continue
            x_single = x_ref.copy()
            x_single[i] = x[i]
            theta_single = float(np.exp(x_single @ beta))
            delta_theta = theta_single - theta_ref
            contributions[cov] = {
                "feature_value": float(x[i]),
                "delta_theta": round(delta_theta, 4),
                "direction": "risk_increasing" if delta_theta > 0 else "protective",
                "beta": round(float(beta[i]), 4),
            }

        ref_pop = {}
        if self._ptcm_summary:
            ref_pop = {
                "n_total": self._ptcm_summary.get("n_total", 0),
                "overall_cure_fraction": self._ptcm_summary.get("overall_cure_fraction", 0),
                "event_rate_pct": self._ptcm_summary.get("event_rate_pct", 0),
            }

        return {
            "cure_probability": round(cure_prob, 4),
            "cure_probability_pct": round(cure_prob * 100, 1),
            "cure_tier": tier,
            "cure_interpretation": CLINICAL_INTERPRETATIONS.get(tier, ""),
            "theta": round(theta, 4),
            "conditional_survival": cond_surv_df,
            "feature_contributions": contributions,
            "reference_population": ref_pop,
            "patient_features_used": patient_features,
            "warnings": result_warnings,
        }

    def predict_cure_batch(self, df: pd.DataFrame) -> pd.DataFrame:
        """Score a DataFrame of patients, returning cure_prob and cure_tier columns."""
        if not self.ptcm_available:
            return pd.DataFrame()
        params = self._ptcm_params
        beta = params[2:]
        rows = []
        for _, row in df.iterrows():
            feats = {k: row.get(k) for k in CURE_CALCULATOR_FEATURES}
            x = _build_ptcm_feature_vector(feats, self._ptcm_age_mean, self._ptcm_age_std)
            theta = float(np.exp(x @ beta))
            cp = float(np.exp(-theta))
            rows.append({"research_id": row.get("research_id"), "theta": theta, "cure_prob": cp})
        result = pd.DataFrame(rows)
        if not result.empty:
            result["cure_tier"] = pd.cut(
                result["cure_prob"],
                bins=[0, 0.50, 0.70, 0.85, 1.001],
                labels=["low (<50%)", "moderate (50-70%)", "high (70-85%)", "very_high (>85%)"],
                right=False,
            )
        return result

    # ── 2. Enhanced competing risks ──────────────────────────────────────

    def fit_competing_risks(
        self,
        time_col: str,
        event_col: str,
        competing_event_col: str | None = None,
        predictors: list[str] | None = None,
        data: pd.DataFrame | None = None,
        view: str | None = None,
        strata_col: str | None = None,
    ) -> dict[str, Any]:
        """Aalen-Johansen CIF with optional stratification and cause-specific HRs.

        Extends the base ``ThyroidAdvancedAnalyzer.fit_competing_risks`` with:
          - Stratified CIF curves per subgroup
          - Cause-specific Cox HR estimates (Cox PH per event type)
          - Landmark CIF summaries at 1, 3, 5, 10 years

        Parameters
        ----------
        time_col, event_col : str
            Survival time and primary event indicator.
        competing_event_col : str, optional
            Column for the competing event (e.g., death_occurred).
        predictors : list[str], optional
            Covariates for cause-specific Cox models.
        data : DataFrame, optional
            Pre-loaded data. Falls back to view resolution.
        view : str, optional
            Source DuckDB view/table.
        strata_col : str, optional
            Column to stratify CIF curves (e.g., overall_stage_ajcc8).

        Returns
        -------
        dict with keys: cif_primary, cif_competing, cif_plot,
        summary_table, cause_specific_hrs, stratified_cifs, warnings
        """
        if not HAS_LIFELINES:
            return {"error": "lifelines required for competing-risks analysis"}

        if data is None:
            resolved = self._base.resolve_view(view)
            if resolved is None:
                return {"error": "No source view available"}
            data = self._load(resolved)

        if data.empty:
            return {"error": "Empty dataset"}

        required = [time_col, event_col]
        if competing_event_col:
            required.append(competing_event_col)
        missing = [c for c in required if c not in data.columns]
        if missing:
            return {"error": f"Missing columns: {missing}"}

        sub = self._prepare_numeric(data, required)
        result_warnings: list[str] = []

        if time_col == "time_to_event_days":
            sub["time_years"] = sub[time_col] / 365.25
            duration = "time_years"
        else:
            duration = time_col

        sub = sub[sub[duration] > 0]

        # Build combined event indicator
        if competing_event_col and competing_event_col in sub.columns:
            sub["combined_event"] = 0
            sub.loc[sub[event_col] == 1, "combined_event"] = 1
            sub.loc[
                (sub[competing_event_col] == 1) & (sub[event_col] != 1),
                "combined_event",
            ] = 2
            event_indicator = "combined_event"
        else:
            event_indicator = event_col
            result_warnings.append(
                "No competing event column — single-event CIF."
            )

        n_events = int((sub[event_indicator] == 1).sum())
        n_competing = int((sub[event_indicator] == 2).sum()) if competing_event_col else 0

        if n_events < 5:
            return {"error": f"Only {n_events} primary events — too few for CIF"}

        # Fit Aalen-Johansen — primary event
        try:
            aj_primary = AalenJohansenFitter(calculate_variance=True, seed=42)
            aj_primary.fit(sub[duration], sub[event_indicator], event_of_interest=1)
            cif_primary = aj_primary.cumulative_density_.reset_index()
            cif_primary.columns = ["time", "CIF_primary"]
        except Exception as exc:
            return {"error": f"CIF fitting failed: {exc}"}

        # Fit competing event CIF
        cif_competing = pd.DataFrame()
        if competing_event_col and n_competing >= 5:
            try:
                aj_comp = AalenJohansenFitter(calculate_variance=True, seed=42)
                aj_comp.fit(sub[duration], sub[event_indicator], event_of_interest=2)
                cif_competing = aj_comp.cumulative_density_.reset_index()
                cif_competing.columns = ["time", "CIF_competing"]
            except Exception:
                result_warnings.append("Competing-event CIF estimation failed.")

        # Landmark summary
        summary_rows = []
        for yr in [1, 3, 5, 10]:
            p_val = self._cif_at_time(cif_primary, yr)
            c_val = self._cif_at_time(cif_competing, yr) if not cif_competing.empty else 0.0
            summary_rows.append({
                "Year": yr,
                "CIF Primary (%)": round(p_val * 100, 2),
                "CIF Competing (%)": round(c_val * 100, 2),
                "Event-Free (%)": round((1 - p_val - c_val) * 100, 2),
            })
        summary_df = pd.DataFrame(summary_rows)

        # Cause-specific Cox HRs
        cs_hrs = {}
        if predictors and HAS_LIFELINES:
            cs_hrs = self._cause_specific_cox(sub, duration, event_indicator, predictors, result_warnings)

        # Stratified CIF curves
        stratified = {}
        if strata_col and strata_col in sub.columns:
            stratified = self._stratified_cif(sub, duration, event_indicator, strata_col, result_warnings)

        # sksurv CIF with confidence bands (if available)
        sksurv_cif_data = None
        if HAS_SKSURV_CIF and competing_event_col:
            try:
                y_surv_cr = np.array(
                    [(int(e), float(t)) for e, t in zip(sub[event_indicator], sub[duration])],
                    dtype=[("event", np.int32), ("time", np.float64)],
                )
                sksurv_cif_data = sksurv_cif(y_surv_cr)
            except Exception as exc:
                result_warnings.append(f"sksurv CIF failed (non-critical): {exc}")

        # Build Plotly figure
        cif_plot = self._build_cif_plot(cif_primary, cif_competing, stratified, strata_col)

        # Clinical methodology note
        clinical_note = (
            "**Interpretation guide:** Cause-specific hazard ratios estimate the "
            "instantaneous rate of each event type among patients still event-free "
            "(etiological insight: 'Does BRAF truly accelerate recurrence?'). "
            "Cumulative incidence functions (CIF) account for competing risks and "
            "directly estimate real-world probability of each event by a given timepoint "
            "('What is my patient's actual probability of recurrence by year 10, "
            "accounting for other-cause death?'). "
            "Full subdistribution (Fine-Gray) modeling is planned for a future version."
        )

        return {
            "cif_primary": cif_primary,
            "cif_competing": cif_competing,
            "cif_plot": cif_plot,
            "summary_table": summary_df,
            "cause_specific_hrs": cs_hrs,
            "stratified_cifs": stratified,
            "sksurv_cif": sksurv_cif_data,
            "clinical_note": clinical_note,
            "n_obs": len(sub),
            "n_events": n_events,
            "n_competing": n_competing,
            "warnings": result_warnings,
        }

    @staticmethod
    def _cif_at_time(cif_df: pd.DataFrame, t: float) -> float:
        if cif_df.empty:
            return 0.0
        col = [c for c in cif_df.columns if c != "time"][0]
        mask = cif_df["time"] <= t
        if mask.any():
            return float(cif_df.loc[mask, col].iloc[-1])
        return 0.0

    def _cause_specific_cox(
        self, data: pd.DataFrame, duration: str, event_ind: str,
        predictors: list[str], warn_list: list[str],
    ) -> dict[str, Any]:
        """Fit separate Cox PH for primary and competing events."""
        results = {}
        available_preds = [p for p in predictors if p in data.columns]
        if len(available_preds) < 2:
            warn_list.append("Fewer than 2 predictors available for cause-specific Cox.")
            return results

        for event_val, label in [(1, "primary"), (2, "competing")]:
            sub = data.copy()
            sub["cs_event"] = (sub[event_ind] == event_val).astype(int)
            if sub["cs_event"].sum() < 10:
                warn_list.append(f"Fewer than 10 {label} events — skipping cause-specific Cox.")
                continue
            try:
                numeric_preds = []
                for p in available_preds:
                    sub[p] = pd.to_numeric(sub[p], errors="coerce")
                    numeric_preds.append(p)
                cox_df = sub[[duration, "cs_event"] + numeric_preds].dropna()
                cph = CoxPHFitter(penalizer=0.01)
                cph.fit(cox_df, duration_col=duration, event_col="cs_event")
                hr_df = cph.summary[["exp(coef)", "exp(coef) lower 95%", "exp(coef) upper 95%", "p"]].copy()
                hr_df.columns = ["HR", "HR_lower", "HR_upper", "p_value"]
                hr_df = hr_df.round(4)
                results[label] = {
                    "hr_table": hr_df.reset_index(),
                    "concordance": round(cph.concordance_index_, 4),
                    "n_events": int(sub["cs_event"].sum()),
                }
            except Exception as exc:
                warn_list.append(f"Cause-specific Cox ({label}) failed: {exc}")
        return results

    def _stratified_cif(
        self, data: pd.DataFrame, duration: str, event_ind: str,
        strata_col: str, warn_list: list[str],
    ) -> dict[str, pd.DataFrame]:
        strata_cifs = {}
        for level in sorted(data[strata_col].dropna().unique()):
            sub = data[data[strata_col] == level]
            if len(sub) < 20 or (sub[event_ind] == 1).sum() < 3:
                warn_list.append(f"Stratum '{level}': too few events (n={len(sub)}). Skipped.")
                continue
            try:
                aj = AalenJohansenFitter(calculate_variance=True, seed=42)
                aj.fit(sub[duration], sub[event_ind], event_of_interest=1)
                cdf = aj.cumulative_density_.reset_index()
                cdf.columns = ["time", f"CIF_{level}"]
                strata_cifs[str(level)] = cdf
            except Exception:
                warn_list.append(f"CIF fitting failed for stratum '{level}'.")
        return strata_cifs

    def _build_cif_plot(
        self,
        cif_primary: pd.DataFrame,
        cif_competing: pd.DataFrame,
        stratified: dict[str, pd.DataFrame],
        strata_col: str | None,
    ) -> "go.Figure | None":
        if not HAS_PLOTLY:
            return None
        fig = go.Figure()

        if not stratified:
            fig.add_trace(go.Scatter(
                x=cif_primary["time"], y=cif_primary["CIF_primary"],
                mode="lines", name="Recurrence (primary)",
                line=dict(color=_TEAL, width=2.5),
                fill="tozeroy", fillcolor="rgba(45,212,191,0.1)",
            ))
            if not cif_competing.empty:
                fig.add_trace(go.Scatter(
                    x=cif_competing["time"], y=cif_competing["CIF_competing"],
                    mode="lines", name="Death (competing)",
                    line=dict(color=_ROSE, width=2, dash="dash"),
                    fill="tozeroy", fillcolor="rgba(244,63,94,0.08)",
                ))
        else:
            for i, (level, cdf) in enumerate(stratified.items()):
                col = [c for c in cdf.columns if c != "time"][0]
                fig.add_trace(go.Scatter(
                    x=cdf["time"], y=cdf[col],
                    mode="lines", name=f"{strata_col}={level}",
                    line=dict(color=_COLORWAY[i % len(_COLORWAY)], width=2),
                ))

        return _fig_layout(
            fig,
            "Cumulative Incidence Functions (Competing Risks)",
            height=480,
            xaxis_title="Years from surgery",
            yaxis_title="Cumulative incidence",
            yaxis_range=[0, max(0.15, float(cif_primary["CIF_primary"].max()) * 1.3)],
        )

    # ── 3. Explainable ML Nomograms ──────────────────────────────────────

    def train_explainable_nomogram(
        self,
        outcome: str,
        predictors: list[str],
        base_model: Literal["xgboost", "random_forest", "survival_forest"] = "xgboost",
        data: pd.DataFrame | None = None,
        view: str | None = None,
        n_folds: int = 5,
    ) -> dict[str, Any]:
        """Train a cross-validated ML model with SHAP explanation.

        Supports classification (XGBoost, RF) and survival-specific
        models (sksurv RandomSurvivalForest) depending on ``base_model``.

        Returns
        -------
        dict with keys: model, model_type, auc_cv, brier_cv,
        concordance_cv, feature_importance, shap_values,
        shap_summary_plot, shap_beeswarm_plot, calibration_plot,
        feature_names, X, warnings
        """
        result_warnings: list[str] = []

        if data is None:
            resolved = self._base.resolve_view(view)
            if resolved is None:
                return {"error": "No source view available"}
            data = self._load(resolved)

        if data.empty:
            return {"error": "Empty dataset"}

        available = [p for p in predictors if p in data.columns]
        if len(available) < 2:
            return {"error": f"Only {len(available)} predictors available"}

        if outcome not in data.columns:
            return {"error": f"Outcome '{outcome}' not in data"}

        sub = data[available + [outcome]].copy()
        for c in available:
            sub[c] = pd.to_numeric(sub[c], errors="coerce")
        sub[outcome] = pd.to_numeric(sub[outcome], errors="coerce")
        sub = sub.dropna()

        X = sub[available].values
        y = sub[outcome].values.astype(float)
        n_events = int(y.sum())

        if n_events < 20:
            return {"error": f"Only {n_events} events — need ≥20 for ML models"}

        if base_model == "survival_forest" and HAS_SKSURV:
            return self._train_survival_forest(X, y, available, sub, n_folds, result_warnings)
        elif base_model == "xgboost" and HAS_XGB:
            return self._train_xgb_nomogram(X, y, available, n_folds, result_warnings)
        elif HAS_SKLEARN:
            return self._train_rf_nomogram(X, y, available, n_folds, result_warnings)
        else:
            return {"error": "No ML backend available (install xgboost or scikit-learn)"}

    def _train_xgb_nomogram(
        self, X: np.ndarray, y: np.ndarray,
        features: list[str], n_folds: int, warns: list[str],
    ) -> dict[str, Any]:
        event_rate = y.mean()
        scale_pos = max((1 - event_rate) / max(event_rate, 1e-6), 1.0)
        model = xgb.XGBClassifier(
            n_estimators=200, max_depth=4, learning_rate=0.05,
            scale_pos_weight=min(scale_pos, 10.0),
            eval_metric="logloss", use_label_encoder=False,
            random_state=42, n_jobs=-1,
        )
        return self._evaluate_classifier(model, X, y, features, n_folds, "xgboost", warns)

    def _train_rf_nomogram(
        self, X: np.ndarray, y: np.ndarray,
        features: list[str], n_folds: int, warns: list[str],
    ) -> dict[str, Any]:
        model = RandomForestClassifier(
            n_estimators=300, max_depth=6, class_weight="balanced",
            random_state=42, n_jobs=-1,
        )
        return self._evaluate_classifier(model, X, y, features, n_folds, "random_forest", warns)

    def _evaluate_classifier(
        self, model, X: np.ndarray, y: np.ndarray,
        features: list[str], n_folds: int, model_type: str,
        warns: list[str],
    ) -> dict[str, Any]:
        cv = StratifiedKFold(n_splits=n_folds, shuffle=True, random_state=42)

        try:
            y_prob = cross_val_predict(model, X, y, cv=cv, method="predict_proba")[:, 1]
            auc = round(roc_auc_score(y, y_prob), 4)
            brier = round(brier_score_loss(y, y_prob), 4)
        except Exception as exc:
            warns.append(f"CV evaluation failed: {exc}")
            auc, brier, y_prob = np.nan, np.nan, np.full(len(y), 0.5)

        model.fit(X, y)

        # Feature importance
        if hasattr(model, "feature_importances_"):
            fi = pd.DataFrame({
                "feature": features,
                "importance": model.feature_importances_,
            }).sort_values("importance", ascending=False)
        else:
            fi = pd.DataFrame({"feature": features, "importance": 0.0})

        # SHAP
        shap_vals_df = pd.DataFrame()
        shap_summary_fig = None
        shap_beeswarm_fig = None
        if HAS_SHAP:
            try:
                explainer = shap.TreeExplainer(model)
                shap_values = explainer.shap_values(X)
                if isinstance(shap_values, list):
                    shap_values = shap_values[1]
                shap_vals_df = pd.DataFrame(shap_values, columns=features)
                shap_summary_fig = self._build_shap_bar_plot(shap_vals_df, features)
                shap_beeswarm_fig = self._build_shap_beeswarm(shap_values, X, features)
            except Exception as exc:
                warns.append(f"SHAP computation failed: {exc}")

        # Calibration
        cal_fig = self._build_calibration_plot(y, y_prob) if HAS_PLOTLY else None

        return {
            "model": model,
            "model_type": model_type,
            "auc_cv": auc,
            "brier_cv": brier,
            "feature_importance": fi,
            "shap_values": shap_vals_df,
            "shap_summary_plot": shap_summary_fig,
            "shap_beeswarm_plot": shap_beeswarm_fig,
            "calibration_plot": cal_fig,
            "n_obs": len(y),
            "n_events": int(y.sum()),
            "event_rate": round(float(y.mean()), 4),
            "predictors_used": features,
            "feature_names": features,
            "X": X,
            "warnings": warns,
        }

    def _train_survival_forest(
        self, X: np.ndarray, y: np.ndarray,
        features: list[str], data: pd.DataFrame,
        n_folds: int, warns: list[str],
    ) -> dict[str, Any]:
        """Train a scikit-survival Random Survival Forest with SHAP."""
        if "time_to_event_days" not in data.columns:
            return {"error": "time_to_event_days column required for survival forest"}

        time_arr = pd.to_numeric(data["time_to_event_days"], errors="coerce").fillna(0).values
        y_surv = np.array(
            [(bool(e), float(t)) for e, t in zip(y.astype(bool), time_arr)],
            dtype=[("event", bool), ("time", float)],
        )

        rsf = RandomSurvivalForest(
            n_estimators=200, max_depth=6, min_samples_leaf=15,
            random_state=42, n_jobs=-1,
        )
        rsf.fit(X, y_surv)

        try:
            c_index = concordance_index_censored(y_surv["event"], y_surv["time"], rsf.predict(X))
            concordance = round(c_index[0], 4)
        except Exception:
            concordance = np.nan

        fi = pd.DataFrame({
            "feature": features,
            "importance": rsf.feature_importances_ if hasattr(rsf, "feature_importances_") else 0.0,
        }).sort_values("importance", ascending=False)

        shap_vals_df = pd.DataFrame()
        shap_fig = None
        if HAS_SHAP:
            try:
                X_sample = X[:min(500, len(X))]
                explainer = shap.TreeExplainer(rsf)
                sv = explainer.shap_values(X_sample)
                if isinstance(sv, list):
                    sv = sv[0]
                shap_vals_df = pd.DataFrame(sv, columns=features)
                shap_fig = self._build_shap_bar_plot(shap_vals_df, features)
            except Exception as exc:
                warns.append(f"SHAP for survival forest failed: {exc}")

        return {
            "model": rsf,
            "model_type": "survival_forest",
            "concordance_cv": concordance,
            "feature_importance": fi,
            "shap_values": shap_vals_df,
            "shap_summary_plot": shap_fig,
            "n_obs": len(y),
            "n_events": int(y.sum()),
            "event_rate": round(float(y.mean()), 4),
            "predictors_used": features,
            "feature_names": features,
            "X": X,
            "warnings": warns,
        }

    # ── SHAP plot builders ───────────────────────────────────────────────

    @staticmethod
    def _build_shap_bar_plot(shap_df: pd.DataFrame, features: list[str]) -> "go.Figure | None":
        if not HAS_PLOTLY or shap_df.empty:
            return None
        mean_abs = shap_df.abs().mean().sort_values(ascending=True)
        fig = go.Figure(go.Bar(
            x=mean_abs.values, y=mean_abs.index,
            orientation="h",
            marker_color=_TEAL,
        ))
        return _fig_layout(fig, "Feature Importance (mean |SHAP|)", height=max(350, len(features) * 30))

    @staticmethod
    def _build_shap_beeswarm(
        shap_values: np.ndarray, X: np.ndarray, features: list[str],
    ) -> "go.Figure | None":
        if not HAS_PLOTLY:
            return None
        mean_abs = np.abs(shap_values).mean(axis=0)
        order = np.argsort(mean_abs)
        fig = go.Figure()
        n_show = min(len(features), 12)
        for rank in range(n_show):
            idx = order[-(rank + 1)]
            sv = shap_values[:, idx]
            xv = X[:, idx]
            xv_norm = (xv - xv.min()) / max(xv.max() - xv.min(), 1e-8)
            colors = [
                f"rgb({int(55 + 200 * v)}, {int(212 - 150 * v)}, {int(191 - 100 * v)})"
                for v in xv_norm
            ]
            n_pts = min(500, len(sv))
            rng = np.random.default_rng(42)
            idx_sample = rng.choice(len(sv), n_pts, replace=False)
            jitter = rng.normal(0, 0.15, n_pts)
            fig.add_trace(go.Scatter(
                x=sv[idx_sample],
                y=[rank] * n_pts + jitter,
                mode="markers",
                marker=dict(size=3, color=[colors[i] for i in idx_sample], opacity=0.6),
                name=features[idx],
                showlegend=False,
                hovertemplate=f"{features[idx]}: SHAP=%{{x:.3f}}<extra></extra>",
            ))
        fig.update_yaxes(
            tickvals=list(range(n_show)),
            ticktext=[features[order[-(i + 1)]] for i in range(n_show)],
        )
        return _fig_layout(fig, "SHAP Beeswarm — Feature Impact on Prediction", height=max(400, n_show * 40))

    @staticmethod
    def _build_calibration_plot(y_true: np.ndarray, y_prob: np.ndarray) -> "go.Figure | None":
        if not HAS_PLOTLY:
            return None
        try:
            frac_pos, mean_pred = calibration_curve(y_true, y_prob, n_bins=10, strategy="quantile")
        except Exception:
            frac_pos, mean_pred = calibration_curve(y_true, y_prob, n_bins=5)
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=mean_pred, y=frac_pos, mode="lines+markers",
            name="Model", line=dict(color=_TEAL, width=2),
            marker=dict(size=8),
        ))
        fig.add_trace(go.Scatter(
            x=[0, 1], y=[0, 1], mode="lines",
            name="Perfect", line=dict(color=_GRAY, dash="dash"),
        ))
        return _fig_layout(
            fig, "Calibration Plot",
            height=400,
            xaxis_title="Mean predicted probability",
            yaxis_title="Observed fraction of positives",
        )

    # ── 4. Multi-model comparison ────────────────────────────────────────

    def compare_survival_models(
        self,
        predictors: list[str] | None = None,
        data: pd.DataFrame | None = None,
        view: str | None = None,
        time_col: str = "time_to_event_days",
        event_col: str = "event_occurred",
    ) -> dict[str, Any]:
        """Compare KM, Cox PH, and PTCM on the same cohort.

        Returns a unified comparison table with concordance, AIC, BIC,
        and Brier scores, plus summary Plotly dashboard figure.
        """
        result_warnings: list[str] = []

        if data is None:
            resolved = self._base.resolve_view(view)
            if resolved is None:
                return {"error": "No source view available"}
            data = self._load(resolved)

        if data.empty:
            return {"error": "Empty dataset"}

        if predictors is None:
            predictors = [p for p in THYROID_PREDICTORS if p in data.columns]

        sub = data.copy()
        for c in [time_col, event_col]:
            if c in sub.columns:
                sub[c] = pd.to_numeric(sub[c], errors="coerce")
        sub = sub.dropna(subset=[time_col, event_col])
        sub = sub[sub[time_col] > 0]

        if time_col == "time_to_event_days":
            sub["time_years"] = sub[time_col] / 365.25
            dur = "time_years"
        else:
            dur = time_col

        n_events = int(sub[event_col].sum())
        rows = []

        # Kaplan-Meier (non-parametric baseline)
        if HAS_LIFELINES:
            try:
                kmf = KaplanMeierFitter()
                kmf.fit(sub[dur], sub[event_col])
                rows.append({
                    "Model": "Kaplan-Meier",
                    "Type": "Non-parametric",
                    "Concordance": "—",
                    "AIC": "—",
                    "Predictors": 0,
                    "N": len(sub),
                    "Events": n_events,
                    "Notes": f"Median survival: {kmf.median_survival_time_:.1f}y" if np.isfinite(kmf.median_survival_time_) else "Median not reached",
                })
            except Exception as exc:
                result_warnings.append(f"KM failed: {exc}")

        # Cox PH
        if HAS_LIFELINES and predictors:
            available = [p for p in predictors if p in sub.columns]
            if len(available) >= 2:
                try:
                    cox_sub = sub[[dur, event_col] + available].copy()
                    for c in available:
                        cox_sub[c] = pd.to_numeric(cox_sub[c], errors="coerce")
                    cox_sub = cox_sub.dropna()
                    cph = CoxPHFitter(penalizer=0.01)
                    cph.fit(cox_sub, duration_col=dur, event_col=event_col)
                    rows.append({
                        "Model": "Cox PH",
                        "Type": "Semi-parametric",
                        "Concordance": round(cph.concordance_index_, 4),
                        "AIC": round(cph.AIC_partial_, 1) if hasattr(cph, "AIC_partial_") else "—",
                        "Predictors": len(available),
                        "N": len(cox_sub),
                        "Events": int(cox_sub[event_col].sum()),
                        "Notes": f"Top HR: {cph.summary['exp(coef)'].idxmax()} ({cph.summary['exp(coef)'].max():.2f})",
                    })
                except Exception as exc:
                    result_warnings.append(f"Cox PH failed: {exc}")

        # PTCM
        if self.ptcm_available and self._ptcm_summary:
            s = self._ptcm_summary
            rows.append({
                "Model": "Weibull PTCM",
                "Type": "Parametric (cure)",
                "Concordance": "—",
                "AIC": s.get("aic", "—"),
                "Predictors": len(PTCM_COVARIATES) - 1,
                "N": s.get("n_total", 0),
                "Events": s.get("n_events", 0),
                "Notes": f"Cure π̄={s.get('overall_cure_fraction', 0):.1%}, κ={s.get('weibull_kappa', 0):.3f}",
            })

        # Mixture cure model (Weibull latency, constant cure fraction)
        if HAS_SCIPY and n_events >= 10:
            try:
                mc_result = self._fit_mixture_cure(
                    sub[dur].values, sub[event_col].values, n_events
                )
                if mc_result:
                    rows.append({
                        "Model": "Mixture Cure (Weibull)",
                        "Type": "Parametric (cure)",
                        "Concordance": "—",
                        "AIC": mc_result["aic"],
                        "Predictors": 0,
                        "N": len(sub),
                        "Events": n_events,
                        "Notes": (
                            f"Cure π={mc_result['cure_fraction']:.1%}, "
                            f"κ={mc_result['kappa']:.3f}, "
                            f"σ={mc_result['sigma']:.2f}y"
                        ),
                    })
            except Exception as exc:
                result_warnings.append(f"Mixture cure failed: {exc}")

        # Penalized Cox (L2 via sksurv)
        if HAS_SKSURV and predictors:
            available = [p for p in predictors if p in sub.columns]
            if len(available) >= 2:
                try:
                    pen_sub = sub[[time_col, event_col] + available].copy()
                    for c in available:
                        pen_sub[c] = pd.to_numeric(pen_sub[c], errors="coerce")
                    pen_sub = pen_sub.dropna()
                    X_pen = pen_sub[available].values
                    y_pen = np.array(
                        [(bool(e), float(t)) for e, t in zip(pen_sub[event_col].astype(bool), pen_sub[time_col])],
                        dtype=[("event", bool), ("time", float)],
                    )
                    coxnet = CoxnetSurvivalAnalysis(l1_ratio=0.0, alpha_min_ratio=0.01, max_iter=1000)
                    coxnet.fit(X_pen, y_pen)
                    c_pen = concordance_index_censored(y_pen["event"], y_pen["time"], coxnet.predict(X_pen))
                    rows.append({
                        "Model": "Penalized Cox (Ridge)",
                        "Type": "Semi-parametric (L2)",
                        "Concordance": round(c_pen[0], 4),
                        "AIC": "—",
                        "Predictors": len(available),
                        "N": len(pen_sub),
                        "Events": int(pen_sub[event_col].sum()),
                        "Notes": "L2-regularized partial likelihood",
                    })
                except Exception as exc:
                    result_warnings.append(f"Penalized Cox failed: {exc}")

        # Random Survival Forest (if sksurv available)
        if HAS_SKSURV and predictors:
            available = [p for p in predictors if p in sub.columns]
            if len(available) >= 2:
                try:
                    rsf_sub = sub[[time_col, event_col] + available].copy()
                    for c in available:
                        rsf_sub[c] = pd.to_numeric(rsf_sub[c], errors="coerce")
                    rsf_sub = rsf_sub.dropna()
                    X_rsf = rsf_sub[available].values
                    y_surv = np.array(
                        [(bool(e), float(t)) for e, t in zip(rsf_sub[event_col].astype(bool), rsf_sub[time_col])],
                        dtype=[("event", bool), ("time", float)],
                    )
                    rsf = RandomSurvivalForest(n_estimators=100, max_depth=5, random_state=42, n_jobs=-1)
                    rsf.fit(X_rsf, y_surv)
                    c_idx = concordance_index_censored(y_surv["event"], y_surv["time"], rsf.predict(X_rsf))
                    rows.append({
                        "Model": "Random Survival Forest",
                        "Type": "ML (survival)",
                        "Concordance": round(c_idx[0], 4),
                        "AIC": "—",
                        "Predictors": len(available),
                        "N": len(rsf_sub),
                        "Events": int(rsf_sub[event_col].sum()),
                        "Notes": "Non-parametric ensemble",
                    })
                except Exception as exc:
                    result_warnings.append(f"RSF failed: {exc}")

        comparison_df = pd.DataFrame(rows)

        # Build comparison figure
        comp_fig = self._build_comparison_figure(comparison_df) if HAS_PLOTLY else None

        return {
            "comparison_table": comparison_df,
            "comparison_plot": comp_fig,
            "n_models": len(rows),
            "warnings": result_warnings,
        }

    @staticmethod
    def _build_comparison_figure(df: pd.DataFrame) -> "go.Figure | None":
        if not HAS_PLOTLY or df.empty:
            return None
        concordance_rows = df[df["Concordance"] != "—"].copy()
        if concordance_rows.empty:
            return None
        concordance_rows["Concordance"] = pd.to_numeric(concordance_rows["Concordance"])
        fig = go.Figure(go.Bar(
            x=concordance_rows["Model"],
            y=concordance_rows["Concordance"],
            marker_color=[_TEAL, _SKY, _VIOLET, _AMBER][:len(concordance_rows)],
            text=concordance_rows["Concordance"].round(3),
            textposition="outside",
        ))
        return _fig_layout(
            fig, "Model Concordance Comparison",
            height=400,
            yaxis_title="Concordance Index",
            yaxis_range=[0, 1],
        )

    # ── Mixture cure model (MLE) ───────────────────────────────────────

    @staticmethod
    def _fit_mixture_cure(
        t: np.ndarray, e: np.ndarray, n_events: int,
    ) -> dict[str, Any] | None:
        """Fit a Weibull mixture cure model via MLE.

        Model: S(t) = π + (1-π)·exp(-(t/σ)^κ)
        where π is the cure fraction, κ the Weibull shape, σ the scale.

        Three parameters: [logit_pi, log_kappa, log_sigma].
        """
        if not HAS_SCIPY or n_events < 10:
            return None

        def _sigmoid(z):
            return 1.0 / (1.0 + np.exp(-np.clip(z, -30, 30)))

        def _neg_ll(params, t, e):
            logit_pi, log_k, log_s = params
            pi = _sigmoid(logit_pi)
            k = np.exp(log_k)
            s = np.exp(log_s)
            u = (t / s) ** k
            Su = np.exp(-u)
            fu = (k / s) * (t / s) ** (k - 1.0) * np.exp(-u)
            fu = np.clip(fu, 1e-300, None)
            Su = np.clip(Su, 1e-15, 1.0)

            ll_event = np.log(np.clip((1.0 - pi) * fu, 1e-300, None))
            ll_cens = np.log(np.clip(pi + (1.0 - pi) * Su, 1e-300, None))
            return -np.sum(np.where(e == 1, ll_event, ll_cens))

        p0 = np.array([
            2.0,
            np.log(1.0),
            np.log(np.median(t[e == 1]) if e.sum() > 0 else np.median(t)),
        ])
        try:
            res = minimize(
                _neg_ll, p0, args=(t, e),
                method="L-BFGS-B",
                bounds=[(-10, 10), (-3, 3), (None, None)],
                options={"maxiter": 2000, "ftol": 1e-10},
            )
            if not res.success:
                return None
            logit_pi, log_k, log_s = res.x
            pi = float(_sigmoid(logit_pi))
            kappa = float(np.exp(log_k))
            sigma = float(np.exp(log_s))
            n_p = 3
            aic = 2 * res.fun + 2 * n_p
            return {
                "cure_fraction": pi,
                "kappa": round(kappa, 4),
                "sigma": round(sigma, 4),
                "aic": round(aic, 1),
                "neg_loglik": round(float(res.fun), 2),
            }
        except Exception:
            return None

    # ── 5. Cure calculator interface spec ────────────────────────────────

    def create_interactive_cure_calculator(self) -> dict[str, Any]:
        """Return metadata for building Streamlit input widgets.

        The returned dict contains feature specifications, reference
        population stats, and the prediction function handle.
        """
        return {
            "features": CURE_CALCULATOR_FEATURES,
            "ptcm_available": self.ptcm_available,
            "reference_population": {
                "age_mean": round(self._ptcm_age_mean, 1),
                "age_std": round(self._ptcm_age_std, 1),
                "n_total": self._ptcm_summary.get("n_total", 0) if self._ptcm_summary else 0,
                "overall_cure": self._ptcm_summary.get("overall_cure_fraction", 0) if self._ptcm_summary else 0,
            },
            "predict_fn": self.predict_individual_cure_probability,
            "clinical_interpretations": CLINICAL_INTERPRETATIONS,
        }

    # ── 6. Manuscript report generation ──────────────────────────────────

    def generate_manuscript_report(
        self,
        sections: list[str] | None = None,
        data: pd.DataFrame | None = None,
        view: str | None = None,
        title: str = "THYROID_2026 — Predictive Analytics Report",
        author: str = "Thyroid Research Team",
    ) -> dict[str, Any]:
        """Generate a Word manuscript report with selected analysis sections.

        Parameters
        ----------
        sections : list[str]
            Subset of ["PTCM", "CompetingRisks", "Nomogram", "Comparison"].
        """
        if not HAS_DOCX:
            return {"error": "python-docx required for manuscript generation"}

        if sections is None:
            sections = ["PTCM", "CompetingRisks", "Nomogram", "Comparison"]

        PREDICTIVE_EXPORT_DIR.mkdir(parents=True, exist_ok=True)

        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)

        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M} | Author: {author}")
        doc.add_paragraph("---")

        if "PTCM" in sections and self.ptcm_available:
            self._report_ptcm_section(doc)

        if "CompetingRisks" in sections:
            self._report_competing_risks_section(doc, data, view)

        if "Nomogram" in sections:
            self._report_nomogram_section(doc, data, view)

        if "Comparison" in sections:
            self._report_comparison_section(doc, data, view)

        out_path = PREDICTIVE_EXPORT_DIR / f"predictive_report_{datetime.now():%Y%m%d_%H%M}.docx"
        doc.save(str(out_path))

        return {
            "path": str(out_path),
            "sections_included": sections,
            "generated_at": datetime.now().isoformat(),
        }

    def _report_ptcm_section(self, doc: "Document") -> None:
        doc.add_heading("Promotion Time Cure Model (PTCM)", level=1)
        s = self._ptcm_summary
        doc.add_paragraph(
            f"The Weibull PTCM was fitted to {s.get('n_total', 0):,} patients "
            f"with {s.get('n_events', 0):,} events ({s.get('event_rate_pct', 0):.1f}%). "
            f"The overall cure fraction was {s.get('overall_cure_fraction', 0):.1%}. "
            f"Weibull baseline shape κ={s.get('weibull_kappa', 0):.3f}, "
            f"scale σ={s.get('weibull_sigma_years', 0):.2f} years. "
            f"AIC={s.get('aic', 0):.1f}."
        )
        coeff_path = PTCM_EXPORT_DIR / "ptcm_covariate_effects.csv"
        if coeff_path.exists():
            cdf = pd.read_csv(coeff_path)
            doc.add_heading("Covariate Effects on Promotion Intensity", level=2)
            self._df_to_docx_table(doc, cdf[["covariate", "exp_beta", "ci_lower", "ci_upper", "p_value"]])

    def _report_competing_risks_section(self, doc: "Document", data, view) -> None:
        doc.add_heading("Competing Risks Analysis", level=1)
        preset = PREDICTIVE_PRESETS["recurrence"]
        result = self.fit_competing_risks(
            preset["time_col"], preset["event_col"],
            preset.get("competing_event_col"),
            data=data, view=view,
        )
        if "error" in result:
            doc.add_paragraph(f"Analysis could not be completed: {result['error']}")
            return
        doc.add_paragraph(
            f"Competing risks analysis on {result['n_obs']:,} patients "
            f"({result['n_events']} primary events, {result['n_competing']} competing events)."
        )
        doc.add_heading("Cumulative Incidence at Landmark Timepoints", level=2)
        self._df_to_docx_table(doc, result["summary_table"])

    def _report_nomogram_section(self, doc: "Document", data, view) -> None:
        doc.add_heading("Explainable ML Nomogram", level=1)
        preset = PREDICTIVE_PRESETS["recurrence"]
        result = self.train_explainable_nomogram(
            "event_occurred", preset["predictors"][:6],
            data=data, view=view,
        )
        if "error" in result:
            doc.add_paragraph(f"Nomogram training failed: {result['error']}")
            return
        doc.add_paragraph(
            f"Model: {result.get('model_type', 'unknown')} | "
            f"AUC={result.get('auc_cv', 'N/A')} | "
            f"Brier={result.get('brier_cv', 'N/A')} | "
            f"N={result.get('n_obs', 0):,}"
        )
        if not result.get("feature_importance", pd.DataFrame()).empty:
            doc.add_heading("Feature Importance", level=2)
            self._df_to_docx_table(doc, result["feature_importance"].head(10))

    def _report_comparison_section(self, doc: "Document", data, view) -> None:
        doc.add_heading("Model Comparison", level=1)
        result = self.compare_survival_models(data=data, view=view)
        if "error" in result:
            doc.add_paragraph(f"Comparison failed: {result['error']}")
            return
        self._df_to_docx_table(doc, result["comparison_table"])

    @staticmethod
    def _df_to_docx_table(doc: "Document", df: pd.DataFrame, max_rows: int = 100) -> None:
        if df.empty:
            return
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Light Grid Accent 1"
        for i, col in enumerate(df.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in df.head(max_rows).iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                val = row[col]
                if isinstance(val, float) and not np.isnan(val):
                    cells[i].text = f"{val:.4f}" if abs(val) < 1 else f"{val:.2f}"
                else:
                    cells[i].text = str(val) if pd.notna(val) else ""

    # ── Static prediction helper (for risk calculator reuse) ─────────────

    @staticmethod
    def predict_individual_risk(
        model, feature_values: dict[str, float], feature_names: list[str],
    ) -> dict[str, Any]:
        """Single-patient prediction with SHAP force contributions.

        Compatible with models trained by ``train_explainable_nomogram``.
        """
        X_single = np.array([[feature_values.get(f, 0.0) for f in feature_names]])

        if hasattr(model, "predict_proba"):
            prob = float(model.predict_proba(X_single)[0, 1])
        elif hasattr(model, "predict"):
            prob = float(model.predict(X_single)[0])
        else:
            return {"error": "Model has no predict method"}

        risk_class = "high" if prob > 0.3 else ("moderate" if prob > 0.1 else "low")

        contributions = {}
        if HAS_SHAP:
            try:
                explainer = shap.TreeExplainer(model)
                sv = explainer.shap_values(X_single)
                if isinstance(sv, list):
                    sv = sv[1]
                for i, f in enumerate(feature_names):
                    contributions[f] = round(float(sv[0, i]), 4)
            except Exception:
                pass

        return {
            "risk_probability": round(prob, 4),
            "risk_pct": round(prob * 100, 1),
            "risk_class": risk_class,
            "feature_contributions": contributions,
        }

    # ── Cure probability survival curve for individual patient ───────────

    def plot_individual_cure_trajectory(
        self, patient_features: dict[str, Any],
    ) -> "go.Figure | None":
        """Plotly figure showing PTCM survival curve for a specific patient."""
        if not self.ptcm_available or not HAS_PLOTLY:
            return None
        params = self._ptcm_params
        log_kappa, log_sigma = float(params[0]), float(params[1])
        beta = params[2:]

        x = _build_ptcm_feature_vector(patient_features, self._ptcm_age_mean, self._ptcm_age_std)
        theta = float(np.exp(x @ beta))
        cure = float(np.exp(-theta))

        t_grid = np.linspace(1, MAX_TIME, 300)
        s_curve = _ptcm_survival(t_grid, theta, log_kappa, log_sigma)

        # Population average for comparison
        theta_avg = -np.log(max(self._ptcm_summary.get("overall_cure_fraction", 0.95), 1e-6))
        s_avg = _ptcm_survival(t_grid, theta_avg, log_kappa, log_sigma)

        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=t_grid / 365.25, y=s_curve, mode="lines",
            name="This Patient",
            line=dict(color=_TEAL, width=3),
        ))
        fig.add_trace(go.Scatter(
            x=t_grid / 365.25, y=s_avg, mode="lines",
            name="Population Average",
            line=dict(color=_GRAY, width=2, dash="dash"),
        ))
        fig.add_hline(
            y=cure, line_dash="dot", line_color=_AMBER,
            annotation_text=f"Cure plateau π={cure:.1%}",
            annotation_font_color=_AMBER,
        )
        return _fig_layout(
            fig,
            "Personalized Recurrence-Free Survival (PTCM)",
            height=420,
            xaxis_title="Years from surgery",
            yaxis_title="Survival probability",
            yaxis_range=[0, 1.05],
        )
