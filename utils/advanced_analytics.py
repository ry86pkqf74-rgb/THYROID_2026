"""
Advanced analytics for the THYROID_2026 research platform (Phase 3).

Competing-risks survival (Aalen-Johansen CIF), longitudinal mixed-effects
with stratification, explainable ML nomograms (XGBoost/RandomForest + SHAP),
interactive risk calculators, and automated manuscript report generation.

Composes with ``ThyroidStatisticalAnalyzer`` for shared presets and
plotting infrastructure.

Usage (standalone)::

    from utils.advanced_analytics import ThyroidAdvancedAnalyzer
    adv = ThyroidAdvancedAnalyzer(con)
    cif = adv.fit_competing_risks("time_days", "event", "death", ["age", "braf"])

Usage (dashboard)::

    Consumed by ``app/advanced_analytics.py`` which wires interactive
    controls around these methods.
"""
from __future__ import annotations

import io
import logging
import warnings
from datetime import datetime
from typing import Any, Literal

import numpy as np
import pandas as pd

try:
    import plotly.graph_objects as go
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

try:
    from lifelines import AalenJohansenFitter
    HAS_LIFELINES = True
except ImportError:
    HAS_LIFELINES = False

try:
    import statsmodels.formula.api as smf  # noqa: F401
    HAS_STATSMODELS = True
except ImportError:
    HAS_STATSMODELS = False

try:
    import xgboost as xgb
    HAS_XGB = True
except ImportError:
    HAS_XGB = False

try:
    from sklearn.ensemble import RandomForestClassifier
    from sklearn.model_selection import StratifiedKFold, cross_val_predict
    from sklearn.metrics import roc_auc_score, brier_score_loss
    from sklearn.calibration import calibration_curve
    HAS_SKLEARN = True
except ImportError:
    HAS_SKLEARN = False

try:
    import shap
    HAS_SHAP = True
except ImportError:
    HAS_SHAP = False

try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


from utils.statistical_analysis import (
    ThyroidStatisticalAnalyzer,
    THYROID_PREDICTORS,
    THYROID_SURVIVAL,
    PL,
)

log = logging.getLogger(__name__)

_TEAL = "#2dd4bf"
_SKY = "#38bdf8"
_VIOLET = "#a78bfa"
_AMBER = "#f59e0b"
_ROSE = "#f43f5e"
_GREEN = "#34d399"
_GRAY = "#4a5568"

# ── Thyroid-specific presets for advanced models ─────────────────────────

COMPETING_RISK_PRESETS: dict[str, dict[str, Any]] = {
    "recurrence_vs_death": {
        "label": "Recurrence (competing: death)",
        "time_col": "time_to_event_days",
        "event_col": "event_occurred",
        "competing_event_col": "death_occurred",
        "event_of_interest": 1,
        "predictors": ["age_at_surgery", "braf_positive", "ln_positive",
                        "largest_tumor_cm", "overall_stage_ajcc8"],
    },
}

NOMOGRAM_PRESETS: dict[str, dict[str, Any]] = {
    "recurrence": {
        "label": "Structural Recurrence",
        "outcome": "event_occurred",
        "predictors": ["age_at_surgery", "sex_numeric", "braf_positive",
                        "largest_tumor_cm", "ln_positive", "ln_ratio",
                        "tumor_1_ete_microscopic_only", "tumor_1_gross_ete"],
    },
    "nsqip_complication": {
        "label": "Any NSQIP Complication",
        "outcome": "any_nsqip_complication",
        "predictors": ["age_at_surgery", "sex_numeric", "braf_mutation_mentioned",
                        "largest_tumor_cm", "ln_positive"],
    },
}

LONGITUDINAL_STRATIFIERS: list[str] = [
    "braf_positive", "overall_stage_ajcc8",
    "recurrence_risk_band", "rai_received",
]

# ── Plot helpers ─────────────────────────────────────────────────────────

_COLORWAY = [_TEAL, _SKY, _VIOLET, _AMBER, _ROSE, _GREEN, "#fb923c"]


def _fig_layout(fig: "go.Figure", title: str, height: int = 500, **kw) -> "go.Figure":
    base = {k: v for k, v in PL.items()}
    base.update(title=title, height=height)
    base.update(kw)
    fig.update_layout(**base)
    return fig


# ── Core class ───────────────────────────────────────────────────────────

class ThyroidAdvancedAnalyzer:
    """Advanced analytics engine for thyroid cancer research.

    Extends capabilities beyond ``ThyroidStatisticalAnalyzer`` with
    competing-risks models, stratified longitudinal analysis, ML
    nomograms with SHAP explainability, and manuscript generation.

    Parameters
    ----------
    con : duckdb.DuckDBPyConnection
        Active DuckDB / MotherDuck connection (read-only sufficient).
    """

    def __init__(self, con: Any) -> None:
        self._con = con
        self._base = ThyroidStatisticalAnalyzer(con)

    # ── Data helpers ─────────────────────────────────────────────────────

    def _load(self, view: str, where: str | None = None) -> pd.DataFrame:
        clause = f" WHERE {where}" if where else ""
        try:
            return self._con.execute(f"SELECT * FROM {view}{clause}").fetchdf()
        except Exception as exc:
            log.warning("Query on %s failed: %s", view, exc)
            return pd.DataFrame()

    def _tbl_exists(self, name: str) -> bool:
        try:
            r = self._con.execute(
                f"SELECT COUNT(*) FROM information_schema.tables "
                f"WHERE table_name='{name}'"
            ).fetchone()
            return bool(r and r[0] > 0)
        except Exception:
            return False

    @staticmethod
    def _prepare_numeric(df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
        """Coerce columns to numeric, drop rows with any NaN in those cols."""
        sub = df.copy()
        for c in cols:
            if c in sub.columns:
                sub[c] = pd.to_numeric(sub[c], errors="coerce")
        return sub.dropna(subset=[c for c in cols if c in sub.columns])

    # ── 1. Competing Risks (Aalen-Johansen CIF) ──────────────────────────

    def fit_competing_risks(
        self,
        time_col: str,
        event_col: str,
        competing_event_col: str | None = None,
        predictors: list[str] | None = None,
        data: pd.DataFrame | None = None,
        view: str | None = None,
    ) -> dict[str, Any]:
        """Fit Aalen-Johansen cumulative incidence curves for competing risks.

        When ``competing_event_col`` is provided, constructs a combined
        event indicator (0=censored, 1=event of interest, 2=competing).
        Without it, treats all events as the primary event.

        Parameters
        ----------
        time_col, event_col : str
            Survival time and primary event indicator columns.
        competing_event_col : str, optional
            Competing event indicator column.
        predictors : list[str], optional
            Strata to produce subgroup CIF curves.
        data : DataFrame, optional
            Pre-loaded data. Falls back to *view* resolution.
        view : str, optional
            Source view/table.

        Returns
        -------
        dict with keys: cif_primary, cif_competing, cif_plot, summary_table,
        n_obs, n_events, n_competing, warnings
        """
        if not HAS_LIFELINES:
            return {"error": "lifelines required for competing-risks models"}

        if data is None:
            resolved = self._base.resolve_view(view)
            if resolved is None:
                return {"error": "No source view available"}
            data = self._load(resolved)

        if data.empty:
            return {"error": "Empty dataset"}

        required = [time_col, event_col]
        if competing_event_col:
            required.append(competing_event_col)
        missing = [c for c in required if c not in data.columns]
        if missing:
            return {"error": f"Missing columns: {missing}"}

        sub = self._prepare_numeric(data, required)

        if time_col == "time_to_event_days":
            sub["time_years"] = sub[time_col] / 365.25
            duration = "time_years"
        else:
            duration = time_col

        sub = sub[sub[duration] > 0]

        result_warnings: list[str] = []

        if competing_event_col and competing_event_col in sub.columns:
            sub["combined_event"] = 0
            sub.loc[sub[event_col] == 1, "combined_event"] = 1
            sub.loc[
                (sub[competing_event_col] == 1) & (sub[event_col] != 1),
                "combined_event",
            ] = 2
            event_indicator = "combined_event"
        else:
            event_indicator = event_col
            result_warnings.append(
                "No competing event column — fitting standard CIF (single-event)."
            )

        n_events = int((sub[event_indicator] == 1).sum())
        n_competing = int((sub[event_indicator] == 2).sum()) if competing_event_col else 0

        if n_events < 5:
            return {"error": f"Only {n_events} primary events — too few for CIF estimation"}

        # Fit Aalen-Johansen for primary event
        try:
            aj_primary = AalenJohansenFitter(calculate_variance=True, seed=42)
            aj_primary.fit(
                sub[duration],
                sub[event_indicator],
                event_of_interest=1,
            )
            cif_primary_df = aj_primary.cumulative_density_.copy()
            cif_primary_df.columns = ["CIF_primary"]
            cif_primary_df = cif_primary_df.reset_index()
            cif_primary_df.columns = ["time", "CIF_primary"]
        except Exception as exc:
            return {"error": f"Aalen-Johansen fitting failed: {exc}"}

        cif_competing_df = pd.DataFrame()
        if competing_event_col and n_competing >= 5:
            try:
                aj_competing = AalenJohansenFitter(calculate_variance=True, seed=42)
                aj_competing.fit(
                    sub[duration],
                    sub[event_indicator],
                    event_of_interest=2,
                )
                cif_competing_df = aj_competing.cumulative_density_.copy()
                cif_competing_df.columns = ["CIF_competing"]
                cif_competing_df = cif_competing_df.reset_index()
                cif_competing_df.columns = ["time", "CIF_competing"]
            except Exception:
                result_warnings.append("Competing-event CIF estimation failed.")

        # Landmark summaries
        summary_rows = []
        for t in [1, 3, 5, 10]:
            mask = cif_primary_df["time"] <= t
            if mask.any():
                cif_val = float(cif_primary_df.loc[mask, "CIF_primary"].iloc[-1])
            else:
                cif_val = 0.0

            comp_val = 0.0
            if not cif_competing_df.empty:
                mask_c = cif_competing_df["time"] <= t
                if mask_c.any():
                    comp_val = float(cif_competing_df.loc[mask_c, "CIF_competing"].iloc[-1])

            summary_rows.append({
                "Timepoint (years)": t,
                "CIF Primary (%)": round(cif_val * 100, 2),
                "CIF Competing (%)": round(comp_val * 100, 2),
                "1 - Both (%)": round((1 - cif_val - comp_val) * 100, 2),
            })
        summary_table = pd.DataFrame(summary_rows)

        # Plot
        fig = None
        if HAS_PLOTLY:
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=cif_primary_df["time"], y=cif_primary_df["CIF_primary"],
                mode="lines", name="Primary Event (CIF)",
                line=dict(color=_TEAL, width=2),
            ))
            if not cif_competing_df.empty:
                fig.add_trace(go.Scatter(
                    x=cif_competing_df["time"], y=cif_competing_df["CIF_competing"],
                    mode="lines", name="Competing Event (CIF)",
                    line=dict(color=_ROSE, width=2, dash="dash"),
                ))
            fig.add_hline(y=0.0, line_dash="dot", line_color=_GRAY, opacity=0.5)
            _fig_layout(fig, "Cumulative Incidence Functions (Aalen-Johansen)",
                        xaxis_title="Time (years)", yaxis_title="Cumulative Incidence",
                        yaxis_range=[0, min(1.0, max(0.2, cif_primary_df["CIF_primary"].max() * 1.3))])

        clinical_note = (
            f"At 5 years: {summary_rows[2]['CIF Primary (%)']:.1f}% cumulative incidence "
            f"of primary event"
        )
        if n_competing > 0:
            clinical_note += f", {summary_rows[2]['CIF Competing (%)']:.1f}% competing event"
        clinical_note += (
            ". Standard KM overestimates event probability when competing risks "
            "are present — CIF provides unbiased estimates."
        )

        return {
            "cif_primary": cif_primary_df,
            "cif_competing": cif_competing_df,
            "cif_plot": fig,
            "summary_table": summary_table,
            "n_obs": len(sub),
            "n_events": n_events,
            "n_competing": n_competing,
            "clinical_note": clinical_note,
            "warnings": result_warnings,
        }

    # ── 2. Stratified Longitudinal Mixed-Effects ──────────────────────────

    def fit_stratified_longitudinal(
        self,
        marker: str = "tg",
        stratify_by: str | None = None,
        view: str | None = None,
    ) -> dict[str, Any]:
        """Fit stratified mixed-effects models for biomarker trajectories.

        Extends ``ThyroidStatisticalAnalyzer.longitudinal_summary`` with
        subgroup stratification — separate slope estimates and trajectory
        plots for each stratum (e.g. BRAF+/- or stage grouping).

        Parameters
        ----------
        marker : str
            Biomarker key (``"tg"``, ``"tsh"``, ``"anti_tg"``).
        stratify_by : str, optional
            Column for subgroup analysis.
        view : str, optional
            Override source view.

        Returns
        -------
        dict with keys: overall_result, strata_results, trajectory_plot,
        comparison_table, warnings
        """
        overall = self._base.longitudinal_summary(marker=marker, view=view)
        if "error" in overall:
            return overall

        result: dict[str, Any] = {
            "overall_result": overall,
            "strata_results": {},
            "trajectory_plot": None,
            "comparison_table": pd.DataFrame(),
            "warnings": [],
        }

        if not stratify_by:
            if HAS_PLOTLY and "data" in overall and not overall["data"].empty:
                result["trajectory_plot"] = self._build_trajectory_plot(
                    overall["data"], overall.get("marker_label", marker), {}
                )
            return result

        data = overall.get("data", pd.DataFrame())
        if data.empty or stratify_by not in data.columns:
            # Try loading from patient-level features to merge stratifier
            for fv in ["risk_enriched_mv", "advanced_features_v3", "survival_cohort_enriched"]:
                if not self._tbl_exists(fv):
                    continue
                try:
                    strat_df = self._con.execute(
                        f"SELECT research_id, {stratify_by} FROM {fv} "
                        f"WHERE {stratify_by} IS NOT NULL"
                    ).fetchdf()
                    if not strat_df.empty:
                        data["research_id"] = pd.to_numeric(data["research_id"], errors="coerce")
                        strat_df["research_id"] = pd.to_numeric(strat_df["research_id"], errors="coerce")
                        data = data.merge(strat_df, on="research_id", how="inner")
                        break
                except Exception:
                    continue

        if stratify_by not in data.columns:
            result["warnings"].append(
                f"Stratification variable '{stratify_by}' not available in data."
            )
            return result

        strata_groups = data[stratify_by].dropna().unique()
        if len(strata_groups) < 2:
            result["warnings"].append(
                f"'{stratify_by}' has fewer than 2 groups — no stratification."
            )
            return result

        if len(strata_groups) > 10:
            result["warnings"].append(
                f"'{stratify_by}' has {len(strata_groups)} levels — showing top 5."
            )
            top5 = data[stratify_by].value_counts().head(5).index
            data = data[data[stratify_by].isin(top5)]
            strata_groups = top5

        comparison_rows = []
        strata_data_for_plot: dict[str, pd.DataFrame] = {}

        for stratum in sorted(strata_groups, key=str):
            stratum_data = data[data[stratify_by] == stratum]
            if len(stratum_data) < 20:
                continue

            strata_data_for_plot[str(stratum)] = stratum_data

            if not HAS_STATSMODELS:
                continue

            try:
                stratum_data = stratum_data.copy()
                stratum_data["days_scaled"] = stratum_data.get(
                    "days_from_surgery",
                    stratum_data.get("days_scaled", pd.Series(dtype=float))
                )
                if "days_scaled" not in stratum_data.columns or stratum_data["days_scaled"].isna().all():
                    continue

                stratum_data["days_scaled"] = stratum_data["days_scaled"] / 365.25
                stratum_data["research_id"] = stratum_data["research_id"].astype(str)

                if "outcome" not in stratum_data.columns:
                    continue

                obs_per_patient = stratum_data.groupby("research_id")["outcome"].count()
                multi_ids = obs_per_patient[obs_per_patient >= 2].index
                fit_data = stratum_data[stratum_data["research_id"].isin(multi_ids)]

                if len(fit_data) < 20 or fit_data["research_id"].nunique() < 5:
                    continue

                with warnings.catch_warnings(record=True):
                    warnings.simplefilter("always")
                    md = smf.mixedlm("outcome ~ days_scaled", fit_data,
                                     groups=fit_data["research_id"])
                    mdf = md.fit(reml=True, method="lbfgs")

                slope = float(mdf.params.get("days_scaled", np.nan))
                p_val = float(mdf.pvalues.get("days_scaled", np.nan))
                ci = mdf.conf_int()
                ci_lo = float(ci.loc["days_scaled", 0]) if "days_scaled" in ci.index else np.nan
                ci_hi = float(ci.loc["days_scaled", 1]) if "days_scaled" in ci.index else np.nan

                comparison_rows.append({
                    "stratum": str(stratum),
                    "n_patients": int(fit_data["research_id"].nunique()),
                    "n_obs": len(fit_data),
                    "slope_per_year": round(slope, 4),
                    "slope_CI_lower": round(ci_lo, 4),
                    "slope_CI_upper": round(ci_hi, 4),
                    "p_value": round(p_val, 6),
                    "direction": "rising" if slope > 0 else "falling",
                })

                result["strata_results"][str(stratum)] = {
                    "slope": slope, "p_value": p_val,
                    "ci": (ci_lo, ci_hi),
                    "n_patients": int(fit_data["research_id"].nunique()),
                }

            except Exception as exc:
                result["warnings"].append(
                    f"Mixed-effects failed for stratum {stratum}: {exc}"
                )

        if comparison_rows:
            result["comparison_table"] = pd.DataFrame(comparison_rows)

        if HAS_PLOTLY and strata_data_for_plot:
            result["trajectory_plot"] = self._build_trajectory_plot(
                data, overall.get("marker_label", marker),
                strata_data_for_plot, stratify_by,
            )

        return result

    @staticmethod
    def _build_trajectory_plot(
        data: pd.DataFrame,
        marker_label: str,
        strata_data: dict[str, pd.DataFrame],
        stratify_by: str | None = None,
    ) -> "go.Figure":
        """Build a multi-stratum spaghetti/trend plot for longitudinal data."""
        fig = go.Figure()
        time_col = "days_from_surgery" if "days_from_surgery" in data.columns else "days_scaled"

        if not strata_data:
            # Single population — median trend
            if time_col in data.columns and "outcome" in data.columns:
                sorted_d = data.sort_values(time_col)
                bins = pd.cut(sorted_d[time_col], bins=20)
                trend = sorted_d.groupby(bins, observed=True)["outcome"].median()
                x_mid = [(iv.left + iv.right) / 2 for iv in trend.index]
                fig.add_trace(go.Scatter(
                    x=x_mid, y=trend.values,
                    mode="lines+markers", name="Population Median",
                    line=dict(color=_TEAL, width=3),
                    marker=dict(size=6),
                ))
        else:
            for i, (stratum_name, sdf) in enumerate(strata_data.items()):
                color = _COLORWAY[i % len(_COLORWAY)]
                if time_col in sdf.columns and "outcome" in sdf.columns:
                    sorted_s = sdf.sort_values(time_col)
                    bins = pd.cut(sorted_s[time_col], bins=15)
                    trend = sorted_s.groupby(bins, observed=True)["outcome"].median()
                    x_mid = [(iv.left + iv.right) / 2 for iv in trend.index]
                    fig.add_trace(go.Scatter(
                        x=x_mid, y=trend.values,
                        mode="lines+markers",
                        name=f"{stratify_by}={stratum_name}",
                        line=dict(color=color, width=2),
                        marker=dict(size=5),
                    ))

        x_label = "Days from Surgery" if time_col == "days_from_surgery" else "Years"
        _fig_layout(fig, f"{marker_label} Trajectories", height=480,
                    xaxis_title=x_label, yaxis_title=f"{marker_label} (transformed)")
        return fig

    # ── 3. ML Nomogram (XGBoost / RandomForest + SHAP) ────────────────────

    def train_ml_nomogram(
        self,
        outcome: str,
        predictors: list[str],
        model_type: Literal["xgboost", "randomforest"] = "xgboost",
        data: pd.DataFrame | None = None,
        view: str | None = None,
        n_folds: int = 5,
    ) -> dict[str, Any]:
        """Train an explainable ML model for binary outcome prediction.

        Uses stratified k-fold cross-validation for unbiased performance
        estimation, then fits a final model on all data for SHAP values.

        Parameters
        ----------
        outcome : str
            Binary outcome column (0/1).
        predictors : list[str]
            Feature columns.
        model_type : str
            ``"xgboost"`` or ``"randomforest"``.
        data : DataFrame, optional
            Pre-loaded cohort.
        view : str, optional
            Source view/table.
        n_folds : int
            Cross-validation folds (default 5).

        Returns
        -------
        dict with keys: model, auc_cv, brier_cv, feature_importance,
        shap_values, shap_summary_plot, calibration_plot, n_obs,
        n_events, predictors_used, model_type, warnings
        """
        if model_type == "xgboost" and not HAS_XGB:
            return {"error": "xgboost not installed"}
        if model_type == "randomforest" and not HAS_SKLEARN:
            return {"error": "scikit-learn not installed"}

        if data is None:
            resolved = self._base.resolve_view(view)
            if resolved is None:
                return {"error": "No source view available"}
            data = self._load(resolved)

        if data.empty:
            return {"error": "Empty dataset"}

        cols_needed = [outcome] + predictors
        present = [c for c in cols_needed if c in data.columns]
        if outcome not in present:
            return {"error": f"Outcome '{outcome}' not found in data"}

        sub = self._prepare_numeric(data, present)

        sub[outcome] = sub[outcome].astype(int)
        pred_used = [p for p in predictors if p in sub.columns]
        if not pred_used:
            return {"error": "No valid predictors after cleaning"}

        X = sub[pred_used].values
        y = sub[outcome].values
        feature_names = pred_used

        n_events = int(y.sum())
        result_warnings: list[str] = []

        if n_events < 20:
            result_warnings.append(
                f"Only {n_events} events — ML model likely underpowered. "
                "Interpret feature importances directionally, not quantitatively."
            )
        if n_events < 10:
            return {"error": f"Only {n_events} events — too few for ML training"}

        # Handle imbalance
        event_rate = n_events / len(y) if len(y) > 0 else 0
        if event_rate < 0.05 or event_rate > 0.95:
            result_warnings.append(
                f"Extreme class imbalance (event rate {event_rate:.1%}). "
                "Consider oversampling or adjusting scale_pos_weight."
            )
        scale_pos = (len(y) - n_events) / n_events if n_events > 0 else 1.0

        # Build model
        if model_type == "xgboost":
            model = xgb.XGBClassifier(
                n_estimators=300, max_depth=4, learning_rate=0.05,
                scale_pos_weight=scale_pos, subsample=0.8,
                colsample_bytree=0.8, min_child_weight=5,
                eval_metric="logloss", random_state=42,
                use_label_encoder=False, verbosity=0,
            )
        else:
            model = RandomForestClassifier(
                n_estimators=300, max_depth=6, min_samples_leaf=10,
                class_weight="balanced", random_state=42, n_jobs=-1,
            )

        # Cross-validated performance
        effective_folds = min(n_folds, n_events, len(y) - n_events)
        if effective_folds < 2:
            effective_folds = 2

        try:
            skf = StratifiedKFold(n_splits=effective_folds, shuffle=True, random_state=42)
            cv_probs = cross_val_predict(model, X, y, cv=skf, method="predict_proba")[:, 1]
            auc_cv = round(float(roc_auc_score(y, cv_probs)), 4)
            brier_cv = round(float(brier_score_loss(y, cv_probs)), 4)
        except Exception as exc:
            result_warnings.append(f"CV evaluation failed: {exc}")
            auc_cv = None
            brier_cv = None
            cv_probs = None

        # Final model on full data
        try:
            model.fit(X, y)
        except Exception as exc:
            return {"error": f"Model training failed: {exc}"}

        # Feature importance (native)
        if model_type == "xgboost":
            raw_imp = model.feature_importances_
        else:
            raw_imp = model.feature_importances_

        imp_df = pd.DataFrame({
            "feature": feature_names,
            "importance": np.round(raw_imp, 4),
        }).sort_values("importance", ascending=False)
        imp_df["rank"] = range(1, len(imp_df) + 1)

        # SHAP values
        shap_vals = None
        shap_summary_fig = None
        shap_feature_fig = None
        if HAS_SHAP:
            try:
                if model_type == "xgboost":
                    explainer = shap.TreeExplainer(model)
                else:
                    explainer = shap.TreeExplainer(model)
                sv = explainer.shap_values(X)

                if isinstance(sv, list):
                    sv = sv[1]

                shap_df = pd.DataFrame(sv, columns=feature_names)
                shap_vals = shap_df

                # Build SHAP summary as Plotly (not matplotlib)
                mean_abs = shap_df.abs().mean().sort_values(ascending=True)
                if HAS_PLOTLY:
                    shap_summary_fig = go.Figure(go.Bar(
                        x=mean_abs.values,
                        y=mean_abs.index,
                        orientation="h",
                        marker_color=[_TEAL if v > mean_abs.median() else _GRAY
                                      for v in mean_abs.values],
                    ))
                    _fig_layout(shap_summary_fig, "SHAP Feature Importance",
                                xaxis_title="Mean |SHAP value|",
                                height=max(300, 30 * len(mean_abs) + 80),
                                margin=dict(l=150, r=16, t=50, b=40))

                    # Per-feature beeswarm approximation (top 8 features)
                    top_feats = mean_abs.tail(8).index.tolist()
                    beeswarm_fig = go.Figure()
                    for i, feat in enumerate(top_feats):
                        feat_idx = feature_names.index(feat)
                        sv_feat = shap_df[feat].values
                        x_feat = X[:, feat_idx]

                        x_norm = (x_feat - x_feat.min()) / (x_feat.max() - x_feat.min() + 1e-10)

                        sample_size = min(500, len(sv_feat))
                        rng = np.random.RandomState(42)
                        idx = rng.choice(len(sv_feat), sample_size, replace=False)

                        colors = [f"rgb({int(255*v)},{int(80*(1-v))},{int(200*(1-v))})"
                                  for v in x_norm[idx]]

                        beeswarm_fig.add_trace(go.Scatter(
                            x=sv_feat[idx],
                            y=[i + rng.uniform(-0.3, 0.3) for _ in range(sample_size)],
                            mode="markers",
                            marker=dict(size=3, color=colors, opacity=0.6),
                            name=feat,
                            showlegend=True,
                            hovertemplate=(
                                f"<b>{feat}</b><br>"
                                "SHAP=%{x:.3f}<br>"
                                f"Feature value=%{{text}}<extra></extra>"
                            ),
                            text=[f"{x_feat[j]:.2f}" for j in idx],
                        ))

                    beeswarm_fig.update_layout(
                        yaxis=dict(
                            tickvals=list(range(len(top_feats))),
                            ticktext=top_feats,
                        ),
                    )
                    _fig_layout(beeswarm_fig, "SHAP Beeswarm (top 8 features)",
                                xaxis_title="SHAP value (impact on prediction)",
                                height=max(350, 50 * len(top_feats) + 80))
                    shap_feature_fig = beeswarm_fig

            except Exception as exc:
                result_warnings.append(f"SHAP computation failed: {exc}")

        # Calibration plot
        cal_fig = None
        if cv_probs is not None and HAS_PLOTLY:
            try:
                n_bins = min(10, max(3, n_events // 10))
                frac_pos, mean_pred = calibration_curve(y, cv_probs, n_bins=n_bins, strategy="quantile")
                cal_fig = go.Figure()
                cal_fig.add_trace(go.Scatter(
                    x=mean_pred, y=frac_pos,
                    mode="lines+markers", name="Model",
                    line=dict(color=_TEAL, width=2),
                    marker=dict(size=8),
                ))
                cal_fig.add_trace(go.Scatter(
                    x=[0, 1], y=[0, 1],
                    mode="lines", name="Perfect",
                    line=dict(color=_GRAY, dash="dash"),
                ))
                _fig_layout(cal_fig, "Calibration Plot",
                            xaxis_title="Mean Predicted Probability",
                            yaxis_title="Observed Fraction",
                            height=400)
            except Exception:
                pass

        return {
            "model": model,
            "model_type": model_type,
            "auc_cv": auc_cv,
            "brier_cv": brier_cv,
            "feature_importance": imp_df,
            "shap_values": shap_vals,
            "shap_summary_plot": shap_summary_fig,
            "shap_beeswarm_plot": shap_feature_fig,
            "calibration_plot": cal_fig,
            "n_obs": len(y),
            "n_events": n_events,
            "event_rate": round(event_rate, 4),
            "predictors_used": pred_used,
            "feature_names": feature_names,
            "X": X,
            "warnings": result_warnings,
        }

    # ── 4. Interactive Risk Calculator ────────────────────────────────────

    @staticmethod
    def predict_individual_risk(
        model: Any,
        feature_values: dict[str, float],
        feature_names: list[str],
    ) -> dict[str, Any]:
        """Predict individual risk from a trained ML model.

        Parameters
        ----------
        model : fitted sklearn/xgboost classifier
            Trained model from ``train_ml_nomogram``.
        feature_values : dict
            Feature name → numeric value mapping.
        feature_names : list[str]
            Ordered feature names matching model training.

        Returns
        -------
        dict with risk_pct, risk_class, feature_contributions (if SHAP available)
        """
        x = np.array([[feature_values.get(f, 0.0) for f in feature_names]])

        try:
            prob = float(model.predict_proba(x)[0, 1])
        except Exception as exc:
            return {"error": f"Prediction failed: {exc}"}

        risk_class = (
            "Low" if prob < 0.10 else
            "Intermediate" if prob < 0.30 else
            "High"
        )

        result: dict[str, Any] = {
            "risk_pct": round(prob * 100, 1),
            "risk_probability": round(prob, 4),
            "risk_class": risk_class,
        }

        if HAS_SHAP:
            try:
                explainer = shap.TreeExplainer(model)
                sv = explainer.shap_values(x)
                if isinstance(sv, list):
                    sv = sv[1]
                contributions = dict(zip(feature_names, np.round(sv[0], 4)))
                result["feature_contributions"] = contributions
                result["base_value"] = round(float(explainer.expected_value
                                                    if not isinstance(explainer.expected_value, list)
                                                    else explainer.expected_value[1]), 4)
            except Exception:
                pass

        return result

    # ── 5. Feature Ranges for Risk Calculator ─────────────────────────────

    @staticmethod
    def compute_feature_ranges(
        X: np.ndarray,
        feature_names: list[str],
    ) -> dict[str, dict[str, float]]:
        """Compute min/max/median/mean for each feature for slider defaults."""
        ranges: dict[str, dict[str, float]] = {}
        for i, name in enumerate(feature_names):
            col = X[:, i]
            valid = col[~np.isnan(col)]
            if len(valid) == 0:
                continue
            ranges[name] = {
                "min": float(np.min(valid)),
                "max": float(np.max(valid)),
                "median": float(np.median(valid)),
                "mean": float(np.mean(valid)),
                "p25": float(np.percentile(valid, 25)),
                "p75": float(np.percentile(valid, 75)),
            }
        return ranges

    # ── 6. Manuscript Report Generation ───────────────────────────────────

    def generate_manuscript_report(
        self,
        sections: list[str] | None = None,
        data: pd.DataFrame | None = None,
        view: str | None = None,
        title: str = "Thyroid Cancer Analytics Report",
        author: str = "THYROID_2026 Research Team",
    ) -> dict[str, Any]:
        """Generate a Word document with publication-ready tables and text.

        Assembles selected analysis sections into a formatted .docx using
        python-docx. Each section runs the relevant analysis, formats
        results as Word tables, and appends clinical interpretation text.

        Parameters
        ----------
        sections : list[str], optional
            Sections to include. Defaults to all available.
            Options: ``"Table1"``, ``"Cox"``, ``"CompetingRisks"``,
            ``"ML_Nomogram"``, ``"Longitudinal"``.
        data : DataFrame, optional
            Pre-loaded cohort.
        view : str, optional
            Source view/table.
        title, author : str
            Report metadata.

        Returns
        -------
        dict with keys: docx_bytes, filename, sections_included, warnings
        """
        if not HAS_DOCX:
            return {"error": "python-docx required — pip install python-docx"}

        if sections is None:
            sections = ["Table1", "Cox", "Longitudinal"]

        if data is None:
            resolved = self._base.resolve_view(view)
            if resolved is None:
                return {"error": "No source view available"}
            data = self._load(resolved)

        if data.empty:
            return {"error": "Empty dataset"}

        doc = Document()
        result_warnings: list[str] = []
        sections_included: list[str] = []

        # Title
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Author: {author}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"N = {len(data):,} patients")
        doc.add_paragraph("")

        if "Table1" in sections:
            try:
                t1_df, meta = self._base.generate_table_one(data=data)
                if "error" not in meta and not t1_df.empty:
                    doc.add_heading("Table 1 — Cohort Characteristics", level=1)
                    self._df_to_docx_table(doc, t1_df.reset_index()
                                           if hasattr(t1_df.index, "names") else t1_df)
                    doc.add_paragraph(
                        f"Values are median [IQR] for non-normal continuous variables "
                        f"and n (%) for categorical. N={meta.get('n_total', len(data)):,}."
                    )
                    sections_included.append("Table1")
            except Exception as exc:
                result_warnings.append(f"Table 1 generation failed: {exc}")

        if "Cox" in sections:
            try:
                preds = [p for p in THYROID_PREDICTORS if p in data.columns]
                time_col = THYROID_SURVIVAL["time_col"]
                event_col = THYROID_SURVIVAL["event_col"]
                if time_col in data.columns and event_col in data.columns and preds:
                    cox_result = self._base.fit_cox_ph(
                        time_col=time_col, event_col=event_col,
                        predictors=preds[:8], data=data,
                    )
                    if "error" not in cox_result:
                        doc.add_heading("Cox Proportional Hazards Analysis", level=1)
                        hr = cox_result["hr_table"]
                        self._df_to_docx_table(doc, hr)
                        doc.add_paragraph(
                            f"Concordance index: {cox_result['concordance']:.3f}. "
                            f"N={cox_result['n_obs']:,}, events={cox_result['n_events']:,}."
                        )
                        snippet = self._base.format_clinical_snippet(
                            cox_result, model_type="HR",
                            outcome_label="disease recurrence",
                        )
                        doc.add_paragraph(snippet.replace("**", "").replace("*", ""))
                        sections_included.append("Cox")
            except Exception as exc:
                result_warnings.append(f"Cox section failed: {exc}")

        if "Longitudinal" in sections:
            try:
                long_result = self._base.longitudinal_summary(marker="tg")
                if "error" not in long_result:
                    doc.add_heading("Longitudinal Thyroglobulin Trajectory", level=1)
                    doc.add_paragraph(long_result.get("model_summary", ""))
                    doc.add_paragraph(long_result.get("clinical_note", ""))

                    pp = long_result.get("per_patient_summary", pd.DataFrame())
                    if not pp.empty:
                        summary_stats = pd.DataFrame([{
                            "Metric": "Patients (≥2 obs)",
                            "Value": str(long_result.get("n_patients", "")),
                        }, {
                            "Metric": "Total observations",
                            "Value": str(long_result.get("n_obs", "")),
                        }, {
                            "Metric": "Rising trajectory (%)",
                            "Value": f"{long_result.get('rising_pct', 0):.1f}%",
                        }, {
                            "Metric": "Slope (β/year)",
                            "Value": f"{long_result.get('slope', 0):.4f}",
                        }])
                        self._df_to_docx_table(doc, summary_stats)
                    sections_included.append("Longitudinal")
            except Exception as exc:
                result_warnings.append(f"Longitudinal section failed: {exc}")

        if "CompetingRisks" in sections:
            try:
                preset = COMPETING_RISK_PRESETS.get("recurrence_vs_death", {})
                if all(c in data.columns for c in [preset.get("time_col", ""), preset.get("event_col", "")]):
                    cr = self.fit_competing_risks(
                        time_col=preset["time_col"],
                        event_col=preset["event_col"],
                        competing_event_col=preset.get("competing_event_col"),
                        data=data,
                    )
                    if "error" not in cr:
                        doc.add_heading("Competing Risks Analysis", level=1)
                        self._df_to_docx_table(doc, cr["summary_table"])
                        doc.add_paragraph(cr.get("clinical_note", ""))
                        sections_included.append("CompetingRisks")
            except Exception as exc:
                result_warnings.append(f"Competing risks section failed: {exc}")

        if "ML_Nomogram" in sections:
            try:
                preset = NOMOGRAM_PRESETS.get("recurrence", {})
                out = preset.get("outcome", "event_occurred")
                preds = [p for p in preset.get("predictors", []) if p in data.columns]
                if out in data.columns and len(preds) >= 2:
                    ml_result = self.train_ml_nomogram(
                        outcome=out, predictors=preds,
                        model_type="xgboost", data=data,
                    )
                    if "error" not in ml_result:
                        doc.add_heading("ML Nomogram — Feature Importance", level=1)
                        self._df_to_docx_table(doc, ml_result["feature_importance"])
                        perf_text = f"Cross-validated AUC: {ml_result.get('auc_cv', 'N/A')}"
                        if ml_result.get("brier_cv"):
                            perf_text += f", Brier score: {ml_result['brier_cv']}"
                        doc.add_paragraph(perf_text)
                        sections_included.append("ML_Nomogram")
            except Exception as exc:
                result_warnings.append(f"ML nomogram section failed: {exc}")

        # Serialize
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        ts = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"thyroid_analytics_report_{ts}.docx"

        return {
            "docx_bytes": buf.getvalue(),
            "filename": filename,
            "sections_included": sections_included,
            "warnings": result_warnings,
        }

    @staticmethod
    def _df_to_docx_table(doc: "Document", df: pd.DataFrame, max_rows: int = 100) -> None:
        """Append a DataFrame as a formatted Word table."""
        display = df.head(max_rows).copy()

        for col in display.columns:
            display[col] = display[col].astype(str).str.replace("nan", "", regex=False)

        table = doc.add_table(rows=1, cols=len(display.columns), style="Table Grid")
        for i, col_name in enumerate(display.columns):
            cell = table.rows[0].cells[i]
            cell.text = str(col_name)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for _, row in display.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph("")

    # ── 7. LaTeX Snippet Generation ───────────────────────────────────────

    @staticmethod
    def generate_latex_table(
        df: pd.DataFrame,
        caption: str = "Results",
        label: str = "tab:results",
    ) -> str:
        """Convert a DataFrame to a LaTeX longtable string."""
        cols = list(df.columns)
        n_cols = len(cols)
        col_spec = "l" + "r" * (n_cols - 1)

        lines = [
            r"\begin{longtable}{" + col_spec + "}",
            r"\caption{" + caption + r"} \label{" + label + r"} \\",
            r"\toprule",
            " & ".join(str(c).replace("_", r"\_") for c in cols) + r" \\",
            r"\midrule",
            r"\endfirsthead",
            r"\multicolumn{" + str(n_cols) + r"}{c}{{\tablename\ \thetable{} -- continued}} \\",
            r"\toprule",
            " & ".join(str(c).replace("_", r"\_") for c in cols) + r" \\",
            r"\midrule",
            r"\endhead",
            r"\bottomrule",
            r"\endfoot",
        ]

        for _, row in df.iterrows():
            vals = []
            for v in row:
                s = str(v).replace("_", r"\_").replace("%", r"\%").replace("nan", "")
                vals.append(s)
            lines.append(" & ".join(vals) + r" \\")

        lines.append(r"\end{longtable}")
        return "\n".join(lines)
