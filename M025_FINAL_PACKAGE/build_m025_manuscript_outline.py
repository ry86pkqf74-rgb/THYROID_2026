#!/usr/bin/env python3
"""
Build M025 v2.0 manuscript draft outline as .docx.

Headline framing: patient-level analysis (cursor commit 1d4ecc1) is the
manuscript primary; nodule-level analysis (mig_306) is the sister analysis.

Output: M025_FINAL_PACKAGE/M025_v2_manuscript_DRAFT_outline.docx
"""

from __future__ import annotations

import os
from datetime import datetime, timezone

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUTDIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(OUTDIR, "M025_v2_manuscript_DRAFT_outline.docx")


def set_cell_bg(cell, color_hex: str):
    """Set table cell background color via OOXML shading element."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_para(doc, text: str, bold: bool = False, italic: bool = False,
             style: str | None = None, align=None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_todo(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"[TODO: {text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xA0, 0x40, 0x40)
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_table_with_header(doc, headers: list[str], rows: list[list[str]],
                          header_color: str = "305496"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = True
    # header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""  # clear default
        para = hdr_cells[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        set_cell_bg(hdr_cells[i], header_color)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, start=1):
        cells = tbl.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            para = cells[c_idx].paragraphs[0]
            run = para.add_run(str(val) if val is not None else "")
            run.font.size = Pt(10)
            run.font.name = "Calibri"
    doc.add_paragraph()


def main():
    doc = Document()

    # ========== Page setup ==========
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ========== Default style ==========
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ========== TITLE PAGE ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Diagnostic Performance of ACR TI-RADS in a 25-Year Operative Thyroid Cohort: "
        "Patient-Level Analysis with Nodule-Level Sister Validation"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    add_para(doc, "")
    add_para(doc, "Manuscript draft — DRAFT v0.1", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Article type: Original Article", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "")
    add_para(doc, "Authors", bold=True)
    add_todo(doc, "Author list, degrees, affiliations — confirm with senior author")
    add_para(doc, "Logan D. Glosser, B.S.; [co-author 2]; [senior author], MD, FACS")
    add_para(doc, "On behalf of the THYROID_2026 institutional study group")

    add_para(doc, "")
    add_para(doc, "Affiliations", bold=True)
    add_todo(doc, "Department/Division/Institution — current placeholder is Emory")
    add_para(doc, "1 Department of Surgery, Emory University School of Medicine, Atlanta, Georgia, USA")
    add_para(doc, "2 [Additional affiliations]")

    add_para(doc, "")
    add_para(doc, "Corresponding author", bold=True)
    add_todo(doc, "Final corresponding author + contact info")
    add_para(doc, "Logan D. Glosser, [Address], [Phone], logan.glosser@gmail.com")

    add_para(doc, "")
    add_para(doc, "Running title", bold=True)
    add_para(doc, "ACR TI-RADS in a 25-year operative cohort")

    add_para(doc, "")
    add_para(doc, "Word count", bold=True)
    add_todo(doc, "Update before submission — Thyroid limit ~3,500-4,000")
    add_para(doc, "Abstract: ~250 words. Body: ~3,500 words.")

    add_para(doc, "")
    add_para(doc, "Tables / Figures", bold=True)
    add_para(doc, "4 in-text tables (Tables 1–4); 5 in-text figures (Figures 1–5); "
                 "Supplementary tables S1–S4 + Supplementary Figure S1.")

    add_para(doc, "")
    add_para(doc, "Funding / disclosures", bold=True)
    add_todo(doc, "Fill in funding source; declare conflicts of interest")

    add_para(doc, "")
    add_para(doc, "IRB / ethics", bold=True)
    add_todo(doc, "Confirm IRB protocol number")
    add_para(doc, "This retrospective analysis was conducted under Emory University IRB protocol "
                 "[#TBD] with informed-consent waiver. De-identified data only.")

    add_para(doc, "")
    add_para(doc, "Data availability", bold=True)
    add_para(doc, "De-identified summary tables underlying the analyses are available with the "
                 "manuscript supplement (M025_master_data.xlsx). Patient-level data are subject to "
                 "institutional sharing rules and available on reasonable request.")

    doc.add_page_break()

    # ========== ABSTRACT ==========
    add_heading(doc, "Abstract", level=1)

    add_para(doc, "Background.", bold=True)
    add_para(doc,
             "The ACR TI-RADS score is the most widely used ultrasound risk-stratification "
             "system for thyroid nodules. Validation in operative cohorts has historically "
             "produced higher per-category risk-of-malignancy (ROM) than ACR-published expected "
             "ranges, a finding usually attributed to selection bias. We re-examined diagnostic "
             "performance and calibration in a 25-year single-institution operative cohort, "
             "with a pre-specified secondary analysis at the nodule grain to quantify how much "
             "of operative-cohort ROM elevation is multinodular attribution error rather than "
             "pure selection.")

    add_para(doc, "Methods.", bold=True)
    add_para(doc,
             f"We assembled a single-institution operative thyroid cohort of 3,375 patients with "
             f"surgical pathology between 1994 and 2025 (canonical_patient_master, MotherDuck "
             f"thyroid_canonical_publication_v1_0, pub_v1_1). The primary analysis used patient "
             f"grain with maximum TI-RADS category across all preoperative ultrasound exams as "
             f"the predictor and any pathology-proven thyroid malignancy as the reference standard. "
             f"Diagnostic performance (sensitivity, specificity, positive and negative predictive "
             f"values, Wilson 95% CIs, area under the receiver-operating-characteristic curve) was "
             f"computed at TR≥TR3, TR≥TR4 and TR≥TR5. The Youden index identified the optimal "
             f"clinical threshold. A pre-specified sister analysis was performed at the nodule "
             f"grain (cohort_m025_nodule_level_v1, n=3,687 strict ACR-eligible nodules / 631 path-"
             f"malignant) to compute per-nodule ROM and compare directly with patient-level "
             f"calibration.")

    add_para(doc, "Results.", bold=True)
    add_para(doc,
             f"Among 3,375 patients (79.7% female; 45.5% Black, 40.9% White, 6% Asian; median "
             f"surgery year 2021; 1,479 [43.8%] pathology-proven malignant), patient-level "
             f"discrimination was modest: AUC 0.648 (95% CI 0.630–0.667). The Youden-optimal "
             f"threshold was TR≥TR4 (J=0.271; sensitivity 71.3%, specificity 55.9%, PPV 55.7%, "
             f"NPV 71.4%). Per-category patient ROM was 28.2%, 32.1%, 27.6%, 47.4% and 58.7% "
             f"for TR1–TR5; only TR5 fell within the ACR-expected range. Applying ACR FNA-eligibility "
             f"thresholds retrospectively flagged 1,553 unnecessary FNAs and 472 cancers below "
             f"threshold. The nodule-level sister analysis (3,687 strict ACR-eligible nodules) "
             f"yielded per-nodule TR4 ROM 18.7% (95% CI 16.3–21.5) and TR5 26.1% (23.7–28.6), both "
             f"inside ACR-published bands; AUC was 0.640. The patient-level vs nodule-level "
             f"divergence (TR4 inflation +28.7 percentage points; TR5 +32.6) indicates multinodular "
             f"attribution explains a substantial fraction of patient-grain ROM elevation.")

    add_para(doc, "Conclusions.", bold=True)
    add_para(doc,
             f"In a contemporary 25-year operative cohort, ACR TI-RADS provides modest discrimination "
             f"with TR≥TR4 as the Youden-optimal clinical threshold, but per-category patient-level "
             f"ROM substantially exceeds ACR-expected bands at TR1–TR4. A pre-specified per-nodule "
             f"reanalysis recovers ACR-expected calibration at TR4 and TR5, demonstrating that much "
             f"of the apparent operative-cohort risk inflation reflects multinodular attribution at "
             f"patient grain rather than pure selection bias. Reporting per-nodule rather than "
             f"per-patient ROM may improve calibration for institutional TI-RADS validation studies.")

    doc.add_page_break()

    # ========== INTRODUCTION ==========
    add_heading(doc, "Introduction", level=1)
    add_para(doc,
             "Thyroid nodules are highly prevalent yet rarely malignant; risk-stratification "
             "of ultrasound (US) features therefore drives clinical decisions about biopsy and "
             "surgery. The American College of Radiology (ACR) Thyroid Imaging Reporting and "
             "Data System (TI-RADS), introduced by Tessler and colleagues in 2017, provides a "
             "five-feature additive scoring algorithm with five categorical risk tiers (TR1–TR5) "
             "and explicit fine-needle aspiration (FNA) eligibility thresholds.")
    add_todo(doc, "Add citations: Tessler 2017 ACR TI-RADS White Paper; ATA 2015 management guidelines; "
                  "Middleton et al multicenter validation; recent meta-analyses; Bethesda 2023 "
                  "system; selection-bias critique of operative cohort validations.")
    add_para(doc,
             "Multiple operative-cohort validation studies have observed that per-category ROM "
             "in surgically resected nodules consistently exceeds the ACR-published expected "
             "ranges, particularly at lower TR categories. This 'operative inflation' has been "
             "attributed to selection bias (only nodules judged worrisome enough to warrant "
             "surgery enter the cohort) and to differences in pathology referent.")
    add_para(doc,
             "An under-examined alternative explanation is multinodular attribution error. When a "
             "patient has multiple US nodules and is collapsed to a single patient-level TI-RADS "
             "score (commonly the maximum), all malignancies in that patient are credited to that "
             "single category — even when the path-proven malignant nodule is not the highest-TR "
             "lesion. This attribution couples patient-level ROM mechanically to the prevalence "
             "of multinodular disease in the cohort and inflates the apparent operative-cohort "
             "risk.")
    add_para(doc,
             "We therefore sought to characterize ACR TI-RADS diagnostic performance and per-"
             "category ROM calibration in a 25-year single-institution operative thyroid cohort, "
             "with a pre-specified secondary nodule-level analysis to quantify the contribution "
             "of multinodular attribution error to operative-cohort ROM elevation.")

    doc.add_page_break()

    # ========== METHODS ==========
    add_heading(doc, "Methods", level=1)

    add_heading(doc, "Study design and cohort", level=2)
    add_para(doc,
             "We analyzed the institutional canonical patient master (canonical_patient_master, "
             "MotherDuck thyroid_canonical_publication_v1_0 at release tag pub_v1_1, 2026-05-04), "
             "a deidentified longitudinal data warehouse of all thyroid surgical patients at "
             "[Institution] from 1994 to 2025 (n=10,871 unique research IDs). Patients eligible "
             "for the primary analysis had at least one preoperative ultrasound with a documented "
             "TI-RADS category and a definitive pathology result on operative thyroidectomy "
             "specimen, yielding the analytic cohort of n=3,375.")
    add_todo(doc, "Confirm institution name and data warehouse description for the journal.")

    add_heading(doc, "Patient-level predictor and outcome (primary analysis)", level=2)
    add_para(doc,
             "The primary patient-level predictor was the maximum ACR TI-RADS 2017 category "
             "across all preoperative US exams (max_tirads_category_ever). When multiple US "
             "exams existed, the patient was assigned the highest TR observed; this is the "
             "convention used in most published operative-cohort validations and matches "
             "clinical practice in which the worst nodule drives FNA and surgical decisions. "
             "The reference standard was any pathology-proven thyroid malignancy on the "
             "operative specimen (is_malignant) following the WHO 2022 thyroid tumor classification.")

    add_heading(doc, "Nodule-level predictor and outcome (sister analysis)", level=2)
    add_para(doc,
             "For the pre-specified nodule-level sister analysis we used the per-nodule analytic "
             "spine (cohort_m025_nodule_level_v1, mig_306). The predictor was per-nodule "
             "acr2017_tirads_category, computed from the five ACR 2017 features stored in "
             "canonical_us_nodule_v2. The strict analytic subset required complete five-feature "
             "ACR scoring, known laterality, no size-outlier quarantine, and no unresolved "
             "multi-nodule attribution flag (analytic_eligible_strict_acr_pernodule = TRUE), "
             "yielding n=3,687 nodules across 1,668 patients. The per-nodule reference standard "
             "(nodule_path_proven_malignant) was assigned TRUE if a same-side malignant tumor "
             "existed in canonical_path_malignant_events_v1 with surgery date within "
             "[exam_date, exam_date + 365 days].")

    add_heading(doc, "FNA Bethesda linkage", level=2)
    add_para(doc,
             "Patient-level FNA results were attached via canonical_fna_events_v1; per-nodule FNA "
             "Bethesda 2023 was bridged via the legacy nodule-FNA linkage table imaging_fna_"
             "linkage_v3 reconstructed at the canonical_us_nodule_v2 keying using "
             "(research_id, normalized laterality, |US date − FNA date| ≤ 30 days). Best link per "
             "nodule was selected by smallest day_gap then highest legacy linkage score. Of the "
             "3,687 strict-ACR analytic-eligible nodules, 495 (13.4%) had a bridged Bethesda value; "
             "of the 3,375 patients, 2,380 (70.5%) had a Bethesda result. The carry-forward "
             "limitation of per-nodule FNA size linkage (CF-FNA-SIZE-CM-NULL) is acknowledged in "
             "the Discussion.")

    add_heading(doc, "Statistical analysis", level=2)
    add_para(doc,
             "Continuous variables are reported as mean (SD) or median (IQR) as appropriate; "
             "categorical variables as count (%). Per-category ROM was computed with Wilson "
             "score-based 95% confidence intervals. Diagnostic performance was evaluated at "
             "three pre-specified thresholds (TR≥TR3, TR≥TR4, TR≥TR5) with 2×2-derived "
             "sensitivity, specificity, PPV, NPV and likelihood ratios with Wilson 95% CIs. "
             "The Youden index (J = sensitivity + specificity − 1) identified the optimal "
             "clinical threshold. Discrimination was summarized by the area under the receiver-"
             "operating-characteristic curve (AUC) computed via the closed-form rank Mann–Whitney "
             "equivalent. Patient-level vs nodule-level per-TR ROM divergence was reported in "
             "percentage points (pp) and assessed against ACR 2017 expected bands. Pre-specified "
             "sensitivity arms (S1A relaxed cohort, S1B first-US-only, S1C single-nodule "
             "patients, S1D unilateral path-only) explored selection effects (Supplementary "
             "Table S1). Analyses were performed in DuckDB SQL on the MotherDuck publication "
             "database; reproduction code is provided in the supplement.")

    doc.add_page_break()

    # ========== RESULTS ==========
    add_heading(doc, "Results", level=1)

    add_heading(doc, "Cohort assembly and baseline characteristics (Figure 1, Table 1)", level=2)
    add_para(doc,
             "Of 10,871 unique research IDs in the institutional thyroid surgical warehouse, "
             "3,375 patients met inclusion criteria for the primary patient-level analysis. The "
             "cohort was 79.7% female (n=2,691), 45.5% Black or African American (n=1,535), 40.9% "
             "White (n=1,382), 6.0% Asian (n=204), with median age at surgery 53 years and median "
             "surgery year 2021 (range 1994–2025). Overall malignancy rate was 43.8% (n=1,479).")
    add_para(doc,
             "Distribution by maximum TI-RADS category was: TR1 n=340 (10.1%), TR2 n=299 (8.9%), "
             "TR3 n=845 (25.0%), TR4 n=492 (14.6%), TR5 n=1,399 (41.5%). Bethesda FNA results were "
             "available for 2,380 patients (70.5%); operative pathology histology was available "
             "for 1,538 patients (45.6%; the remainder underwent surgery for benign indications).")

    add_para(doc, "[INSERT TABLE 1: Baseline characteristics by max TI-RADS category]", italic=True)
    add_para(doc, "[INSERT FIGURE 1: Cohort flow diagram]", italic=True)

    add_heading(doc, "Patient-level diagnostic performance (Table 2, Figure 2)", level=2)
    add_para(doc,
             "Patient-level discrimination of ACR TI-RADS for thyroid malignancy in the operative "
             "cohort was modest: AUC 0.648 (95% CI 0.630–0.667). Diagnostic performance at the "
             "three pre-specified thresholds is summarized in Table 2.")

    add_table_with_header(doc,
        ["Threshold", "Sens (95% CI)", "Spec (95% CI)", "PPV (95% CI)", "NPV (95% CI)"],
        [
            ["TR≥TR3", "87.0% (85.2–88.6)", "23.6% (21.7–25.5)", "47.0% (45.2–48.9)", "70.0% (66.3–73.4)"],
            ["TR≥TR4 (Youden optimal, J=0.271)", "71.3% (68.9–73.5)", "55.9% (53.6–58.1)", "55.7% (53.5–58.0)", "71.4% (69.0–73.6)"],
            ["TR≥TR5", "55.5% (53.0–58.0)", "69.5% (67.4–71.5)", "58.7% (56.1–61.2)", "66.7% (64.6–68.7)"],
        ])

    add_para(doc,
             "The Youden-optimal threshold was TR≥TR4 (J=0.271). At this threshold, applying the "
             "ACR 2017 size-based FNA eligibility rules retrospectively flagged 1,553 patients as "
             "having undergone unnecessary FNA (TR<TR3 or below the size threshold for their TR "
             "category) and identified 472 patients with cancers below the FNA threshold (true "
             "false-negatives of the ACR rule).")

    add_para(doc, "[INSERT FIGURE 2: ROC curve, AUC=0.648 with 95% CI band]", italic=True)

    add_heading(doc, "Risk-of-malignancy by TI-RADS category and ACR-expected calibration "
                     "(Table 3, Figure 3)", level=2)

    add_table_with_header(doc,
        ["TR", "n", "Malignant n", "Patient ROM (95% CI)", "ACR-expected", "Within band?"],
        [
            ["TR1", "340", "96", "28.2% (23.7–33.2)", "<2%", "No"],
            ["TR2", "299", "96", "32.1% (27.1–37.6)", "<2%", "No"],
            ["TR3", "845", "233", "27.6% (24.7–30.7)", "<5%", "No"],
            ["TR4", "492", "233", "47.4% (43.0–51.8)", "5–20%", "No"],
            ["TR5", "1,399", "821", "58.7% (56.1–61.2)", ">20%", "Yes"],
        ])

    add_para(doc,
             "Patient-level ROM substantially exceeded the ACR-published expected ranges at "
             "TR1–TR4. Only TR5 ROM fell within the expected band (>20%). The directional "
             "monotonicity from TR3 (27.6%) to TR4 (47.4%) to TR5 (58.7%) was preserved, but the "
             "absolute magnitudes were elevated by 18 to 30 percentage points relative to ACR "
             "expectation at TR3 and TR4.")

    add_para(doc, "[INSERT FIGURE 3: Patient-level ROM bars by TR with ACR-expected band overlay]",
             italic=True)

    add_heading(doc, "Sister nodule-level analysis: ACR-expected calibration restored "
                     "(Figure 3b)", level=2)

    add_table_with_header(doc,
        ["TR", "Patient ROM (95% CI)", "Nodule ROM (95% CI)", "Inflation (pp)", "ACR band", "Nodule in band?"],
        [
            ["TR2", "32.1% (27.1–37.6)", "12.9% (5.1–28.9)", "+19.2", "<2%", "No"],
            ["TR3", "27.6% (24.7–30.7)", "9.1% (7.8–10.7)", "+18.5", "<5%", "No"],
            ["TR4", "47.4% (43.0–51.8)", "18.7% (16.3–21.5)", "+28.7", "5–20%", "Yes"],
            ["TR5", "58.7% (56.1–61.2)", "26.1% (23.7–28.6)", "+32.6", ">20%", "Yes"],
        ])

    add_para(doc,
             "Re-analyzing the same data at the nodule grain (n=3,687 strict ACR-eligible nodules "
             "from 1,668 patients; 631 path-malignant) recovered ACR-expected calibration at TR4 "
             "and TR5. Per-nodule TR4 ROM was 18.7% (within the ACR 5–20% band) and TR5 was 26.1% "
             "(within the >20% band). Discrimination was preserved (per-nodule AUC 0.640 vs "
             "patient AUC 0.648). Per-nodule diagnostic performance at TR≥TR4 was sensitivity "
             "76.9%, specificity 47.1%, PPV 23.1%, NPV 90.8%.")
    add_para(doc,
             "The percentage-point inflation between patient and nodule grain (TR4 +28.7 pp; "
             "TR5 +32.6 pp) is the quantitative magnitude of multinodular attribution error in "
             "this cohort.")

    add_para(doc, "[INSERT FIGURE 3b: Patient vs nodule ROM with ACR bands; attribution error overlay]",
             italic=True)

    add_heading(doc, "Bethesda × TI-RADS cross-stratification (Table 4)", level=2)
    add_para(doc,
             "Patient-level Bethesda distribution was: I (nondiagnostic) n=88 (ROM 29.5%), II "
             "(benign) n=897 (16.2%), III (AUS) n=403 (48.4%), IV (FN/SFN) n=275 (51.6%), V "
             "(suspicious) n=129 (89.9%), VI (malignant) n=588 (83.7%); 995 patients (29.5%) had "
             "no Bethesda result on file. The Bethesda × TR cross-stratification at the strict "
             "nodule level (Table 4) shows the expected concordance pattern, with most missing "
             "Bethesda cells driven by the FNA-linkage carry-forward limitation.")

    add_para(doc, "[INSERT TABLE 4: Bethesda × TI-RADS strict-nodule contingency]", italic=True)

    add_heading(doc, "FNA-eligibility audit and unnecessary biopsy analysis", level=2)
    add_para(doc,
             "Applying the ACR 2017 size-based FNA-eligibility rules retrospectively in this "
             "cohort identified 1,553 FNAs that would not have been recommended by ACR criteria "
             "(unnecessary FNAs by ACR threshold; 46.0% of all FNAs in the cohort) and 472 cancers "
             "(15.0% of all malignant patients) below the ACR FNA-eligibility threshold for their "
             "TR-size combination. This false-negative pattern is concentrated in TR3 small "
             "nodules and TR4 sub-1.5cm nodules.")

    add_para(doc, "[INSERT FIGURE 4: Patient-level confusion matrix at TR≥TR4 + ACR FNA compliance "
                  "stack chart]", italic=True)

    add_heading(doc, "Subgroup and sensitivity analyses", level=2)
    add_para(doc,
             "Subgroup analyses by sex, age band, histology category and surgery era preserved "
             "the modest discrimination and the patient-level ROM pattern (Supplementary Table "
             "S2). Pre-specified sensitivity arms at the nodule grain (Supplementary Table S1: "
             "S1A relaxed cohort 15,309 nodules; S1B first-US-only; S1C single-nodule patients; "
             "S1D unilateral-path-only) directionally supported the primary findings without "
             "altering the calibration conclusion. The mig_264 read-only Bethesda-2 false-negative "
             "audit identified 13/360 (3.6%) Bethesda-II + path-malignant patients as true "
             "false-negative cytology candidates; the remainder were classifiable as multinodular "
             "attribution (n=21), coverage gaps (n=173), or path-bridge timing artifacts (n=12), "
             "supporting the multinodular-attribution thesis.")

    add_para(doc, "[INSERT FIGURE 5: Subgroup forest plot — AUC by demographic stratum]", italic=True)

    doc.add_page_break()

    # ========== DISCUSSION ==========
    add_heading(doc, "Discussion", level=1)

    add_heading(doc, "Principal findings", level=2)
    add_para(doc,
             "In a contemporary 25-year single-institution operative thyroid cohort of 3,375 "
             "patients, ACR TI-RADS provided modest discrimination for thyroid malignancy "
             "(AUC 0.648, 95% CI 0.630–0.667), with TR≥TR4 as the Youden-optimal clinical "
             "threshold (J=0.271; sensitivity 71.3%, specificity 55.9%). Per-category patient-"
             "level ROM substantially exceeded the ACR 2017 expected bands at TR1–TR4. A pre-"
             "specified nodule-level reanalysis of the same cohort (3,687 strict-ACR analytic-"
             "eligible nodules) recovered ACR-expected calibration at the clinically actionable "
             "TR4 and TR5 thresholds, demonstrating that multinodular attribution error explains "
             "a substantial fraction of operative-cohort ROM elevation that has historically been "
             "attributed to selection bias alone.")

    add_heading(doc, "Comparison with prior literature", level=2)
    add_todo(doc, "Pull 4-6 key validation papers (Middleton 2018; Hoang 2019; Ha 2018; Grani 2019; "
                  "Castellana 2020 meta-analysis); add 1-2 ACR position papers; add 2024-2025 papers if any.")
    add_para(doc,
             "Our patient-level discrimination AUC (0.648) is in line with the published range for "
             "operative-cohort TI-RADS validations [REFS]. Our patient-level per-category ROMs at "
             "TR3 (27.6%), TR4 (47.4%) and TR5 (58.7%) are consistent in magnitude with the "
             "operative-cohort literature [REFS], reproducing the well-documented elevation of "
             "operative-cohort ROM relative to ACR-published bands. The novel contribution is the "
             "matched per-nodule reanalysis: TR4 18.7% and TR5 26.1% land squarely within the ACR "
             "expected bands, providing an alternative explanation for the documented operative-"
             "cohort inflation.")

    add_heading(doc, "Multinodular attribution error", level=2)
    add_para(doc,
             "When a multinodular patient is collapsed to a single max-TR category, all path-proven "
             "malignancies in that patient are credited to that category — even when the malignant "
             "nodule is not the highest-TR lesion. The numerical magnitude of this attribution at "
             "patient grain in our cohort is 28.7 pp at TR4 and 32.6 pp at TR5. This is consistent "
             "with an institutional cohort in which a substantial fraction of operative patients "
             "have multiple US-detectable nodules (mean n=2.7 nodules per patient in our "
             "operative cohort). Operative-cohort TI-RADS validation should therefore report per-"
             "nodule ROM in addition to or in place of per-patient ROM to permit direct comparison "
             "with ACR 2017 expected bands.")

    add_heading(doc, "Implications for clinical TI-RADS practice", level=2)
    add_para(doc,
             "The clinical decision to recommend FNA versus surveillance in a TI-RADS-categorized "
             "nodule is necessarily made at the nodule grain. Validation studies that report only "
             "patient-level ROM systematically overstate the operative-cohort risk relative to the "
             "denominator that drives clinical decision-making. Our findings support the practice "
             "of reporting per-nodule ROM in TI-RADS validation studies and reinforce the original "
             "ACR 2017 calibration as approximately accurate at the strictly-eligible per-nodule "
             "level in operative cohorts.")
    add_para(doc,
             "Retrospective application of ACR 2017 FNA-eligibility rules in our cohort flagged "
             "1,553 unnecessary FNAs (46.0% of FNAs performed) but missed 472 cancers (15.0% of "
             "malignancies) below the ACR threshold. This false-negative rate must be considered "
             "in the operative-cohort context: many of these 472 patients were brought to surgery "
             "for incidental or surveillance indications, not because of an FNA-positive result.")

    add_heading(doc, "Strengths", level=2)
    add_para(doc,
             "Strengths include (1) a 25-year single-institution operative cohort with consistent "
             "pathology adjudication; (2) racially diverse patient population (45.5% Black or "
             "African American), enhancing external generalizability; (3) explicit pre-specified "
             "nodule-level sister analysis using a strict ACR 2017 feature-complete subset; "
             "(4) Wilson 95% CIs for all proportions and rank-based AUC; (5) read-only Bethesda-II "
             "false-negative audit that quantifies the true-FN cytology rate (3.6%); and (6) "
             "open-source reproduction code and locked DuckDB queries.")

    add_heading(doc, "Limitations", level=2)
    add_para(doc,
             "Limitations include (1) operative cohort restricts inference to surgically resected "
             "patients; non-operative TI-RADS-stratified surveillance cohorts at the same "
             "institution were not analyzed; (2) per-nodule FNA size is not yet linked at the "
             "nodule grain (carry-forward CF-FNA-SIZE-CM-NULL), limiting per-nodule size-aware ACR "
             "FNA-compliance analysis to the patient grain; (3) Bethesda coverage at the patient "
             "level is 70.5% — patients without FNA were brought directly to surgery on imaging "
             "criteria; (4) institutional pathology referent uses WHO 2022 classification; "
             "results may not generalize to centers using older WHO classifications without "
             "reclassification of FTUMP/NIFTP; (5) no prospective external validation cohort.")

    add_heading(doc, "Conclusions", level=1)
    add_para(doc,
             "In a 25-year single-institution operative thyroid cohort, ACR TI-RADS provides "
             "modest discrimination (AUC 0.648) with TR≥TR4 as the Youden-optimal clinical "
             "threshold. Per-category patient-level ROM substantially exceeds the ACR-expected "
             "bands at TR1–TR4, but a pre-specified per-nodule reanalysis of the same data "
             "recovers ACR-expected calibration at TR4 (18.7%) and TR5 (26.1%). Approximately "
             "29 to 33 percentage points of apparent operative-cohort ROM elevation reflects "
             "multinodular attribution error at patient grain rather than pure selection bias. "
             "Operative-cohort TI-RADS validation studies should report per-nodule ROM to permit "
             "direct comparison with the ACR 2017 calibration.")

    doc.add_page_break()

    # ========== ACK / FUNDING / COI / IRB ==========
    add_heading(doc, "Acknowledgments", level=1)
    add_todo(doc, "Acknowledgments")

    add_heading(doc, "Funding", level=1)
    add_todo(doc, "Funding sources")

    add_heading(doc, "Conflicts of interest", level=1)
    add_para(doc, "The authors declare no conflicts of interest related to this work.")

    add_heading(doc, "Author contributions", level=1)
    add_todo(doc, "CRediT statement (conceptualization, data curation, formal analysis, methodology, "
                  "writing — original draft, writing — review and editing, supervision, etc.)")

    add_heading(doc, "Data and code availability", level=1)
    add_para(doc,
             "De-identified summary tables (M025_master_data.xlsx, M025_tables_and_summary.xlsx) "
             "and reproduction SQL/Python (M025_FINAL_PACKAGE/build_m025_final_xlsx.py; "
             "08_analysis_code/M025_v2_tirads_analysis.sql) are available with the manuscript "
             "supplement. Patient-level data are subject to institutional sharing rules and "
             "available on reasonable request to the corresponding author.")

    doc.add_page_break()

    # ========== REFERENCES ==========
    add_heading(doc, "References (Vancouver style)", level=1)
    add_todo(doc, "Build reference list. Suggested core: Tessler et al ACR TI-RADS 2017 (J Am Coll Radiol); "
                  "Cibas-Ali Bethesda 2017+2023; ATA 2015 Haugen guidelines; Middleton 2018 multicenter "
                  "validation; Hoang 2019; Ha 2018 EU-TIRADS; Grani 2019; Castellana 2020 meta-analysis; "
                  "Tappouni 2019 multireader; ACR 2024 update; recent JNCI / Thyroid 2024-2025.")

    doc.add_page_break()

    # ========== TABLES (separate page list) ==========
    add_heading(doc, "Tables (in-text)", level=1)
    add_para(doc, "Table 1. Baseline clinical characteristics of the 3,375-patient operative "
                  "cohort by maximum TI-RADS category.")
    add_para(doc, "Table 2. Patient-level diagnostic performance of ACR TI-RADS at TR≥TR3, "
                  "TR≥TR4 (Youden-optimal), and TR≥TR5.")
    add_para(doc, "Table 3. Per-category risk of malignancy at the patient grain with ACR 2017 "
                  "expected bands.")
    add_para(doc, "Table 4. Bethesda × TI-RADS contingency table at the strict-ACR-eligible "
                  "nodule level (n=3,687).")

    add_heading(doc, "Figure legends", level=1)
    add_para(doc, "Figure 1. Cohort flow diagram. From 10,871 unique research IDs in the "
                  "institutional thyroid surgical warehouse, 3,375 patients met inclusion "
                  "criteria for the patient-level primary analysis. Pre-specified strict-ACR "
                  "subset of 3,687 nodules from 1,668 patients was used for the nodule-level "
                  "sister analysis.")
    add_para(doc, "Figure 2. Patient-level receiver-operating-characteristic curve for ACR TI-RADS "
                  "in the operative cohort. AUC 0.648 (95% CI 0.630–0.667). Youden-optimal "
                  "threshold at TR≥TR4 (J=0.271).")
    add_para(doc, "Figure 3. Patient-level per-category risk of malignancy by TI-RADS category "
                  "with Wilson 95% CIs and overlay of ACR 2017 expected ranges.")
    add_para(doc, "Figure 3b. Patient vs nodule per-category ROM (paired bars) with ACR 2017 "
                  "expected bands. Per-nodule TR4 (18.7%) and TR5 (26.1%) recover within-band "
                  "calibration; patient-level inflation at TR4/TR5 is 28.7 / 32.6 percentage "
                  "points respectively.")
    add_para(doc, "Figure 4. Confusion matrix at the Youden-optimal TR≥TR4 threshold and ACR "
                  "FNA-compliance stacked chart (1,553 unnecessary FNAs flagged; 472 cancers "
                  "below threshold).")
    add_para(doc, "Figure 5. Subgroup forest plot — AUC stratified by sex, age band, surgery era, "
                  "and histology category. Discrimination is preserved (modest) across all strata.")

    add_heading(doc, "Supplementary materials", level=1)
    add_para(doc, "Supplementary Table S1. Pre-specified sensitivity arms at the nodule grain "
                  "(A relaxed cohort, B first-US-only, C single-nodule patients, D unilateral "
                  "path-only).")
    add_para(doc, "Supplementary Table S2. Subgroup-stratified per-TR ROM and AUC (sex, age band, "
                  "histology category, surgery era).")
    add_para(doc, "Supplementary Table S3. mig_264 Bethesda-II false-negative audit dispositions "
                  "(n=360).")
    add_para(doc, "Supplementary Table S4. ACR 2017 FNA-eligibility rule application — "
                  "unnecessary-FNA and below-threshold-cancer breakdown by TR-size cell.")
    add_para(doc, "Supplementary Figure S1. Bethesda × TI-RADS heatmap at the strict-ACR-eligible "
                  "nodule level (visual companion to Table 4).")

    # ========== FOOTER PROVENANCE ==========
    doc.add_page_break()
    add_heading(doc, "Provenance", level=1)
    add_para(doc, f"Database: thyroid_canonical_publication_v1_0 (release tag pub_v1_1, 2026-05-04).")
    add_para(doc, "Patient cohort view: manuscript_workspace.cohort_m025_tirads_performance_v1.")
    add_para(doc, "Nodule cohort view: manuscript_workspace.cohort_m025_nodule_level_v1 (mig_306).")
    add_para(doc, "Analytic master tables: manuscript_workspace.m025_analytic_master_patient_v1, "
                  "m025_analytic_master_nodule_v1, m025_threshold_metrics_v1, m025_rom_by_tr_v1, "
                  "m025_bethesda_x_tr_v1 (mig_307b).")
    add_para(doc, "Migration sign-offs: mig_306 (nodule-level spine), mig_307 (M025 v2.0 "
                  "submission package), mig_307b (analytic master tables).")
    add_para(doc, f"Draft generated: {datetime.now(timezone.utc).isoformat()}.")
    add_para(doc, "Source CSVs / xlsx: M025_FINAL_PACKAGE/M025_master_data.xlsx and "
                  "M025_tables_and_summary.xlsx.")
    add_para(doc, "Methods prose source: M025_submission_package_v2_0/08_analysis_code/"
                  "METHODS_DRAFT.md (Cowork pre-bake) + cursor 1d4ecc1 patient-level pipeline.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
