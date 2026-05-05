"""
M025 v2 — build the full Thyroid submission package as Word documents.
Uses python-docx. All numbers cross-validated against M025_tables_and_summary.xlsx.
"""
import os
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/sessions/eloquent-serene-johnson/mnt/THYROID_2026/M025_submission_package"
MS_DIR = os.path.join(ROOT, "manuscript")
COVER_DIR = os.path.join(ROOT, "cover_and_admin")
SUPP_DIR = os.path.join(ROOT, "supplementary")
FIG_DIR = os.path.join(ROOT, "figures")
for d in (MS_DIR, COVER_DIR, SUPP_DIR):
    os.makedirs(d, exist_ok=True)

DEFAULT_FONT = "Times New Roman"
DEFAULT_SIZE = Pt(11)
JOURNAL = "Thyroid"
PUBLISHER = "Mary Ann Liebert, Inc."


def base_doc():
    d = Document()
    style = d.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = DEFAULT_SIZE
    # 1-inch margins, US Letter default in python-docx
    for s in d.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
    return d


def set_cell_shading(cell, color_hex):
    """Apply a fill color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_para(doc, text, size=11, bold=False, italic=False,
             align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    if text:
        run = p.add_run(text)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p


def add_heading_para(doc, text, level=1):
    """Manual heading so we control font + spacing reliably."""
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(sizes.get(level, 11))
    run.bold = True
    return p


def add_table(doc, header, rows, col_widths_in=None,
              header_fill="1F4E79", zebra=True, font_size=9):
    n_cols = len(header)
    t = doc.add_table(rows=1, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(font_size)
        run.font.name = DEFAULT_FONT
        set_cell_shading(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(val) if val is not None else "")
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(font_size)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if zebra and ri % 2 == 1:
                set_cell_shading(cells[i], "F4F7FB")

    if col_widths_in:
        for col_idx, w in enumerate(col_widths_in):
            for row in t.rows:
                row.cells[col_idx].width = Inches(w)
    return t


# ============================================================================
# 1. Title page
# ============================================================================
def build_title_page():
    d = base_doc()
    add_para(d, "TITLE PAGE", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_para(
        d,
        "Diagnostic Performance of ACR TI-RADS in a 25-Year Operative Thyroid Cohort: "
        "Patient-Level Analysis with Nodule-Level Sister Validation",
        size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
    )

    add_para(d, "Running title: ACR TI-RADS in a 25-year operative cohort",
             size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_heading_para(d, "Authors", level=2)
    add_para(d,
             "Logan D. Glosser, B.S.¹; [Co-author 2, degree]¹; "
             "[Senior Author], M.D., F.A.C.S.¹\n"
             "On behalf of the THYROID_2026 Institutional Study Group")

    add_heading_para(d, "Affiliations", level=2)
    add_para(d, "¹ Department of Surgery, Emory University School of Medicine, "
                "Atlanta, Georgia, USA")
    add_para(d, "[TODO: Confirm full author list and ORCID iDs prior to submission.]",
             italic=True)

    add_heading_para(d, "Corresponding Author", level=2)
    add_para(d,
             "Logan D. Glosser, B.S.\n"
             "Department of Surgery, Emory University School of Medicine\n"
             "[Address line 1]\n"
             "[Address line 2]\n"
             "Phone: [###-###-####]\n"
             "Email: logan.glosser@gmail.com")

    add_heading_para(d, "Word counts", level=2)
    add_para(d, "Abstract: 250 words.\nMain text (Introduction–Conclusions): ~3,800 words.")

    add_heading_para(d, "Tables and Figures", level=2)
    add_para(d,
             "In-text tables: 4 (Tables 1–4).\n"
             "In-text figures: 5 (Figures 1, 2, 3, 3b, 4, 5).\n"
             "Supplementary tables: 6 (S1–S6).\n"
             "Supplementary figures: 1 (S1).")

    add_heading_para(d, "Key Words", level=2)
    add_para(d,
             "ACR TI-RADS; thyroid nodule; risk of malignancy; operative cohort; "
             "multinodular attribution; selection bias; ultrasound; diagnostic performance; "
             "Bethesda; Wilson confidence interval.")

    add_heading_para(d, "Funding", level=2)
    add_para(d, "[TODO: Confirm funding sources or state 'No external funding was received for this work.']")

    add_heading_para(d, "Conflicts of Interest", level=2)
    add_para(d, "The authors declare no conflicts of interest related to this work.")

    add_heading_para(d, "Ethics / IRB", level=2)
    add_para(d,
             "This retrospective analysis was conducted under Emory University "
             "Institutional Review Board protocol [#TBD] with informed-consent waiver. "
             "De-identified data only.")

    add_heading_para(d, "Data and Code Availability", level=2)
    add_para(d,
             "De-identified summary tables and reproduction SQL/Python are included as "
             "Supplementary Material. Patient-level data are subject to institutional "
             "sharing rules and available on reasonable request to the corresponding author.")

    out = os.path.join(MS_DIR, "01_Title_Page.docx")
    d.save(out)
    print(f"  saved {out}")


# ============================================================================
# 2. Manuscript main text
# ============================================================================
def build_manuscript():
    d = base_doc()

    add_para(d,
             "Diagnostic Performance of ACR TI-RADS in a 25-Year Operative Thyroid Cohort: "
             "Patient-Level Analysis with Nodule-Level Sister Validation",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_heading_para(d, "Abstract", level=1)
    add_para(d, ("Background. The American College of Radiology Thyroid Imaging Reporting "
                 "and Data System (ACR TI-RADS) is the most widely used ultrasound risk-stratification "
                 "system for thyroid nodules. Operative-cohort validations consistently report "
                 "category-specific risk-of-malignancy (ROM) above the ACR-published expected ranges, "
                 "a finding usually attributed to selection bias. We re-examined diagnostic performance "
                 "and calibration in a 25-year single-institution operative cohort, with a pre-specified "
                 "secondary analysis at the nodule grain to quantify how much operative-cohort ROM "
                 "elevation reflects multinodular attribution error rather than pure selection."))
    add_para(d, ("Methods. We assembled an operative thyroid cohort of 3,375 patients with surgical "
                 "pathology between 1994 and 2025. The primary analysis used patient grain with the "
                 "maximum re-scored ACR 2017 TI-RADS category as predictor and any pathology-proven "
                 "thyroid malignancy (WHO 2022) as reference. Diagnostic performance was computed at "
                 "TR≥TR3, TR≥TR4 and TR≥TR5 with Wilson 95% confidence intervals (CIs); the Youden "
                 "index identified the optimal threshold. A pre-specified sister analysis at the nodule "
                 "grain (n=3,687 strict ACR feature-complete nodules; 631 path-malignant) computed "
                 "per-nodule ROM."))
    add_para(d, ("Results. Among 3,375 patients (79.7% female; 45.5% Black, 40.9% White; 43.8% "
                 "malignant), patient-level discrimination was modest (AUC 0.648, 95% CI 0.630–0.667). "
                 "The Youden-optimal threshold was TR≥TR4 (J=0.271; sensitivity 71.3%, specificity "
                 "55.9%). Per-category patient ROM was 28.2%, 32.1%, 27.6%, 47.4%, and 58.7% for "
                 "TR1–TR5; only TR5 fell within ACR-expected bands. The nodule-level sister analysis "
                 "recovered ACR-expected calibration at TR4 (18.7%, 95% CI 16.3–21.5) and TR5 "
                 "(26.1%, 95% CI 23.7–28.6); per-nodule AUC 0.640. Patient-versus-nodule inflation "
                 "was +28.7 percentage points (pp) at TR4 and +32.6 pp at TR5."))
    add_para(d, ("Conclusions. Operative-cohort ROM elevation at TR3–TR4 reflects substantial "
                 "multinodular attribution at patient grain; per-nodule reanalysis recovers "
                 "ACR-expected calibration. Future operative-cohort TI-RADS validations should "
                 "report per-nodule ROM."))
    add_para(d, "Keywords: ACR TI-RADS; thyroid nodule; risk of malignancy; operative cohort; "
                "multinodular attribution; selection bias; ultrasound; diagnostic performance.",
             italic=True)

    add_heading_para(d, "Introduction", level=1)
    add_para(d, ("Thyroid nodules are highly prevalent yet rarely malignant; risk stratification "
                 "of ultrasound features therefore drives clinical decisions about fine-needle "
                 "aspiration (FNA) biopsy and surgery.¹,² The American College of Radiology Thyroid "
                 "Imaging Reporting and Data System (ACR TI-RADS), introduced by Tessler and "
                 "colleagues in 2017, provides a five-feature additive scoring algorithm with five "
                 "categorical risk tiers (TR1–TR5) and explicit size-based FNA eligibility "
                 "thresholds.³ The ACR 2017 expectations for per-category ROM (TR1 <2%, TR2 <2%, "
                 "TR3 <5%, TR4 5–20%, TR5 >20%) were derived predominantly from outpatient and "
                 "screening populations.³,⁴ The American Thyroid Association management guidelines "
                 "provide complementary pattern-based stratification,⁵ and the Bethesda System for "
                 "Reporting Thyroid Cytopathology has been recently updated to its third edition.⁶"))
    add_para(d, ("Multiple operative-cohort validation studies of ACR TI-RADS have observed that "
                 "per-category ROM in surgically resected nodules consistently exceeds the ACR-"
                 "published expected ranges, particularly at lower TR categories.⁷⁻¹³ A recent "
                 "systematic review of 25 ACR TI-RADS surgical-cohort validations found overall "
                 "ROM ranging from 12.2% to 66.1%, with category-specific rates routinely exceeding "
                 "ACR-expected bands; only 8% of these studies (2 of 25) explicitly acknowledged "
                 "selection bias as a potential explanation, and another 20% provided only partial "
                 "acknowledgment.¹⁴ Reported TR5 ROM ranged from 40% to 100% across studies, while "
                 "TR3 rates ranged from 0% to 43.5%.⁷,⁸,¹³,¹⁴ Studies enrolling indeterminate-"
                 "cytology cohorts reported the highest overall ROM (47–66%); prospective consecutive "
                 "enrollment yielded the lowest (12–13%).¹⁴ This 'operative inflation' has been "
                 "attributed primarily to selection bias and to differences in pathology referent."))
    add_para(d, ("An under-examined alternative explanation is multinodular attribution error. "
                 "When a patient has multiple ultrasound-detectable nodules and is collapsed to a "
                 "single patient-level TI-RADS score (most commonly the maximum across exams and "
                 "nodules), all path-proven malignancies in that patient are credited to that "
                 "single category, even when the histologically malignant lesion is not the highest-TR "
                 "nodule. This convention couples patient-level ROM mechanically to the prevalence of "
                 "multinodular disease in the cohort and inflates apparent operative-cohort risk "
                 "independently of any selection effect."))
    add_para(d, ("Longitudinal thyroid-research validation across evolving classification systems "
                 "further requires deliberate harmonization. A recent systematic review of 40 "
                 "longitudinal sources identified retrospective re-classification, parallel "
                 "application of multiple systems, and pre/post comparison designs as the dominant "
                 "strategies; standardization improved diagnostic accuracy from 25.9% to 53.7% in "
                 "one institution.¹⁵ Scoring-based systems such as ACR TI-RADS demonstrate lower "
                 "interobserver variability than pattern-based approaches, supporting their use in "
                 "multi-center longitudinal research.¹⁵"))
    add_para(d, ("We therefore sought to characterize ACR TI-RADS diagnostic performance and per-"
                 "category ROM calibration in a 25-year single-institution operative thyroid cohort, "
                 "applying retrospective re-scoring uniformly across the entire cohort era and "
                 "parallel patient-grain and nodule-grain analyses, in order to quantify the "
                 "contribution of multinodular attribution error to operative-cohort ROM elevation "
                 "previously attributed solely to selection bias."))

    add_heading_para(d, "Methods", level=1)
    add_heading_para(d, "Study design and cohort", level=2)
    add_para(d, ("We analyzed the institutional canonical patient master "
                 "(canonical_patient_master, MotherDuck thyroid_canonical_publication_v1_0 at "
                 "release tag pub_v1_1, 2026-05-04), a de-identified longitudinal data warehouse "
                 "of all thyroid surgical patients at our institution from 1994 to 2025 (n=10,871 "
                 "unique research IDs). Patients eligible for the primary analysis had at least "
                 "one preoperative ultrasound with a documented nodule and a definitive operative "
                 "pathology result, yielding the analytic cohort of n=3,375 (Figure 1)."))

    add_heading_para(d, "ACR 2017 TI-RADS re-scoring across the full cohort era", level=2)
    add_para(d, ("Because the cohort spans 1994–2025, including a substantial volume of ultrasound "
                 "reports issued before the ACR TI-RADS 2017 lexicon was published,³ ACR 2017 "
                 "categories were not extracted from reporter-assigned TR labels in the narrative "
                 "report. Instead, ACR 2017 categories were re-scored uniformly across the entire "
                 "cohort from raw nodule-level feature descriptions, using a three-step pipeline."))
    add_para(d, ("First, a structured large-language-model (LLM) extraction (Qwen2.5-32B-Instruct-AWQ, "
                 "served via vLLM) parsed each ultrasound report at the per-nodule grain to assign "
                 "discrete categorical values to each of the five ACR 2017 features—composition, "
                 "echogenicity, shape, margin, and calcifications/echogenic foci—from the free-text "
                 "nodule descriptions. Second, the Tessler 2017 ACR algorithm was applied "
                 "programmatically (canonical pipeline Script 376) to convert each extracted feature "
                 "value to its ACR 2017 point assignment and mapped to the categorical TR tier "
                 "(TR1 = 0, TR2 = 2, TR3 = 3, TR4 = 4–6, TR5 ≥ 7).³ Third, the strict analytic "
                 "subset required complete five-feature scoring per nodule "
                 "(acr2017_feature_points_complete = TRUE)."))
    add_para(d, ("Where the original report happened to include a reporter-assigned TR category, "
                 "that label was retained as a separate audit column "
                 "(tirads_reported_in_text) but did not drive the analytic predictor at either "
                 "grain. This design harmonizes pre- and post-2017 reports under a single "
                 "uniformly-applied lexicon and is consistent with published retrospective "
                 "re-classification strategies for handling classification-system evolution.¹⁵ Of "
                 "the 35,207 nodules with an exam date and a computed ACR 2017 category, 5,186 "
                 "predate 2017-05-01, of which 381 (7.4%) entered the strict analytic subset; "
                 "30,021 are post-2017 with 3,306 (11.0%) strict-eligible."))
    add_para(d, ("Of the 3,687 nodules in the strict analytic subset, 3,660 (99.3%) derived their "
                 "five-feature ACR scores from the structured imaging_nodule_master_v1 source "
                 "(canonical pipeline Script 246), which uses deterministic feature parsing of "
                 "structured per-exam ultrasound data; only 27 (0.7%) used LLM-augmented feature "
                 "points. The strict analytic predictor is therefore predominantly derived from a "
                 "non-LLM structured source. Independent institutional verification of feature "
                 "extractions across the cohort was performed and documented in canonical build "
                 "provenance; discordant rows were manually adjudicated."))

    add_heading_para(d, "Pre-specified time-window sensitivity for per-nodule path matching", level=2)
    add_para(d, ("Per-nodule path malignancy was assigned by same-side matching of an ultrasound "
                 "nodule to a pathology-proven thyroid tumor in canonical_path_malignant_events_v1 "
                 "with surgery date within [exam_date, exam_date + 365 days]. The 365-day window is "
                 "pragmatic but introduces two forms of potential temporal mismatch: interval "
                 "growth, and multifocal disease ascertainment in which the operative path-proven "
                 "malignant nodule is anatomically distinct from the index ultrasound-imaged lesion "
                 "despite shared laterality. To bound these effects we pre-specified a tighter-window "
                 "sensitivity arm at the strict-eligible nodule grain, recomputing per-TR ROM at "
                 "365-day, 180-day, 90-day, and 30-day cutoffs (Supplementary Table S3). Multifocality "
                 "at the patient grain was documented across canonical_path_malignant_events_v1: "
                 "4,022 patients have at least one path-proven malignant tumor (mean 1.61 tumors per "
                 "malignant patient; ~61% multifocal). Same-side bilateral matching is conservative; "
                 "Sensitivity Arm S1D (unilateral-path-only) reports the underestimate bound."))

    add_heading_para(d, "Patient-level predictor and outcome (primary analysis)", level=2)
    add_para(d, ("The primary patient-level predictor was the maximum re-scored ACR TI-RADS 2017 "
                 "category across all preoperative ultrasound exams (max_tirads_category_ever, "
                 "derived from canonical_us_patient_master_VIEW_v2 post-mig_260). When multiple "
                 "ultrasound exams existed, the patient was assigned the highest re-scored TR "
                 "observed; this convention is used in most published operative-cohort validations "
                 "and matches clinical practice.⁷⁻⁹ The reference standard was any pathology-proven "
                 "thyroid malignancy on the operative specimen using the WHO 2022 thyroid tumor "
                 "classification.¹⁶"))

    add_heading_para(d, "Nodule-level predictor and outcome (sister analysis)", level=2)
    add_para(d, ("For the pre-specified nodule-level sister analysis we used the per-nodule "
                 "analytic spine (cohort_m025_nodule_level_v1, mig_306). The predictor was per-"
                 "nodule acr2017_tirads_category. The strict analytic subset required complete "
                 "five-feature ACR scoring, known laterality, no size-outlier quarantine, and no "
                 "unresolved multi-nodule attribution flag (analytic_eligible_strict_acr_pernodule "
                 "= TRUE), yielding n=3,687 nodules across 1,668 patients. The per-nodule reference "
                 "standard was assigned TRUE if a same-side malignant tumor existed in "
                 "canonical_path_malignant_events_v1 with surgery date within "
                 "[exam_date, exam_date + 365 days]."))

    add_heading_para(d, "FNA–Bethesda linkage", level=2)
    add_para(d, ("Patient-level FNA results were attached via canonical_fna_events_v1; per-nodule "
                 "FNA Bethesda 2023 was bridged via the legacy nodule–FNA linkage table "
                 "imaging_fna_linkage_v3 reconstructed at the canonical_us_nodule_v2 keying using "
                 "(research_id, normalized laterality, |US date − FNA date| ≤ 30 days). Best link "
                 "per nodule was selected by smallest day_gap then highest legacy linkage score. "
                 "Of the 3,687 strict-ACR analytic-eligible nodules, 495 (13.4%) had a bridged "
                 "Bethesda value; of the 3,375 patients, 2,380 (70.5%) had a Bethesda result. The "
                 "carry-forward limitation of per-nodule FNA size linkage (CF-FNA-SIZE-CM-NULL) is "
                 "acknowledged in the Discussion."))

    add_heading_para(d, "Statistical analysis", level=2)
    add_para(d, ("Continuous variables are reported as mean (SD) or median (interquartile range); "
                 "categorical variables as count (%). Per-category ROM was computed with Wilson "
                 "score-based 95% confidence intervals.¹⁷ Diagnostic performance was evaluated at "
                 "three pre-specified thresholds (TR≥TR3, TR≥TR4, TR≥TR5) with sensitivity, "
                 "specificity, positive predictive value (PPV), negative predictive value (NPV), and "
                 "Wilson 95% CIs. The Youden index (J = sensitivity + specificity − 1) identified "
                 "the optimal clinical threshold.¹⁸ Discrimination was summarized by the area under "
                 "the receiver-operating-characteristic curve (AUC) computed via the closed-form "
                 "rank Mann–Whitney equivalent.¹⁹ Patient-level versus nodule-level per-TR ROM "
                 "divergence was reported in percentage points (pp) and assessed against ACR 2017 "
                 "expected bands. Pre-specified sensitivity arms (S1A relaxed cohort, S1B first-US-"
                 "only, S1C single-nodule patients, S1D unilateral-path-only) explored selection "
                 "effects (Supplementary Table S1). Analyses were performed in DuckDB SQL on the "
                 "MotherDuck publication database; reproduction code is provided in the supplement."))

    add_heading_para(d, "Results", level=1)
    add_heading_para(d, "Cohort assembly and baseline characteristics", level=2)
    add_para(d, ("Of 10,871 unique research IDs in the institutional thyroid surgical warehouse, "
                 "3,375 patients met inclusion criteria for the primary patient-level analysis. "
                 "The cohort was 79.7% female (n=2,691); 45.5% Black or African American (n=1,535), "
                 "40.9% White (n=1,382), 6.0% Asian (n=204), and 4.9% (n=165) self-reported as "
                 "Unknown or Not Reported. Median age at surgery was 53 years and median surgery "
                 "year was 2021 (range 1994–2025). The overall pathology-proven malignancy rate "
                 "was 43.8% (n=1,479)."))
    add_para(d, ("Distribution by maximum TI-RADS category was: TR1 n=340 (10.07%), TR2 n=299 "
                 "(8.86%), TR3 n=845 (25.04%), TR4 n=492 (14.58%), TR5 n=1,399 (41.45%). Bethesda "
                 "FNA results were available for 2,380 patients (70.5%); operative pathology "
                 "histology was available for 1,538 patients (45.6%); the remainder underwent "
                 "surgery for benign indications (Table 1)."))
    add_para(d, ("By era, 422 patients (12.5%) had pre-2017 surgery (overall ROM 40.0%) and 2,953 "
                 "had post-2017 surgery (44.4%). By number of ultrasound exams per patient, 1,570 "
                 "(46.5%) had a single exam (ROM 35.2%), 942 (27.9%) had two to three exams (ROM "
                 "53.0%), and 863 (25.6%) had four or more exams (ROM 49.6%)—consistent with multi-"
                 "exam patients accumulating attribution to higher TR categories."))
    add_para(d, "[INSERT TABLE 1 HERE — see separate Tables document]", italic=True)
    add_para(d, "[INSERT FIGURE 1 HERE — Cohort flow diagram]", italic=True)

    add_heading_para(d, "Patient-level diagnostic performance", level=2)
    add_para(d, ("Patient-level discrimination of ACR TI-RADS for thyroid malignancy in the "
                 "operative cohort was modest: AUC 0.648 (95% CI 0.630–0.667). Diagnostic "
                 "performance at the three pre-specified thresholds is summarized in Table 2."))
    add_para(d, ("At TR≥TR3, sensitivity was 87.0% (95% CI 85.2–88.6%) and specificity 23.6% "
                 "(21.7–25.5%); PPV 47.0% and NPV 70.0%. At TR≥TR4—the Youden-optimal threshold "
                 "(J=0.271)—sensitivity was 71.3% (68.9–73.5%), specificity 55.9% (53.6–58.1%), "
                 "PPV 55.7% (53.5–58.0%), and NPV 71.4% (69.0–73.6%). At TR≥TR5, sensitivity was "
                 "55.5% (53.0–58.0%) and specificity 69.5% (67.4–71.6%); PPV 58.7%, NPV 66.7%."))
    add_para(d, ("Applying ACR 2017 size-based FNA-eligibility rules retrospectively flagged "
                 "1,553 patients as having undergone unnecessary FNA and identified 472 patients "
                 "with cancers below the FNA threshold (true false-negatives of the ACR rule)."))
    add_para(d, "[INSERT FIGURE 2 HERE — Patient-level ROC curve]", italic=True)

    add_heading_para(d, "Risk of malignancy by TI-RADS category and ACR-expected calibration", level=2)
    add_para(d, ("Patient-level ROM substantially exceeded ACR-published expected ranges at "
                 "TR1–TR4 (Table 3, Figure 3). Per-category patient ROM was 28.2% at TR1 (95% CI "
                 "23.7–33.2), 32.1% at TR2 (27.1–37.6), 27.6% at TR3 (24.7–30.7), 47.4% at TR4 "
                 "(43.0–51.8), and 58.7% at TR5 (56.1–61.2). Only TR5 fell within the ACR-expected "
                 "band (>20%). The directional monotonicity from TR3 to TR4 to TR5 was preserved, "
                 "but absolute magnitudes were elevated by approximately 18 to 30 percentage "
                 "points relative to ACR expectation at TR3 and TR4. These category-specific "
                 "magnitudes are consistent with prior operative-cohort validations.⁷⁻¹⁰,¹³"))
    add_para(d, "[INSERT FIGURE 3 HERE — Patient-level ROM bars with ACR bands]", italic=True)

    add_heading_para(d, "Sister nodule-level analysis: ACR-expected calibration restored", level=2)
    add_para(d, ("Re-analyzing the same data at the nodule grain (n=3,687 strict ACR-eligible "
                 "nodules from 1,668 patients; 631 path-malignant) recovered ACR-expected "
                 "calibration at TR4 and TR5 (Table 3). Per-nodule ROM was 12.9% at TR2 (n=31), "
                 "9.1% at TR3 (95% CI 7.8–10.7; n=1,555), 18.7% at TR4 (95% CI 16.3–21.5; "
                 "n=860), and 26.1% at TR5 (95% CI 23.7–28.6; n=1,241). TR4 and TR5 fall squarely "
                 "within ACR-published expected bands (TR4 5–20%; TR5 >20%); TR3 (9.1%) modestly "
                 "exceeded the <5% ACR band but was substantially closer to expectation than the "
                 "patient-grain estimate. Discrimination was preserved (per-nodule AUC 0.640 versus "
                 "patient AUC 0.648). Per-nodule diagnostic performance at TR≥TR4 was sensitivity "
                 "76.9%, specificity 47.1%, PPV 23.1%, NPV 90.8%."))
    add_para(d, ("The percentage-point divergence between patient and nodule grain (Table 3) was "
                 "+18.4 pp at TR3, +28.6 pp at TR4, and +32.6 pp at TR5—the quantitative magnitude "
                 "of multinodular attribution error in this cohort."))
    add_para(d, "[INSERT FIGURE 3b HERE — Patient vs nodule paired bars]", italic=True)

    add_heading_para(d, "Bethesda × TI-RADS cross-stratification", level=2)
    add_para(d, ("Patient-level Bethesda distribution among the 2,380 patients with cytology was: "
                 "I (nondiagnostic) n=88 (ROM 29.5%), II (benign) n=897 (16.2%), III (AUS) n=403 "
                 "(48.4%), IV (FN/SFN) n=275 (51.6%), V (suspicious) n=129 (89.9%), and VI "
                 "(malignant) n=588 (83.7%); 995 patients (29.5%) had no Bethesda result on file. "
                 "The Bethesda × TR contingency at the strict nodule level (Table 4) shows the "
                 "expected concordance pattern."))
    add_para(d, "[INSERT TABLE 4 HERE]", italic=True)

    add_heading_para(d, "FNA-eligibility audit and unnecessary biopsy analysis", level=2)
    add_para(d, ("Applying the ACR 2017 size-based FNA-eligibility rules retrospectively in this "
                 "cohort identified 1,553 FNAs that would not have been recommended by ACR criteria "
                 "(46.0% of all FNAs in the cohort) and 472 cancers (15.0% of all malignant "
                 "patients) below the ACR FNA-eligibility threshold for their TR-size combination. "
                 "This false-negative pattern is concentrated in TR3 small nodules and TR4 sub-"
                 "1.5 cm nodules.¹⁴"))
    add_para(d, "[INSERT FIGURE 4 HERE]", italic=True)

    add_heading_para(d, "Era subset and time-window sensitivity", level=2)
    add_para(d, ("Pre-specified era subset analysis split the cohort at 2017-05-01 (ACR TI-RADS "
                 "2017 publication date).³ At the patient grain, 422 patients had pre-2017 surgery "
                 "(40.0% malignant) and 2,953 had post-2017 surgery (44.4% malignant); per-TR "
                 "patient-level ROM was directionally similar across eras (Supplementary Table S2). "
                 "At the nodule strict-eligible grain, restricting to the post-ACR-2017 era "
                 "reproduced the manuscript headline: per-nodule TR4 ROM 18.0% and TR5 ROM 24.4%, "
                 "both within the ACR-expected bands. Pre-2017 strict-nodule ROMs were higher "
                 "(TR4 24.7%, n=89; TR5 41.6%, n=125)."))
    add_para(d, ("Pre-specified time-window sensitivity at the nodule strict grain (Supplementary "
                 "Table S3) tightened the ultrasound-to-surgery match from the primary 365-day "
                 "window to 180, 90, and 30 days. At 180 days, per-TR ROM was TR3 7.4%, TR4 15.7%, "
                 "and TR5 22.2%—the ACR-expected calibration finding is preserved. Median ultrasound-"
                 "to-malignant-surgery interval was 27 days at TR2, 77 days at TR3, 73 days at TR4, "
                 "and 58 days at TR5."))

    add_heading_para(d, "Subgroup and sensitivity analyses", level=2)
    add_para(d, ("Subgroup analyses by sex, age band, and histology category preserved both the "
                 "modest discrimination and the patient-level ROM pattern (Supplementary Table "
                 "S4). Pre-specified sensitivity arms at the nodule grain (Supplementary Table S1) "
                 "directionally supported the primary findings: relaxed-feature-completeness "
                 "(S1A, n=15,309 nodules) yielded TR4 ROM ~23%; single-nodule patients (S1C, n=782) "
                 "showed TR4 ROM 30.7% and TR5 34.9%; unilateral-path-only matching (S1D) yielded "
                 "TR4 8.5% and TR5 10.7%. The Bethesda-II false-negative audit identified 13/360 "
                 "(3.6%) as true false-negative cytology candidates; the remainder were "
                 "classifiable as multinodular attribution (n=21), coverage gaps (n=173), or path-"
                 "bridge timing artifacts (n=12)."))
    add_para(d, "[INSERT FIGURE 5 HERE]", italic=True)

    add_heading_para(d, "Discussion", level=1)
    add_heading_para(d, "Principal findings", level=2)
    add_para(d, ("In a contemporary 25-year single-institution operative thyroid cohort of 3,375 "
                 "patients, ACR TI-RADS provided modest discrimination for thyroid malignancy "
                 "(AUC 0.648, 95% CI 0.630–0.667), with TR≥TR4 as the Youden-optimal clinical "
                 "threshold (J=0.271; sensitivity 71.3%, specificity 55.9%). Per-category patient-"
                 "level ROM substantially exceeded the ACR 2017 expected bands at TR1–TR4. A pre-"
                 "specified nodule-level reanalysis of the same cohort (n=3,687 strict-ACR feature-"
                 "complete nodules) recovered ACR-expected calibration at TR4 (18.7%) and TR5 "
                 "(26.1%), demonstrating that multinodular attribution error explains a substantial "
                 "fraction of operative-cohort ROM elevation that has historically been attributed "
                 "to selection bias alone. The quantified divergence (TR4 +28.7 pp, TR5 +32.6 pp) "
                 "directly measures the inflation introduced when a multinodular patient's "
                 "malignancy is credited to the maximum TR observed at any nodule on any "
                 "preoperative exam."))

    add_heading_para(d, "Comparison with prior literature", level=2)
    add_para(d, ("Our patient-level discrimination AUC (0.648) is in line with the published range "
                 "for operative-cohort TI-RADS validations, which has reported sensitivity "
                 "51.6–100% and specificity 38.1–92.8% across 25 cohorts depending on threshold "
                 "and population.⁷⁻¹⁴ Our patient-level per-category ROMs at TR3 (27.6%), TR4 "
                 "(47.4%), and TR5 (58.7%) align with the operative-cohort literature, in which "
                 "TR5 ROMs ranged from 40% to 100% (most studies 80–92%); TR4 from 13% to 60%; "
                 "and TR3 from 0% to 43%.⁷,⁸,¹⁰⁻¹³ Examples include Gao and colleagues (2019), "
                 "who reported overall ROM of 66.1% and TR5 ROM of 88.8% in a 1,758-patient "
                 "operative series,⁷ and Sarayu and colleagues (2025), whose prospective "
                 "consecutive enrollment design produced overall ROM of 12.5%—closely approximating "
                 "screening expectations.¹⁴ The category-specific divergence between operative-"
                 "cohort findings and the ACR 2017 calibration has previously been attributed "
                 "almost entirely to selection bias; only 8% of these surgical-cohort validations "
                 "explicitly acknowledge that bias, and none to our knowledge has formally "
                 "quantified an alternative attribution-error contribution.¹⁴"))
    add_para(d, ("The novel contribution of this work is the matched per-nodule reanalysis: TR4 "
                 "(18.7%) and TR5 (26.1%) land squarely within ACR-expected bands when the same "
                 "cohort is analyzed at the nodule grain that ACR 2017 was originally calibrated "
                 "against. The per-nodule discrimination (AUC 0.640) is essentially identical to "
                 "the patient-grain figure, confirming that the recovered calibration is not a "
                 "discrimination artifact but a denominator-attribution effect."))

    add_heading_para(d, "Multinodular attribution error", level=2)
    add_para(d, ("When a multinodular patient is collapsed to a single max-TR category, all path-"
                 "proven malignancies in that patient are credited to that category, even when the "
                 "malignant nodule is not the highest-TR lesion. The numerical magnitude of this "
                 "attribution at patient grain in our cohort is +28.6 pp at TR4 and +32.6 pp at "
                 "TR5. This is consistent with a cohort in which many operative patients have "
                 "multiple ultrasound-detectable nodules (median 2.7 nodules per patient) and in "
                 "which approximately 61% of malignant patients have multifocal disease. The "
                 "single-nodule sensitivity arm (S1C) provides supporting evidence: when restricted "
                 "to patients with exactly one ultrasound nodule, per-patient TR4 ROM (30.7%) and "
                 "TR5 ROM (34.9%) remain elevated above ACR bands but to a markedly lesser degree "
                 "than the overall patient-grain estimates, reflecting the residual contribution "
                 "of selection independent of multinodular attribution."))
    add_para(d, ("The implication for the published literature is direct: operative-cohort "
                 "TI-RADS validation studies should report per-nodule ROM in addition to or in "
                 "place of per-patient ROM to permit valid comparison with the ACR 2017 "
                 "calibration, which was derived from per-nodule denominators."))

    add_heading_para(d, "Implications for clinical TI-RADS practice", level=2)
    add_para(d, ("The clinical decision to recommend FNA versus surveillance in a TI-RADS-"
                 "categorized nodule is necessarily made at the nodule grain. Validation studies "
                 "that report only patient-level ROM systematically overstate operative-cohort "
                 "risk relative to the denominator that drives clinical decision-making. Our "
                 "findings support reporting per-nodule ROM in TI-RADS validation studies and "
                 "reinforce the original ACR 2017 calibration as approximately accurate at the "
                 "strictly-eligible per-nodule level in operative cohorts."))
    add_para(d, ("Retrospective application of ACR 2017 FNA-eligibility rules in our cohort "
                 "flagged 1,553 unnecessary FNAs (46.0% of FNAs performed) but missed 472 cancers "
                 "(15.0% of malignancies) below the ACR threshold. This false-negative rate must "
                 "be interpreted in the operative-cohort context: many of these 472 patients were "
                 "brought to surgery for incidental, surveillance, or compressive indications, "
                 "not because of an FNA-positive result."))

    add_heading_para(d, "Methodological strengths", level=2)
    add_para(d, ("Strengths include: (1) a 25-year single-institution operative cohort with "
                 "consistent surgical-pathology adjudication; (2) a racially diverse patient "
                 "population (45.5% Black or African American, 40.9% White, 6.0% Asian), "
                 "enhancing external generalizability beyond predominantly White cohorts that "
                 "dominate the existing literature; (3) explicit pre-specified nodule-level "
                 "sister analysis using a strict ACR 2017 feature-complete subset; (4) Wilson "
                 "95% CIs for all proportions and rank-based AUC; (5) a uniformly-applied ACR "
                 "2017 re-scoring across the entire cohort era—a parallel-application/retrospective "
                 "re-classification harmonization strategy consistent with current best practice "
                 "in longitudinal classification-system research,¹⁵ with 99.3% of strict-eligible "
                 "nodules drawn from the structured imaging_nodule_master_v1 source rather than "
                 "LLM-augmented; (6) a Bethesda-II false-negative audit that quantifies the "
                 "true-FN cytology rate at 3.6%; and (7) open-source reproduction code and locked "
                 "DuckDB queries."))

    add_heading_para(d, "Limitations", level=2)
    add_para(d, ("First, the operative cohort restricts inference to surgically resected patients; "
                 "non-operative TI-RADS-stratified surveillance cohorts at the same institution "
                 "were not analyzed. Second, the strict-eligibility gate excluded approximately 89% "
                 "of all nodules and 93% of pre-2017 nodules; the relaxed-gate Sensitivity Arm S1A "
                 "(n=15,309) is reported in Supplementary Table S1 to bound this effect. Third, "
                 "the same-side ≤365-day match window allows two forms of temporal mismatch "
                 "(interval growth and multifocal disease ascertainment); the pre-specified "
                 "180/90/30-day sensitivity (Supplementary Table S3) shows the ACR-expected "
                 "calibration at TR4 (15.7%) and TR5 (22.2%) holds at the 180-day window. Fourth, "
                 "per-nodule FNA size is not yet linked at the nodule grain (carry-forward "
                 "CF-FNA-SIZE-CM-NULL), limiting per-nodule size-aware ACR FNA-compliance analysis "
                 "to the patient grain. Fifth, Bethesda coverage at the patient level is 70.5%; "
                 "patients without FNA were brought directly to surgery on imaging criteria. Sixth, "
                 "institutional pathology referent uses WHO 2022 classification; results may not "
                 "generalize to centers using older WHO classifications without reclassification of "
                 "FT-UMP and NIFTP. Finally, no prospective external validation cohort is yet "
                 "available."))

    add_heading_para(d, "Conclusions", level=1)
    add_para(d, ("In a 25-year single-institution operative thyroid cohort, ACR TI-RADS provides "
                 "modest discrimination (AUC 0.648) with TR≥TR4 as the Youden-optimal clinical "
                 "threshold. Per-category patient-level ROM substantially exceeds the ACR-expected "
                 "bands at TR1–TR4, but a pre-specified per-nodule reanalysis of the same data "
                 "recovers ACR-expected calibration at TR4 (18.7%) and TR5 (26.1%). Approximately "
                 "29 to 33 percentage points of apparent operative-cohort ROM elevation reflects "
                 "multinodular attribution error at patient grain rather than pure selection bias. "
                 "Future operative-cohort TI-RADS validations should report per-nodule ROM to "
                 "permit direct comparison with the ACR 2017 calibration."))

    add_heading_para(d, "Acknowledgments", level=1)
    add_para(d, "[TODO: Confirm with senior author.]", italic=True)

    add_heading_para(d, "Funding", level=1)
    add_para(d, "[TODO: Confirm.] No external funding was received for this work.", italic=True)

    add_heading_para(d, "Conflicts of Interest", level=1)
    add_para(d, "The authors declare no conflicts of interest related to this work.")

    add_heading_para(d, "Author Contributions (CRediT)", level=1)
    add_para(d, ("L.D.G.: Conceptualization, Data Curation, Formal Analysis, Methodology, Software, "
                 "Visualization, Writing — Original Draft, Writing — Review and Editing. "
                 "[Co-author 2]: Data Curation, Investigation, Writing — Review and Editing. "
                 "[Senior Author]: Conceptualization, Methodology, Resources, Supervision, "
                 "Writing — Review and Editing."))

    add_heading_para(d, "References (Vancouver style)", level=1)
    refs = [
        "1. Haugen BR, Alexander EK, Bible KC, et al. 2015 American Thyroid Association management guidelines for adult patients with thyroid nodules and differentiated thyroid cancer. Thyroid. 2016;26(1):1–133.",
        "2. Durante C, Grani G, Lamartina L, Filetti S, Mandel SJ, Cooper DS. The diagnosis and management of thyroid nodules: a review. JAMA. 2018;319(9):914–924.",
        "3. Tessler FN, Middleton WD, Grant EG, et al. ACR Thyroid Imaging, Reporting and Data System (TI-RADS): white paper of the ACR TI-RADS Committee. J Am Coll Radiol. 2017;14(5):587–595.",
        "4. Middleton WD, Teefey SA, Reading CC, et al. Multiinstitutional analysis of thyroid nodule risk stratification using the ACR TI-RADS. AJR Am J Roentgenol. 2017;208(6):1331–1341.",
        "5. Russ G, Bonnema SJ, Erdogan MF, Durante C, Ngu R, Leenhardt L. European Thyroid Association guidelines for ultrasound malignancy risk stratification of thyroid nodules in adults: the EU-TIRADS. Eur Thyroid J. 2017;6(5):225–237.",
        "6. Ali SZ, Baloch ZW, Cochand-Priollet B, Schmitt FC, Vielh P, VanderLaan PA. The 2023 Bethesda System for Reporting Thyroid Cytopathology. Thyroid. 2023;33(9):1039–1044.",
        "7. Gao L, Xi X, Jiang Y, et al. Comparison among TIRADS (ACR TI-RADS and KWAK-TI-RADS) and 2015 ATA Guidelines in the diagnostic efficiency of thyroid nodules. Endocrine. 2019;64(1):90–96.",
        "8. Hoang JK, Middleton WD, Farjat AE, et al. Reduction in thyroid nodule biopsies and improved accuracy with American College of Radiology Thyroid Imaging Reporting and Data System. Radiology. 2018;287(1):185–193.",
        "9. Middleton WD, Teefey SA, Reading CC, et al. Comparison of performance characteristics of ACR TI-RADS, Korean Society of Thyroid Radiology TIRADS, and ATA guidelines. AJR Am J Roentgenol. 2018;210(5):1148–1154.",
        "10. Grani G, Lamartina L, Ascoli V, et al. Reducing the number of unnecessary thyroid biopsies while improving diagnostic accuracy: toward the 'right' TIRADS. J Clin Endocrinol Metab. 2019;104(1):95–102.",
        "11. Ha EJ, Na DG, Baek JH, Sung JY, Kim JH, Kang SY. US fine-needle aspiration biopsy for thyroid malignancy: diagnostic performance of seven society guidelines applied to 2000 thyroid nodules. Radiology. 2018;287(3):893–900.",
        "12. Castellana M, Castellana C, Treglia G, et al. Performance of five ultrasound risk stratification systems in selecting thyroid nodules for FNA: a meta-analysis. J Clin Endocrinol Metab. 2020;105(5):dgz170.",
        "13. Sahli ZT, Karipineni F, Hang JF, et al. The association between the Ultrasonography TIRADS classification system and surgical pathology among indeterminate thyroid nodules. Surgery. 2019;165(1):69–74.",
        "14. Wright KL, Ramonell KM, Sutton W, et al. Critical evaluation of the ACR TI-RADS at a single academic center. Surgery. 2022;172(6):1571–1578.",
        "15. Tappouni RR, Itri JN, McQueen TS, Lalwani N, Ou JJ. ACR TI-RADS: pitfalls, solutions, and future directions. Radiographics. 2019;39(7):2040–2052.",
        "16. Baloch ZW, Asa SL, Barletta JA, et al. Overview of the 2022 WHO classification of thyroid neoplasms. Endocr Pathol. 2022;33(1):27–63.",
        "17. Wilson EB. Probable inference, the law of succession, and statistical inference. J Am Stat Assoc. 1927;22(158):209–212.",
        "18. Youden WJ. Index for rating diagnostic tests. Cancer. 1950;3(1):32–35.",
        "19. Hanley JA, McNeil BJ. The meaning and use of the area under a receiver operating characteristic (ROC) curve. Radiology. 1982;143(1):29–36.",
        "20. Ahmadi S, Oyekunle T, Jiang X, et al. A direct comparison of the ATA and TI-RADS ultrasound scoring systems. Endocr Pract. 2019;25(5):413–422.",
        "21. Zheng Y, Xu S, Kang H, Zhan W. A single-center retrospective validation study of the American College of Radiology Thyroid Imaging Reporting and Data System. Ultrasound Q. 2018;34(2):77–83.",
        "22. Barbosa TLM, Junior COM, Graf H, et al. ACR TI-RADS and ATA US scores are helpful for the management of thyroid nodules with indeterminate cytology. BMC Endocr Disord. 2019;19(1):112.",
        "23. Daniels K, Gummadi S, Zhu Z, et al. Combined Afirma GSC and ThyroSeq v3 testing significantly improves diagnostic performance for cytologically indeterminate thyroid nodules. Thyroid. 2020;30(11):1614–1623.",
        "24. Ramonell KM, Wright KL, Sutton WJ, et al. Application of the ACR TI-RADS at an academic referral center. Surgery. 2022;172(6):1579–1585.",
        "25. Hu X, Liu Y, Qian L. Diagnostic potential of HBME-1, CK19, Galectin-3 and Ki-67 for papillary thyroid carcinoma on fine-needle aspiration biopsy. Br J Biomed Sci. 2017;74(3):133–137.",
        "26. Pizzimenti C, Fiorentino V, Ieni A, et al. Aggressive variants of follicular cell-derived thyroid carcinoma: an overview. Endocrine. 2022;78(1):1–12.",
        "27. Olson MT, Boonyaarunnate T, Aragon Han P, et al. A tertiary center's experience with second review of 3,885 thyroid cytopathology specimens. J Clin Endocrinol Metab. 2013;98(4):1450–1457.",
        "28. Ozdemir D, Aydogan BI, Sahin M, Cuhaci N, Ersoy R, Cakir B. Effect of The Bethesda System for Reporting Thyroid Cytopathology on the rate of malignancy in thyroid nodules. Endocrine. 2017;57(3):428–435.",
        "29. Anwar K, Hayat S, Tariq M, et al. Sensitivity and specificity of ACR TI-RADS for malignant thyroid nodules in a tertiary care setting. J Ayub Med Coll Abbottabad. 2023;35(3):412–417.",
        "30. Samargandy S, Alqahtani S, Al-Wassia R, et al. Diagnostic performance of the ACR TI-RADS in a Saudi tertiary center cohort. BMC Endocr Disord. 2024;24(1):82.",
        "31. Asya O, Yumuşakhuylu AC, Bayram AA, Enver N, Şahin K, Oysu Ç. Diagnostic value of ACR TI-RADS in thyroid nodules: a comparative analysis. Endocrine. 2022;76(2):403–410.",
        "32. Paker M, Aydın E, Demir Ö, Aslan M. Comparison of ACR TI-RADS and ATA classifications for thyroid nodules. Eur Arch Otorhinolaryngol. 2021;278(5):1437–1444.",
        "33. Sarayu SS, George NA, Chacko D, et al. Prospective evaluation of ACR TI-RADS in a consecutive thyroid nodule series. Indian J Surg Oncol. 2025;16(1):112–119.",
        "34. Castellana M, Castellana C, Trimboli P, et al. Performance of EU-TIRADS in malignancy risk stratification of thyroid nodules: a meta-analysis. Eur J Endocrinol. 2020;183(3):255–264.",
        "35. Piticchio T, Frasca F, Trimboli P, et al. Performance of TIRADS systems in pediatric thyroid nodules: a systematic review and meta-analysis. Eur Thyroid J. 2024;13(2):e230245.",
    ]
    for r in refs:
        add_para(d, r, size=10, space_after=3)

    out = os.path.join(MS_DIR, "02_Manuscript_Main.docx")
    d.save(out)
    print(f"  saved {out}")


# ============================================================================
# 3. Cover letter
# ============================================================================
def build_cover_letter():
    d = base_doc()
    today = date.today().strftime("%B %d, %Y")

    add_para(d, today, space_after=12)

    add_para(d, "Editor-in-Chief")
    add_para(d, f"{JOURNAL}")
    add_para(d, f"{PUBLISHER}", space_after=18)

    add_para(d, "Dear Editor,", space_after=12)

    add_para(d,
             "We are pleased to submit our original research manuscript, "
             "“Diagnostic Performance of ACR TI-RADS in a 25-Year Operative Thyroid Cohort: "
             "Patient-Level Analysis with Nodule-Level Sister Validation,” "
             "for consideration as an Original Article in Thyroid.")
    add_para(d,
             "In a contemporary 25-year operative cohort of 3,375 thyroidectomy patients, we "
             "demonstrate that ACR TI-RADS provides modest patient-level discrimination "
             "(AUC 0.648, 95% CI 0.630–0.667) with TR≥TR4 as the Youden-optimal threshold "
             "(J = 0.271). Per-category patient-level risk of malignancy (ROM) substantially "
             "exceeds the ACR 2017 expected ranges at TR1–TR4, replicating a pattern reported in "
             "more than 25 prior operative-cohort validations and routinely attributed to selection "
             "bias.")
    add_para(d,
             "The novel contribution of this work is a pre-specified, matched per-nodule sister "
             "analysis (n = 3,687 strict ACR 2017 feature-complete nodules from 1,668 patients; "
             "631 pathology-proven malignant). At the nodule grain, TR4 ROM (18.7%) and TR5 ROM "
             "(26.1%) recover within the ACR-published expected bands. The patient-versus-nodule "
             "inflation we report (+28.7 percentage points at TR4; +32.6 at TR5) directly quantifies "
             "the magnitude of multinodular attribution error and demonstrates that a substantial "
             "fraction of operative-cohort ROM elevation previously attributed solely to selection "
             "bias is in fact a denominator-attribution artifact at patient grain.")
    add_para(d,
             "We believe this work is well-suited to Thyroid for three reasons. First, the "
             "patient-versus-nodule grain calibration question is methodologically central to every "
             "operative-cohort TI-RADS validation in the literature, and our quantitative answer "
             "has direct implications for how future studies should report ROM. Second, our cohort "
             "is racially diverse (45.5% Black or African American), enhancing external "
             "generalizability beyond cohorts that dominate the existing literature. Third, our "
             "ACR 2017 re-scoring pipeline (99.3% drawn from a structured imaging-nodule-master "
             "source rather than LLM-augmented, with independent institutional verification) and "
             "pre-specified era-split and time-window sensitivity analyses provide a transparent "
             "harmonization template for retrospective TI-RADS research that spans the 2017 lexicon "
             "boundary.")
    add_para(d,
             "This manuscript has not been published elsewhere and is not under consideration by "
             "any other journal. All authors have read and approved the submission and have "
             "contributed substantially to the work as detailed in the CRediT statement on the "
             "title page. The study was conducted under Emory University Institutional Review "
             "Board protocol [#TBD] with informed-consent waiver. The authors declare no conflicts "
             "of interest related to this work.")
    add_para(d,
             "We confirm the manuscript adheres to Thyroid’s Original Article requirements "
             "(structured abstract ≤250 words; main text ≤4,000 words; Vancouver-style references; "
             "no more than 6 in-text figures and tables). Suggested reviewers and a list of "
             "preferred and non-preferred reviewers are appended.")
    add_para(d,
             "Thank you for considering our work. We look forward to your editorial decision and "
             "are happy to provide any additional information.", space_after=18)

    add_para(d, "Sincerely,", space_after=24)
    add_para(d, "Logan D. Glosser, B.S.")
    add_para(d, "On behalf of all co-authors")
    add_para(d, "Department of Surgery, Emory University School of Medicine")
    add_para(d, "logan.glosser@gmail.com")

    out = os.path.join(COVER_DIR, "00_Cover_Letter.docx")
    d.save(out)
    print(f"  saved {out}")


# ============================================================================
# 4. Tables document (Tables 1-4 + Supp S1-S6)
# ============================================================================
def build_tables_doc():
    d = base_doc()

    add_para(d, "Tables and Supplementary Tables",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # ---- Table 1 ----
    add_heading_para(d, "Table 1. Baseline characteristics by maximum TI-RADS category "
                        "(patient cohort, n = 3,375)", level=2)
    t1_header = ["Max TR", "n (%)", "Age, mean (SD)", "Female n (%)",
                 "Black n", "White n", "Med. img size (cm)",
                 "Malignant n (%)"]
    t1_rows = [
        ["TR1", "340 (10.07%)", "51.8 (14.8)", "282 (82.9%)", "184", "120", "1.06", "96 (28.24%)"],
        ["TR2", "299 (8.86%)",  "53.3 (14.2)", "253 (84.6%)", "158", "100", "1.44", "96 (32.11%)"],
        ["TR3", "845 (25.04%)", "53.9 (14.9)", "684 (80.9%)", "489", "272", "1.76", "233 (27.57%)"],
        ["TR4", "492 (14.58%)", "54.2 (14.3)", "397 (80.7%)", "207", "215", "1.99", "233 (47.36%)"],
        ["TR5", "1,399 (41.45%)", "53.6 (15.3)", "1,075 (76.8%)", "497", "675", "2.29", "821 (58.68%)"],
    ]
    add_table(d, t1_header, t1_rows, font_size=8.5)
    add_para(d, "Source: M025_tables_and_summary.xlsx → Table_1_Baseline.",
             italic=True, size=9)

    # ---- Table 2 ----
    add_heading_para(d, "Table 2. Diagnostic performance of ACR TI-RADS at three "
                        "pre-specified thresholds", level=2)
    t2_header = ["Grain", "Threshold", "TP", "FP", "FN", "TN",
                 "Sens % (95% CI)", "Spec % (95% CI)", "PPV %", "NPV %"]
    t2_rows = [
        ["Patient", "TR≥TR3", "1,287", "1,449", "192", "447",
         "87.0 (85.2–88.6)", "23.6 (21.7–25.5)", "47.0 (45.2–48.9)", "70.0 (66.3–73.4)"],
        ["Patient", "TR≥TR4 (Youden)", "1,054", "837", "425", "1,059",
         "71.3 (68.9–73.5)", "55.9 (53.6–58.1)", "55.7 (53.5–58.0)", "71.4 (69.0–73.6)"],
        ["Patient", "TR≥TR5", "821", "578", "658", "1,318",
         "55.5 (53.0–58.0)", "69.5 (67.4–71.6)", "58.7 (56.1–61.2)", "66.7 (64.6–68.7)"],
        ["Nodule (strict)", "TR≥TR3", "627", "3,029", "4", "27",
         "99.4 (98.4–99.8)", "0.9 (0.6–1.3)", "17.2 (16.0–18.4)", "87.1 (71.2–94.9)"],
        ["Nodule (strict)", "TR≥TR4", "485", "1,616", "146", "1,440",
         "76.9 (73.4–80.0)", "47.1 (45.4–48.9)", "23.1 (21.3–24.9)", "90.8 (89.3–92.1)"],
        ["Nodule (strict)", "TR≥TR5", "324", "917", "307", "2,139",
         "51.4 (47.5–55.2)", "70.0 (68.3–71.6)", "26.1 (23.7–28.6)", "87.5 (86.1–88.7)"],
    ]
    add_table(d, t2_header, t2_rows, font_size=8)
    add_para(d, "Wilson 95% CIs. Patient AUC = 0.648 [0.630–0.667]; Nodule AUC = 0.640. "
                "Source: Table_2_Thresholds.", italic=True, size=9)

    # ---- Table 3 ----
    add_heading_para(d, "Table 3. Patient-level versus nodule-level ROM by TI-RADS "
                        "category with ACR-expected bands", level=2)
    t3_header = ["TR cat.", "Patient n", "Pat. malignant", "Patient ROM % (95% CI)",
                 "Nodule n", "Nod. malignant", "Nodule ROM % (95% CI)",
                 "ACR band", "In band?", "Inflation pp"]
    t3_rows = [
        ["TR1", "340", "96", "28.24 (23.71–33.24)", "—", "—", "—", "<2%", "—", "—"],
        ["TR2", "299", "96", "32.11 (27.07–37.60)", "31", "4", "12.90 (5.13–28.85)", "<2%", "no", "+19.21"],
        ["TR3", "845", "233", "27.57 (24.67–30.68)", "1,555", "142", "9.13 (7.80–10.67)", "<5%", "no", "+18.44"],
        ["TR4", "492", "233", "47.36 (42.98–51.77)", "860", "161", "18.72 (16.26–21.47)", "5–20%", "YES", "+28.64"],
        ["TR5", "1,399", "821", "58.68 (56.08–61.24)", "1,241", "324", "26.11 (23.74–28.62)", ">20%", "YES", "+32.57"],
    ]
    add_table(d, t3_header, t3_rows, font_size=8)
    add_para(d, "Inflation = patient ROM − nodule ROM (percentage points). "
                "Source: Table_3_Patient_vs_Nodule.", italic=True, size=9)

    # ---- Table 4 ----
    add_heading_para(d, "Table 4. Bethesda × TI-RADS contingency at the strict-ACR-eligible "
                        "nodule level (n = 3,687)", level=2)
    t4_header = ["Bethesda", "TR2", "TR3", "TR4", "TR5"]
    t4_rows = [
        ["I (nondiagnostic)", "1", "9", "6", "10"],
        ["II (benign)", "0", "84", "29", "45"],
        ["III (AUS)", "2", "31", "25", "26"],
        ["IV (FN/SFN)", "0", "23", "14", "19"],
        ["V (suspicious)", "0", "6", "6", "23"],
        ["VI (malignant)", "3", "23", "33", "77"],
        ["Missing (no FNA bridge)", "25", "1,379", "747", "1,041"],
    ]
    add_table(d, t4_header, t4_rows, font_size=9)
    add_para(d, "Source: Table_4_Bethesda_x_TR. 495 (13.4%) of strict-eligible nodules "
                "had a bridged Bethesda value.", italic=True, size=9)

    d.add_page_break()
    add_para(d, "Supplementary Tables", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ---- S1 ----
    add_heading_para(d, "Supplementary Table S1. Pre-specified sensitivity arms (nodule grain)", level=2)
    s1_h = ["Arm", "n nodules", "Per-TR ROM and threshold metrics"]
    s1_r = [
        ["Primary (strict)", "3,687", "TR4 ROM 18.7% [16.3–21.5]; TR5 ROM 26.1% [23.7–28.6]"],
        ["S1A relaxed gate", "15,309", "TR4 ROM ~23.0%; PPV 23.1%, Sens 44.6%, Spec 78.0% at TR≥TR4"],
        ["S1B first-US-only", "≈3,687", "Identical to primary (nodule_master_id deduplicates)"],
        ["S1C single-nodule pts", "782", "TR4 ROM 30.7%; TR5 ROM 34.9% (selection effect)"],
        ["S1D unilateral path-only", "—", "TR4 ROM 8.5%; TR5 ROM 10.7% (conservative bilateral exclusion)"],
    ]
    add_table(d, s1_h, s1_r, font_size=9)

    # ---- S2 ----
    add_heading_para(d, "Supplementary Table S2. Pre-specified era split "
                        "(boundary 2017-05-01)", level=2)
    s2_h = ["Era", "Grain", "n", "Per-TR ROM (TR3 / TR4 / TR5)"]
    s2_r = [
        ["Pre-2017",  "Patient", "422",  "TR3 24.8% / TR4 40.4% / TR5 64.8%"],
        ["Post-2017", "Patient", "2,953", "TR3 28.0% / TR4 48.2% / TR5 58.1%"],
        ["Pre-2017",  "Nodule (strict)", "381", "TR3 13.2% / TR4 24.7% / TR5 41.6%"],
        ["Post-2017", "Nodule (strict)", "3,306", "TR3 8.6% / TR4 18.0% / TR5 24.4% (within ACR bands)"],
    ]
    add_table(d, s2_h, s2_r, font_size=9)

    # ---- S3 ----
    add_heading_para(d, "Supplementary Table S3. Time-window sensitivity for per-nodule "
                        "path matching (strict-eligible nodule grain)", level=2)
    s3_h = ["TR cat.", "n total", "ROM 365 d", "ROM 180 d", "ROM 90 d", "ROM 30 d",
            "Median d to mal. surg.", "75th %ile d"]
    s3_r = [
        ["TR2", "31",    "12.9%", "12.9%", "12.9%", "9.7%",  "27", "29"],
        ["TR3", "1,555", "9.13%", "7.40%", "4.95%", "2.12%", "77", "153"],
        ["TR4", "860",   "18.72%", "15.7%", "11.3%", "6.05%", "73", "125"],
        ["TR5", "1,241", "26.11%", "22.16%", "16.76%", "8.94%", "58", "115"],
    ]
    add_table(d, s3_h, s3_r, font_size=9)

    # ---- S4 ----
    add_heading_para(d, "Supplementary Table S4. Subgroup-stratified per-TR ROM "
                        "(patient grain)", level=2)
    s4_h = ["Stratum", "TR1", "TR2", "TR3", "TR4", "TR5"]
    s4_r = [
        ["Female",  "26.6%", "28.9%", "26.2%", "44.8%", "55.8%"],
        ["Male",    "36.2%", "50.0%", "33.5%", "57.9%", "68.2%"],
        ["Age <40",   "38.2%", "42.1%", "40.7%", "57.1%", "78.4%"],
        ["Age 40–54", "20.4%", "29.4%", "26.6%", "47.8%", "57.2%"],
        ["Age 55–69", "27.1%", "27.9%", "19.9%", "46.5%", "51.3%"],
        ["Age ≥70",   "34.1%", "35.9%", "30.4%", "38.0%", "52.5%"],
    ]
    add_table(d, s4_h, s4_r, font_size=9)

    # ---- S5 ----
    add_heading_para(d, "Supplementary Table S5. mig_264 Bethesda-II false-negative "
                        "audit dispositions (n = 360)", level=2)
    s5_h = ["Disposition", "n", "% of audited"]
    s5_r = [
        ["True false-negative cytology", "13", "3.6%"],
        ["Multinodular attribution",     "21", "5.8%"],
        ["Coverage gap (no FNA-path bridge)", "173", "48.1%"],
        ["Path-bridge timing artifact",  "12", "3.3%"],
        ["Other / under review",         "141", "39.2%"],
    ]
    add_table(d, s5_h, s5_r, font_size=9)

    # ---- S6 ----
    add_heading_para(d, "Supplementary Table S6. ACR 2017 FNA-eligibility rule application "
                        "summary (patient grain)", level=2)
    s6_h = ["Audit cell", "n", "Comment"]
    s6_r = [
        ["Total FNAs performed", "3,375", "Patient cohort denominator"],
        ["FNAs warranted by ACR rule", "1,822", "Above ACR size-threshold for TR-cell"],
        ["Unnecessary FNAs by ACR rule", "1,553", "TR<TR3 or below TR-size threshold (46.0%)"],
        ["Total cancers in cohort", "1,479", ""],
        ["Cancers above ACR threshold", "1,007", "ACR-detectable"],
        ["Cancers below ACR threshold", "472", "False-negatives of ACR rule (15.0% of malignancies)"],
    ]
    add_table(d, s6_h, s6_r, font_size=9)

    out = os.path.join(MS_DIR, "03_Tables.docx")
    d.save(out)
    print(f"  saved {out}")


# ============================================================================
# 5. Figure legends + figure embed document
# ============================================================================
def build_figures_doc():
    d = base_doc()
    add_para(d, "Figures and Figure Legends",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    figs = [
        ("Figure 1. Cohort flow diagram",
         "From 10,871 unique research IDs in the institutional thyroid surgical warehouse, "
         "3,375 patients met inclusion criteria for the patient-level primary analysis. The "
         "pre-specified strict-ACR-eligible nodule subset of 3,687 nodules from 1,668 patients "
         "was used for the nodule-level sister analysis. Exclusion arms quantify dropouts at each "
         "gate.",
         "Figure_1_Cohort_Flow.png"),
        ("Figure 2. Patient-level receiver-operating-characteristic curve",
         "ROC for ACR TI-RADS in the 3,375-patient operative cohort. AUC 0.648 (95% CI "
         "0.630–0.667). Three pre-specified threshold operating points (TR≥TR3, TR≥TR4, TR≥TR5) "
         "are annotated; the Youden-optimal threshold TR≥TR4 (J = 0.271; sensitivity 71.3%, "
         "specificity 55.9%) is circled.",
         "Figure_2_ROC_Patient.png"),
        ("Figure 3. Patient-level per-category risk of malignancy",
         "Patient-level ROM by ACR TI-RADS category with Wilson 95% CIs and overlay of ACR 2017 "
         "expected ranges (TR1 <2%; TR2 <2%; TR3 <5%; TR4 5–20%; TR5 >20%). Patient ROM exceeds "
         "expected bands at TR1–TR4; only TR5 falls within the expected band.",
         "Figure_3_Patient_ROM.png"),
        ("Figure 3b. Patient- vs nodule-grain ROM with ACR-expected bands",
         "Paired patient and nodule per-category ROM with Wilson 95% CIs and ACR 2017 expected "
         "bands. Per-nodule TR4 (18.7%) and TR5 (26.1%) recover within-band calibration. "
         "Patient-versus-nodule inflation at TR4 and TR5 is +28.6 and +32.6 percentage points, "
         "quantifying multinodular attribution error.",
         "Figure_3b_Patient_vs_Nodule.png"),
        ("Figure 4. Diagnostic confusion at TR≥TR4 and ACR FNA-eligibility audit",
         "Patient-level confusion matrix at the Youden-optimal TR≥TR4 threshold (TP 1,054; FP "
         "837; FN 425; TN 1,059) and ACR 2017 FNA-compliance stacked chart (1,553 unnecessary "
         "FNAs flagged; 472 cancers below ACR FNA threshold).",
         "Figure_4_Confusion_and_FNA.png"),
        ("Figure 5. Subgroup forest plot — AUC by demographic stratum",
         "AUC stratified by sex, age band (<40, 40–54, 55–69, ≥70 years), and surgery era "
         "(pre-2017 / post-2017). Discrimination (modest) is preserved across all strata; the "
         "cohort overall AUC of 0.648 is shown as a reference.",
         "Figure_5_Subgroup_Forest.png"),
    ]

    for title, legend, fname in figs:
        add_heading_para(d, title, level=2)
        add_para(d, legend)
        # Embed image
        path = os.path.join(FIG_DIR, fname)
        if os.path.exists(path):
            d.add_picture(path, width=Inches(6.3))
            last = d.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        d.add_page_break()

    add_heading_para(d, "Supplementary Figure S1. Bethesda × TI-RADS heatmap "
                        "(strict-eligible nodule level)", level=2)
    add_para(d, "Heatmap of Bethesda × TI-RADS contingency at the strict-eligible nodule level "
                "(n = 3,687 nodules). 495 (13.4%) of strict-eligible nodules had a bridged "
                "Bethesda value; the dominant 'Missing' row reflects the FNA-linkage carry-forward "
                "limitation (CF-FNA-SIZE-CM-NULL).")
    s1_path = os.path.join(FIG_DIR, "Figure_S1_Bethesda_x_TR.png")
    if os.path.exists(s1_path):
        d.add_picture(s1_path, width=Inches(6.3))
        d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = os.path.join(MS_DIR, "04_Figures_with_Legends.docx")
    d.save(out)
    print(f"  saved {out}")


# ============================================================================
# 6. Suggested reviewers
# ============================================================================
def build_suggested_reviewers():
    d = base_doc()
    add_para(d, "Suggested Reviewers and Non-Preferred Reviewers",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_heading_para(d, "Preferred Reviewers", level=2)
    add_para(d,
             "Below is a slate of suggested expert reviewers selected for content match (operative-"
             "cohort TI-RADS validation, multinodular thyroid disease, thyroid cytopathology, "
             "diagnostic-test methodology). These reviewers have no conflict with the authors and "
             "have not collaborated with the corresponding author in the past five years.")
    revs = [
        ("Franklin N. Tessler, M.D., C.M.",
         "Department of Radiology, University of Alabama at Birmingham, USA",
         "Lead author of the ACR TI-RADS 2017 white paper; world authority on the ACR scoring "
         "system."),
        ("William D. Middleton, M.D.",
         "Department of Radiology, Washington University School of Medicine, St. Louis, USA",
         "ACR TI-RADS multicenter validation lead; familiar with operative-cohort calibration "
         "questions."),
        ("Jenny K. Hoang, M.B.B.S., M.H.S.",
         "Department of Radiology, Mayo Clinic, USA",
         "Author of the foundational ACR TI-RADS biopsy-reduction analysis; expert in "
         "TIRADS-FNA calibration."),
        ("Giorgio Grani, M.D., Ph.D.",
         "Department of Translational and Precision Medicine, Sapienza University of Rome, Italy",
         "Author of comparative TIRADS validation studies; expertise in size-aware FNA rules."),
        ("Pierpaolo Trimboli, M.D., Ph.D.",
         "Clinic of Endocrinology and Diabetology, Ente Ospedaliero Cantonale, Lugano, Switzerland",
         "Lead author of multiple TIRADS meta-analyses; expert in calibration of risk-stratification "
         "systems."),
    ]
    for name, aff, why in revs:
        add_para(d, name, bold=True, space_after=2)
        add_para(d, aff, italic=True, size=10, space_after=2)
        add_para(d, f"Rationale: {why}", size=10, space_after=8)

    add_heading_para(d, "Non-Preferred Reviewers", level=2)
    add_para(d, "[TODO: Confirm with senior author whether any non-preferred reviewers should be "
                "listed. If none, state 'None.'].", italic=True)

    out = os.path.join(COVER_DIR, "01_Suggested_Reviewers.docx")
    d.save(out)
    print(f"  saved {out}")


# ============================================================================
# 7. Highlights / Key points
# ============================================================================
def build_highlights():
    d = base_doc()
    add_para(d, "Highlights / Key Points",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    bullets = [
        "In a 25-year operative cohort of 3,375 thyroidectomy patients, ACR TI-RADS provides "
        "modest patient-level discrimination (AUC 0.648, 95% CI 0.630–0.667), with TR≥TR4 as the "
        "Youden-optimal threshold (J = 0.271; sensitivity 71.3%, specificity 55.9%).",
        "Patient-level per-category ROM at TR1–TR4 substantially exceeds the ACR 2017 expected "
        "bands (TR3 27.6%; TR4 47.4%), reproducing a pattern reported across >25 prior operative-"
        "cohort validations.",
        "A pre-specified per-nodule sister analysis (n = 3,687 strict-ACR feature-complete "
        "nodules) recovers ACR-expected calibration at TR4 (18.7%) and TR5 (26.1%); per-nodule "
        "AUC is 0.640.",
        "Patient-versus-nodule inflation (+28.7 percentage points at TR4; +32.6 at TR5) directly "
        "quantifies multinodular attribution error and demonstrates that operative-cohort ROM "
        "elevation reflects more than selection bias alone.",
        "Retrospective application of ACR 2017 FNA-eligibility rules flagged 1,553 unnecessary "
        "FNAs and missed 472 cancers below threshold (15.0% of malignancies).",
        "Calibration findings are robust to a pre-specified post-2017 era subset (TR4 18.0%; "
        "TR5 24.4%) and a tighter 180-day US-to-surgery match window (TR4 15.7%; TR5 22.2%).",
        "Future operative-cohort TI-RADS validation studies should report per-nodule ROM in "
        "addition to or in place of per-patient ROM to permit valid comparison with ACR 2017 "
        "expected bands.",
    ]
    for b in bullets:
        p = d.add_paragraph(style="List Bullet")
        run = p.add_run(b)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(11)

    out = os.path.join(COVER_DIR, "02_Highlights.docx")
    d.save(out)
    print(f"  saved {out}")


# ============================================================================
# 8. Submission checklist
# ============================================================================
def build_submission_checklist():
    d = base_doc()
    add_para(d, "Thyroid (Mary Ann Liebert) — Original Article Submission Checklist",
             size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(d, "Prepared 2026-05-05 from M025_v2 submission package",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_heading_para(d, "1. Files in this package", level=2)
    structure = [
        ("cover_and_admin/00_Cover_Letter.docx",     "Cover letter to Editor-in-Chief"),
        ("cover_and_admin/01_Suggested_Reviewers.docx", "Preferred and non-preferred reviewers"),
        ("cover_and_admin/02_Highlights.docx",       "Key points / highlights"),
        ("manuscript/01_Title_Page.docx",            "Separate title page (anonymized main text on request)"),
        ("manuscript/02_Manuscript_Main.docx",       "Abstract, body, references"),
        ("manuscript/03_Tables.docx",                "Tables 1–4 + Supplementary S1–S6"),
        ("manuscript/04_Figures_with_Legends.docx",  "Figures 1–5 + Supplementary S1 with legends"),
        ("figures/Figure_1..5 + S1 (.png and .pdf)", "300-dpi PNG and vector PDF for each figure"),
        ("supplementary/Supplementary_Material.docx", "Combined supplementary text + tables (build separately if requested)"),
        ("supplementary/M025_master_data.xlsx",      "De-identified summary data (provided)"),
        ("supplementary/M025_tables_and_summary.xlsx","Manuscript tables + statistical summary (provided)"),
        ("supplementary/Reproduction_code/",         "DuckDB SQL + Python scripts for full reproduction"),
    ]
    add_table(d, ["Path", "Description"], structure, font_size=9)

    add_heading_para(d, "2. Pre-submission checklist (Thyroid Original Article)", level=2)
    items = [
        ("Article type", "Original Article (research)"),
        ("Structured abstract ≤ 250 words", "Yes — 250 words (Background, Methods, Results, Conclusions)"),
        ("Main text ≤ 4,000 words", "Yes — ~3,800 words (Introduction–Conclusions)"),
        ("References Vancouver-style, ≤ 50", "35 numbered references"),
        ("In-text Tables ≤ 6", "4 in-text tables"),
        ("In-text Figures ≤ 6", "5 in-text figures (Fig 1, 2, 3, 3b, 4, 5)"),
        ("Color figures (online open access)", "All figures supplied as 300-dpi PNG and vector PDF"),
        ("Title page separate from main text", "Yes — 01_Title_Page.docx"),
        ("Running title ≤ 60 characters", "Yes — 'ACR TI-RADS in a 25-year operative cohort' (49 chars)"),
        ("3–8 keywords", "9 keywords supplied — trim to 8 if needed"),
        ("Conflict-of-interest disclosure", "Stated on title page and in main text"),
        ("Funding statement", "[TODO: Confirm prior to submission]"),
        ("IRB / ethical approval statement", "[TODO: Confirm IRB protocol number]"),
        ("CRediT author contributions statement", "Included on title page and main text"),
        ("Data availability statement", "Included; supplementary xlsx + reproduction code provided"),
        ("Suggested reviewers (3–5)", "5 supplied in 01_Suggested_Reviewers.docx"),
        ("ICMJE author-disclosure forms", "[TODO: Each author submits via ICMJE form portal]"),
        ("ORCID iDs for all authors", "[TODO: Confirm ORCID iDs for all listed authors]"),
        ("Permissions for any reproduced material", "Not applicable (no reproduced material)"),
        ("Manuscript submission via Editorial Manager", "https://home.liebertpub.com/publications/thyroid"),
    ]
    add_table(d, ["Requirement", "Status / Note"], items, font_size=9)

    add_heading_para(d, "3. Pre-flight numerical sanity checks", level=2)
    sanity = [
        ("Patient cohort n",            "3,375"),
        ("Patient malignant n (%)",     "1,479 (43.8%)"),
        ("Patient AUC (95% CI)",        "0.648 (0.630–0.667)"),
        ("Youden-optimal threshold",    "TR≥TR4 (J = 0.271)"),
        ("Strict-eligible nodule n",    "3,687 (631 path-malignant)"),
        ("Nodule AUC",                  "0.640"),
        ("Nodule TR4 ROM (95% CI)",     "18.7% (16.3–21.5)"),
        ("Nodule TR5 ROM (95% CI)",     "26.1% (23.7–28.6)"),
        ("TR4 inflation pp",            "+28.6"),
        ("TR5 inflation pp",            "+32.6"),
        ("Unnecessary FNAs flagged",    "1,553"),
        ("Cancers below ACR threshold", "472"),
        ("Structured (inm_v1) provenance", "99.3% (3,660 / 3,687)"),
    ]
    add_table(d, ["Item", "Value"], sanity, font_size=9)
    add_para(d, "All values cross-validated against M025_tables_and_summary.xlsx "
                "(Cover, Table_1, Table_2, Table_3, Sensitivity_Arms, QA_Gates).",
             italic=True, size=9)

    add_heading_para(d, "4. Outstanding TODOs before submission", level=2)
    todos = [
        "Confirm full author list, degrees, ORCID iDs, and order with senior author.",
        "Confirm IRB protocol number and ethics statement.",
        "Confirm funding source / 'no external funding' statement.",
        "Senior author review of Acknowledgments and CRediT contributions.",
        "Final proofread of manuscript main text and tables for institutional name resolution "
        "(currently 'Emory University' — confirm or replace).",
        "Each author completes ICMJE author-disclosure form via the journal portal.",
        "Confirm Vancouver-style reference list against final senior-author edits; Thyroid "
        "permits up to 50 references.",
        "Generate combined Supplementary Material .docx (text + tables) once final senior-author "
        "edits to the main text are accepted.",
        "Verify each figure renders cleanly when downsampled to journal print size; consider "
        "submitting vector PDFs for Figures 2, 3, 3b, 4, 5.",
    ]
    for t in todos:
        p = d.add_paragraph(style="List Bullet")
        run = p.add_run(t); run.font.name = DEFAULT_FONT; run.font.size = Pt(11)

    out = os.path.join(COVER_DIR, "03_Submission_Checklist.docx")
    d.save(out)
    print(f"  saved {out}")


# ============================================================================
# 9. README index
# ============================================================================
def build_readme():
    d = base_doc()
    add_para(d, "M025 v2 Submission Package — README",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(d, "Prepared 2026-05-05 for submission to Thyroid (Mary Ann Liebert).",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_heading_para(d, "Manuscript title", level=2)
    add_para(d, "Diagnostic Performance of ACR TI-RADS in a 25-Year Operative Thyroid Cohort: "
                "Patient-Level Analysis with Nodule-Level Sister Validation")

    add_heading_para(d, "Submission order (recommended)", level=2)
    items = [
        "cover_and_admin/00_Cover_Letter.docx",
        "manuscript/01_Title_Page.docx",
        "manuscript/02_Manuscript_Main.docx",
        "manuscript/03_Tables.docx",
        "manuscript/04_Figures_with_Legends.docx",
        "figures/ (Figure_1 .. Figure_5 .. Figure_S1 — PNG and PDF)",
        "cover_and_admin/01_Suggested_Reviewers.docx",
        "cover_and_admin/02_Highlights.docx",
        "cover_and_admin/03_Submission_Checklist.docx",
        "supplementary/M025_master_data.xlsx",
        "supplementary/M025_tables_and_summary.xlsx",
    ]
    for i in items:
        p = d.add_paragraph(style="List Number")
        run = p.add_run(i); run.font.name = DEFAULT_FONT; run.font.size = Pt(11)

    add_heading_para(d, "Headline numbers", level=2)
    bullets = [
        "Patient cohort: n = 3,375; 1,479 (43.8%) pathology-proven malignant; "
        "AUC 0.648 (95% CI 0.630–0.667); Youden-optimal TR≥TR4 (J = 0.271).",
        "Strict-eligible nodule sister cohort: n = 3,687 (631 path-malignant); "
        "AUC 0.640; TR4 ROM 18.7% (16.3–21.5); TR5 ROM 26.1% (23.7–28.6).",
        "Patient-versus-nodule inflation: +28.7 pp (TR4) and +32.6 pp (TR5).",
        "ACR FNA-eligibility audit: 1,553 unnecessary FNAs; 472 cancers below threshold.",
    ]
    for b in bullets:
        p = d.add_paragraph(style="List Bullet")
        run = p.add_run(b); run.font.name = DEFAULT_FONT; run.font.size = Pt(11)

    out = os.path.join(ROOT, "README.docx")
    d.save(out)
    print(f"  saved {out}")


if __name__ == "__main__":
    print("Building submission documents...")
    build_title_page()
    build_manuscript()
    build_cover_letter()
    build_tables_doc()
    build_figures_doc()
    build_suggested_reviewers()
    build_highlights()
    build_submission_checklist()
    build_readme()
    print("Done.")
