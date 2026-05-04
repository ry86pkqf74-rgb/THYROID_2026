#!/usr/bin/env python3
"""mig_295 — patch M044 02_manuscript.docx / 03_supplement.docx for v1.1 numbers.

Residual v1.0 tokens can appear without thousands separators (4128). Re-run safely:
  .venv/bin/python scripts/mig_295_apply_docx_patches.py --apply
"""
from __future__ import annotations

import argparse
import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parents[1]
PKG = REPO_ROOT / "M044_submission_package_v1_0"
MANUSCRIPT = PKG / "02_manuscript.docx"
SUPPLEMENT = PKG / "03_supplement.docx"
APPLY_LOG = REPO_ROOT / "scripts" / "output" / "mig_295_apply_log.txt"

# Ordered: longer / more specific first. Do not replace bare "4128" globally (SVG/coords risk in other assets).
NUMERIC_PATCHES: list[tuple[str, str]] = [
    ("4,128/4,128", "4,012/4,012"),
    ("(4128/4128),", "(4,012/4,012),"),
    ("(n = 3,789)", "(n = 3,750)"),
    ("(n = 3,756; events = 139)", "(n = 3,750; events = 193)"),
    ("(n = 3,756)", "(n = 3,750)"),
    ("Cox subset n = 2,025.", "Cox subset n = 2,511 (events = 178)."),
    ("(n=4128),", "(n=4,012),"),
    ("n=0/4128),", "n=0/4,012),"),
    ("4,128", "4,012"),
]

DISCUSSION_MEDIAN_FU_OLD = (
    "disproportionately in the microscopic stratum, diluting the time-to-event signal."
)
DISCUSSION_MEDIAN_FU_NEW = (
    "disproportionately in the microscopic stratum "
    "(median follow-up 3.2 years in the Cox-eligible subset), "
    "diluting the time-to-event signal."
)

BAD_TOKENS = [
    "4,128",
    "4128",
    "1.80",
    "2.34",
    "3,789",
    "3,756",
    "2,025",
    "139 events",
    "events = 139",
]


def _replace_paragraph_full_text(paragraph: Paragraph, new_text: str) -> None:
    p = paragraph._element
    for child in list(p):
        qname = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if qname == "r":
            p.remove(child)
    paragraph.add_run(new_text)


def apply_replacements(doc: Document, pairs: list[tuple[str, str]]) -> list[dict]:
    log: list[dict] = []
    for old, new in pairs:
        changed = 0
        for para in doc.paragraphs:
            if old not in para.text:
                continue
            merged = para.text.replace(old, new)
            _replace_paragraph_full_text(para, merged)
            changed += 1
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old not in para.text:
                            continue
                        merged = para.text.replace(old, new)
                        _replace_paragraph_full_text(para, merged)
                        changed += 1
        log.append({"old": old, "new": new, "paragraph_or_cell_hits": changed})
    return log


def blob_all_text(doc: Document) -> str:
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def residual_scan(paths: list[Path]) -> dict[str, list[str]]:
    out: dict[str, list[str]] = {}
    for pth in paths:
        if not pth.exists():
            continue
        doc = Document(pth)
        blob = blob_all_text(doc)
        hits = [t for t in BAD_TOKENS if t in blob]
        if hits:
            out[str(pth.relative_to(REPO_ROOT))] = hits
    return out


def patch_discussion_median_fu(doc: Document) -> bool:
    for para in doc.paragraphs:
        if DISCUSSION_MEDIAN_FU_OLD in para.text and DISCUSSION_MEDIAN_FU_NEW not in para.text:
            merged = para.text.replace(DISCUSSION_MEDIAN_FU_OLD, DISCUSSION_MEDIAN_FU_NEW)
            _replace_paragraph_full_text(para, merged)
            return True
    return False


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--apply", action="store_true", help="Write docx files")
    args = ap.parse_args()

    paths = [MANUSCRIPT, SUPPLEMENT]
    report: dict = {
        "applied_at_utc": datetime.now(timezone.utc).isoformat(),
        "dry_run": not args.apply,
        "files": [],
    }

    if not args.apply:
        print("Dry run — pass --apply to write.")
    for pth in paths:
        if not pth.exists():
            report["files"].append({"path": str(pth), "error": "missing"})
            continue
        doc = Document(pth)
        rep_log = apply_replacements(doc, NUMERIC_PATCHES)
        disc = False
        if pth == MANUSCRIPT:
            disc = patch_discussion_median_fu(doc)
        entry = {
            "file": str(pth.relative_to(REPO_ROOT)),
            "replacements": rep_log,
            "discussion_median_fu_patch": disc,
        }
        if args.apply:
            doc.save(pth)
        report["files"].append(entry)

    report["residual_scan"] = residual_scan(paths)

    APPLY_LOG.parent.mkdir(parents=True, exist_ok=True)
    APPLY_LOG.write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
